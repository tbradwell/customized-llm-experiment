"""Tests for document processor functionality."""

import pytest
import tempfile
import os
from pathlib import Path
from docx import Document

from src.core.document_processor import DocumentProcessor, PlaceholderInfo


class TestDocumentProcessor:
    """Test suite for DocumentProcessor class."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.processor = DocumentProcessor()
        self.test_data_dir = Path("test_data")
        self.test_data_dir.mkdir(exist_ok=True)
    
    def teardown_method(self):
        """Clean up test fixtures."""
        # Clean up any test files if needed
        pass
    
    def create_test_document(self, content_lines):
        """Helper to create a test document with given content."""
        doc = Document()
        for line in content_lines:
            doc.add_paragraph(line)
        return doc
    
    def test_find_placeholders_various_formats(self):
        """Test finding placeholders in different formats."""
        test_content = [
            "Contract between {{client_name}} and {{provider_name}}",
            "Amount: [CONTRACT_VALUE]",
            "Date: {start_date}",
            "Service: <service_type>"
        ]
        
        doc = self.create_test_document(test_content)
        placeholders = self.processor.find_placeholders(doc)
        
        # Should find all 4 placeholders
        assert len(placeholders) == 4
        
        # Check specific placeholders
        field_names = [p.field_name for p in placeholders]
        assert "client_name" in field_names
        assert "provider_name" in field_names
        assert "CONTRACT_VALUE" in field_names
        assert "start_date" in field_names
        assert "service_type" in field_names
    
    def test_find_placeholders_empty_document(self):
        """Test finding placeholders in empty document."""
        doc = Document()
        placeholders = self.processor.find_placeholders(doc)
        assert len(placeholders) == 0
    
    def test_find_placeholders_no_placeholders(self):
        """Test finding placeholders when none exist."""
        test_content = [
            "This is a regular contract.",
            "No placeholders here.",
            "Just normal text."
        ]
        
        doc = self.create_test_document(test_content)
        placeholders = self.processor.find_placeholders(doc)
        assert len(placeholders) == 0
    
    def test_replace_placeholders_basic(self):
        """Test basic placeholder replacement."""
        test_content = [
            "Client: {{client_name}}",
            "Provider: {{provider_name}}",
            "Value: {{contract_value}}"
        ]
        
        doc = self.create_test_document(test_content)
        
        data = {
            "client_name": "Test Client Corp",
            "provider_name": "Test Provider LLC", 
            "contract_value": "$50,000"
        }
        
        result_doc = self.processor.replace_placeholders(doc, data)
        text_content = self.processor.extract_text_content(result_doc)
        
        assert "Test Client Corp" in text_content
        assert "Test Provider LLC" in text_content
        assert "$50,000" in text_content
        assert "{{client_name}}" not in text_content
    
    def test_replace_placeholders_case_insensitive(self):
        """Test that placeholder replacement is case-insensitive."""
        test_content = ["Client: {{CLIENT_NAME}}"]
        doc = self.create_test_document(test_content)
        
        data = {"client_name": "Test Client"}  # lowercase key
        
        result_doc = self.processor.replace_placeholders(doc, data)
        text_content = self.processor.extract_text_content(result_doc)
        
        assert "Test Client" in text_content
        assert "{{CLIENT_NAME}}" not in text_content
    
    def test_replace_placeholders_missing_data(self):
        """Test replacement when some data is missing."""
        test_content = [
            "Client: {{client_name}}",
            "Provider: {{provider_name}}"
        ]
        
        doc = self.create_test_document(test_content)
        data = {"client_name": "Test Client"}  # Missing provider_name
        
        result_doc = self.processor.replace_placeholders(doc, data)
        text_content = self.processor.extract_text_content(result_doc)
        
        assert "Test Client" in text_content
        assert "{{provider_name}}" in text_content  # Should remain unchanged
    
    def test_create_document_copy(self):
        """Test creating a copy of a document."""
        # Create a temporary source document
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_source:
            doc = Document()
            doc.add_paragraph("Original document content")
            doc.save(tmp_source.name)
            source_path = tmp_source.name
        
        # Create a copy
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_copy:
            copy_path = tmp_copy.name
        
        try:
            copy_doc = self.processor.create_document_copy(source_path, copy_path)
            
            # Verify copy exists and has content
            assert os.path.exists(copy_path)
            copy_text = self.processor.extract_text_content(copy_doc)
            assert "Original document content" in copy_text
            
        finally:
            # Clean up
            if os.path.exists(source_path):
                os.unlink(source_path)
            if os.path.exists(copy_path):
                os.unlink(copy_path)
    
    def test_validate_document_structure(self):
        """Test document structure validation."""
        test_content = [
            "Contract Title",
            "Some content here",
            "More content"
        ]
        
        doc = self.create_test_document(test_content)
        validation_result = self.processor.validate_document_structure(doc)
        
        assert validation_result["is_valid"] is True
        assert validation_result["statistics"]["paragraph_count"] > 0
        assert len(validation_result["issues"]) == 0
    
    def test_validate_empty_document(self):
        """Test validation of empty document."""
        doc = Document()
        validation_result = self.processor.validate_document_structure(doc)
        
        assert validation_result["is_valid"] is False
        assert "no paragraphs" in " ".join(validation_result["issues"]).lower()
    
    def test_extract_text_content(self):
        """Test text content extraction."""
        test_content = [
            "First paragraph",
            "Second paragraph",
            "Third paragraph"
        ]
        
        doc = self.create_test_document(test_content)
        extracted_text = self.processor.extract_text_content(doc)
        
        for content in test_content:
            assert content in extracted_text
    
    def test_get_document_metadata(self):
        """Test document metadata extraction."""
        test_content = [
            "First paragraph",
            "Second paragraph"
        ]
        
        doc = self.create_test_document(test_content)
        metadata = self.processor.get_document_metadata(doc)
        
        assert metadata["paragraph_count"] == 2
        assert metadata["word_count"] > 0
        assert metadata["character_count"] > 0
        assert "has_headers" in metadata
        assert "has_footers" in metadata


@pytest.fixture
def sample_document_with_placeholders():
    """Fixture providing a document with various placeholder formats."""
    doc = Document()
    doc.add_paragraph("SERVICE AGREEMENT")
    doc.add_paragraph("Client: {{client_name}}")
    doc.add_paragraph("Provider: {{provider_name}}")
    doc.add_paragraph("Value: [CONTRACT_VALUE]")
    doc.add_paragraph("Start Date: {start_date}")
    doc.add_paragraph("Service Type: <service_description>")
    return doc


def test_comprehensive_placeholder_processing(sample_document_with_placeholders):
    """Integration test for complete placeholder processing."""
    processor = DocumentProcessor()
    
    # Find placeholders
    placeholders = processor.find_placeholders(sample_document_with_placeholders)
    assert len(placeholders) == 5
    
    # Replace placeholders
    data = {
        "client_name": "Acme Corporation",
        "provider_name": "Expert Services LLC",
        "CONTRACT_VALUE": "$100,000",
        "start_date": "2024-01-01",
        "service_description": "Professional consulting services"
    }
    
    result_doc = processor.replace_placeholders(sample_document_with_placeholders, data)
    final_text = processor.extract_text_content(result_doc)
    
    # Verify all replacements
    assert "Acme Corporation" in final_text
    assert "Expert Services LLC" in final_text
    assert "$100,000" in final_text
    assert "2024-01-01" in final_text
    assert "Professional consulting services" in final_text
    
    # Verify no placeholders remain
    assert "{{" not in final_text
    assert "}}" not in final_text
    assert "[" not in final_text
    assert "]" not in final_text
    assert "<service_description>" not in final_text