"""Tests for quality assurance pipeline functionality."""

import pytest
import tempfile
import os
from unittest.mock import Mock, patch, MagicMock
from pathlib import Path

from src.core.quality_pipeline import (
    QualityAssurancePipeline, QualityGateStatus, QualityGateResult, PipelineResult
)
from src.core.content_generator import GenerationContext, GenerationResult
from src.evaluation.metrics import EvaluationResult


class TestQualityAssurancePipeline:
    """Test suite for QualityAssurancePipeline class."""
    
    def setup_method(self):
        """Set up test fixtures."""
        with patch('src.core.quality_pipeline.settings') as mock_settings:
            # Mock settings to avoid dependency on actual configuration
            mock_settings.min_bleu_score = 0.8
            mock_settings.min_rouge_average = 0.85
            mock_settings.min_meteor_score = 0.9
            mock_settings.min_comet_score = 0.8
            mock_settings.min_llm_judge_score = 4.5
            mock_settings.max_redundancy_score = 0.1
            mock_settings.min_completeness_score = 0.98
            mock_settings.max_regeneration_attempts = 3
            mock_settings.quality_gate_enabled = True
            
            self.pipeline = QualityAssurancePipeline()
    
    def create_test_skeleton_file(self, content="Test contract with {{client_name}} and {{provider_name}}"):
        """Helper to create a test skeleton file."""
        from docx import Document
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            doc = Document()
            doc.add_paragraph(content)
            doc.save(tmp_file.name)
            return tmp_file.name
    
    def test_pre_generation_validation_valid_input(self):
        """Test pre-generation validation with valid input."""
        skeleton_path = self.create_test_skeleton_file()
        contract_data = {
            "client_name": "Test Client",
            "provider_name": "Test Provider",
            "contract_value": "$50,000"
        }
        checklist = ["Include payment terms", "Add confidentiality clause"]
        
        try:
            result = self.pipeline._pre_generation_validation(
                skeleton_path, contract_data, checklist
            )
            
            assert result["valid"] is True
            assert len(result["errors"]) == 0
        finally:
            os.unlink(skeleton_path)
    
    def test_pre_generation_validation_invalid_skeleton(self):
        """Test pre-generation validation with invalid skeleton."""
        result = self.pipeline._pre_generation_validation(
            "nonexistent_file.docx", {"client_name": "Test"}, []
        )
        
        assert result["valid"] is False
        assert len(result["errors"]) > 0
        assert "Invalid skeleton document" in result["errors"][0]
    
    def test_pre_generation_validation_empty_contract_data(self):
        """Test pre-generation validation with empty contract data."""
        skeleton_path = self.create_test_skeleton_file()
        
        try:
            result = self.pipeline._pre_generation_validation(
                skeleton_path, {}, []
            )
            
            assert result["valid"] is False
            assert "No contract data provided" in result["errors"]
        finally:
            os.unlink(skeleton_path)
    
    def test_pre_generation_validation_missing_required_fields(self):
        """Test validation with missing required fields."""
        skeleton_path = self.create_test_skeleton_file()
        contract_data = {"some_field": "value"}  # Missing client_name and provider_name
        
        try:
            result = self.pipeline._pre_generation_validation(
                skeleton_path, contract_data, []
            )
            
            # Should still be valid but with warnings
            assert result["valid"] is True
            assert len(result["warnings"]) >= 2  # Missing client_name and provider_name
        finally:
            os.unlink(skeleton_path)
    
    @patch('src.core.quality_pipeline.QualityAssurancePipeline._run_quality_gates')
    @patch('src.core.quality_pipeline.IntelligentContentGenerator.generate_contract_content')
    def test_iterative_generation_success_first_attempt(self, mock_generate, mock_quality_gates):
        """Test successful generation on first attempt."""
        skeleton_path = self.create_test_skeleton_file()
        
        # Mock successful generation
        mock_generate.return_value = GenerationResult(
            generated_content="Generated contract content",
            filled_placeholders={"{{client_name}}": "Test Client"},
            generation_metadata={"tokens": 100},
            warnings=[],
            success=True
        )
        
        # Mock passing quality gates
        quality_gate_result = QualityGateResult(
            gate_name="TEST",
            status=QualityGateStatus.PASSED,
            score=0.9,
            threshold=0.8,
            details={},
            recommendations=[]
        )
        
        mock_quality_gates.return_value = ([quality_gate_result], {"overall": 0.9})
        
        contract_data = {"client_name": "Test Client", "provider_name": "Test Provider"}
        
        try:
            result = self.pipeline._iterative_generation_with_quality_gates(
                skeleton_path, contract_data, [], []
            )
            
            assert result.success is True
            assert result.iterations == 1
            assert len(result.quality_gates) > 0
        finally:
            os.unlink(skeleton_path)
    
    @patch('src.core.quality_pipeline.QualityAssurancePipeline._run_quality_gates')
    @patch('src.core.quality_pipeline.IntelligentContentGenerator.generate_contract_content')
    def test_iterative_generation_max_iterations_reached(self, mock_generate, mock_quality_gates):
        """Test generation when max iterations is reached."""
        skeleton_path = self.create_test_skeleton_file()
        
        # Mock generation that always returns content
        mock_generate.return_value = GenerationResult(
            generated_content="Generated contract content",
            filled_placeholders={"{{client_name}}": "Test Client"},
            generation_metadata={"tokens": 100},
            warnings=[],
            success=True
        )
        
        # Mock failing quality gates
        quality_gate_result = QualityGateResult(
            gate_name="TEST",
            status=QualityGateStatus.FAILED,
            score=0.5,
            threshold=0.8,
            details={},
            recommendations=["Improve quality"]
        )
        
        mock_quality_gates.return_value = ([quality_gate_result], {"overall": 0.5})
        
        contract_data = {"client_name": "Test Client", "provider_name": "Test Provider"}
        
        try:
            result = self.pipeline._iterative_generation_with_quality_gates(
                skeleton_path, contract_data, [], []
            )
            
            assert result.success is False  # Failed to meet quality gates
            assert result.iterations == 3  # Max iterations reached
            assert "Quality gates not passed" in result.warnings[-1]
        finally:
            os.unlink(skeleton_path)
    
    def test_check_quality_gates_pass_all_passed(self):
        """Test quality gate checking when all gates pass."""
        gates = [
            QualityGateResult("GATE1", QualityGateStatus.PASSED, 0.9, 0.8, {}, []),
            QualityGateResult("GATE2", QualityGateStatus.PASSED, 0.85, 0.8, {}, [])
        ]
        
        result = self.pipeline._check_quality_gates_pass(gates)
        assert result is True
    
    def test_check_quality_gates_pass_some_failed(self):
        """Test quality gate checking when some gates fail."""
        gates = [
            QualityGateResult("GATE1", QualityGateStatus.PASSED, 0.9, 0.8, {}, []),
            QualityGateResult("GATE2", QualityGateStatus.FAILED, 0.7, 0.8, {}, [])
        ]
        
        result = self.pipeline._check_quality_gates_pass(gates)
        assert result is False
    
    def test_check_quality_gates_pass_empty_list(self):
        """Test quality gate checking with empty gate list."""
        result = self.pipeline._check_quality_gates_pass([])
        assert result is False
    
    def test_calculate_overall_quality_score(self):
        """Test overall quality score calculation."""
        scores = {
            "bleu": 0.8,
            "rouge": 0.85,
            "meteor": 0.9,
            "comet": 0.8,
            "redundancy": 0.05,  # Low redundancy is good
            "completeness": 0.98,
            "llm_judge": 4.5
        }
        
        overall_score = self.pipeline._calculate_overall_quality_score(scores)
        
        assert isinstance(overall_score, float)
        assert 0.0 <= overall_score <= 1.0
        assert overall_score > 0.8  # Should be high with good scores
    
    def test_calculate_overall_quality_score_empty(self):
        """Test overall quality score calculation with empty scores."""
        result = self.pipeline._calculate_overall_quality_score({})
        assert result == 0.0
    
    def test_prepare_improvement_feedback(self):
        """Test improvement feedback preparation."""
        failed_gates = [
            QualityGateResult(
                "BLEU", QualityGateStatus.FAILED, 0.7, 0.8, {}, 
                ["Score 0.7 below threshold 0.8"]
            ),
            QualityGateResult(
                "Completeness", QualityGateStatus.FAILED, 0.9, 0.98, {}, 
                ["Missing required elements"]
            )
        ]
        
        feedback = self.pipeline._prepare_improvement_feedback(failed_gates)
        
        assert len(feedback) == 2
        assert "bleu" in feedback[0].lower()
        assert "completeness" in feedback[1].lower()
    
    @patch('src.core.quality_pipeline.QualityAssurancePipeline._prepare_final_output')
    @patch('src.core.quality_pipeline.QualityAssurancePipeline._iterative_generation_with_quality_gates')
    @patch('src.core.quality_pipeline.QualityAssurancePipeline._pre_generation_validation')
    def test_process_contract_full_pipeline(self, mock_validation, mock_generation, mock_output):
        """Test full contract processing pipeline."""
        # Mock successful validation
        mock_validation.return_value = {"valid": True, "errors": [], "warnings": []}
        
        # Mock successful generation
        mock_generation.return_value = PipelineResult(
            success=True,
            final_contract="Generated contract text",
            quality_scores={"overall": 0.9},
            quality_gates=[],
            iterations=1,
            total_time=0,
            warnings=[],
            metadata={}
        )
        
        # Mock output preparation
        mock_output.return_value = "/path/to/output.docx"
        
        skeleton_path = self.create_test_skeleton_file()
        contract_data = {"client_name": "Test Client"}
        
        try:
            result = self.pipeline.process_contract(
                skeleton_path, contract_data, output_path="/output/path.docx"
            )
            
            assert result.success is True
            assert result.total_time > 0
            assert "output_path" in result.metadata
        finally:
            os.unlink(skeleton_path)
    
    @patch('src.evaluation.metrics.MetricsCalculator.calculate_bleu_score')
    @patch('src.evaluation.metrics.MetricsCalculator.calculate_rouge_scores')
    @patch('src.evaluation.llm_judge.LLMJudge.evaluate_contract')
    def test_run_quality_gates_comprehensive(self, mock_llm_judge, mock_rouge, mock_bleu):
        """Test running all quality gates."""
        # Mock evaluation results
        mock_bleu.return_value = EvaluationResult("BLEU", 0.85, {}, True, 0.8)
        mock_rouge.return_value = EvaluationResult("ROUGE", 0.87, {}, True, 0.85)
        mock_llm_judge.return_value = EvaluationResult("LLM_Judge", 4.6, {}, True, 4.5)
        
        contract_text = "High quality contract text with proper legal language."
        reference_contracts = ["Reference contract text."]
        context = GenerationContext(
            contract_type="service_agreement",
            skeleton_text="Template text",
            placeholders=["client_name"],
            contract_data={"client_name": "Test"}
        )
        
        gates, scores = self.pipeline._run_quality_gates(
            contract_text, reference_contracts, context
        )
        
        assert len(gates) > 0
        assert len(scores) > 0
        assert "bleu" in scores
        assert "rouge" in scores
        assert "llm_judge" in scores
    
    @patch('src.core.quality_pipeline.settings')
    def test_quality_gates_disabled(self, mock_settings):
        """Test pipeline behavior when quality gates are disabled."""
        mock_settings.quality_gate_enabled = False
        pipeline = QualityAssurancePipeline()
        
        contract_text = "Test contract text"
        context = Mock()
        
        gates, scores = pipeline._run_quality_gates(contract_text, [], context)
        
        assert len(gates) == 0
        assert len(scores) == 0


class TestQualityGateResult:
    """Test suite for QualityGateResult dataclass."""
    
    def test_quality_gate_result_creation(self):
        """Test QualityGateResult creation."""
        result = QualityGateResult(
            gate_name="TEST_GATE",
            status=QualityGateStatus.PASSED,
            score=0.85,
            threshold=0.8,
            details={"metric_details": "test"},
            recommendations=["Keep up the good work"]
        )
        
        assert result.gate_name == "TEST_GATE"
        assert result.status == QualityGateStatus.PASSED
        assert result.score == 0.85
        assert result.threshold == 0.8
        assert result.details["metric_details"] == "test"
        assert len(result.recommendations) == 1


class TestPipelineResult:
    """Test suite for PipelineResult dataclass."""
    
    def test_pipeline_result_creation(self):
        """Test PipelineResult creation."""
        result = PipelineResult(
            success=True,
            final_contract="Contract content",
            quality_scores={"overall": 0.9},
            quality_gates=[],
            iterations=2,
            total_time=5.5,
            warnings=["Minor warning"],
            metadata={"test": "data"}
        )
        
        assert result.success is True
        assert result.final_contract == "Contract content"
        assert result.quality_scores["overall"] == 0.9
        assert result.iterations == 2
        assert result.total_time == 5.5
        assert len(result.warnings) == 1
        assert result.metadata["test"] == "data"


def test_pipeline_integration_with_mocked_components():
    """Integration test with mocked components."""
    with patch('src.core.quality_pipeline.settings') as mock_settings:
        mock_settings.min_bleu_score = 0.5  # Lower thresholds for testing
        mock_settings.min_rouge_average = 0.5
        mock_settings.min_meteor_score = 0.5
        mock_settings.min_comet_score = 0.5
        mock_settings.min_llm_judge_score = 3.0
        mock_settings.max_redundancy_score = 0.5
        mock_settings.min_completeness_score = 0.5
        mock_settings.max_regeneration_attempts = 1
        mock_settings.quality_gate_enabled = True
        
        pipeline = QualityAssurancePipeline()
        
        # Create test skeleton
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Contract with {{client_name}} and {{provider_name}}")
            doc.save(tmp_file.name)
            skeleton_path = tmp_file.name
        
        try:
            with patch.object(pipeline.content_generator, 'generate_contract_content') as mock_gen:
                mock_gen.return_value = GenerationResult(
                    generated_content="Generated contract",
                    filled_placeholders={"{{client_name}}": "Test Client"},
                    generation_metadata={},
                    warnings=[],
                    success=True
                )
                
                result = pipeline.process_contract(
                    skeleton_path,
                    {"client_name": "Test Client", "provider_name": "Test Provider"}
                )
                
                assert isinstance(result, PipelineResult)
                assert result.iterations >= 1
        finally:
            os.unlink(skeleton_path)