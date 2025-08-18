"""Tests for API endpoints and functionality."""

import pytest
import json
import tempfile
import os
from unittest.mock import Mock, patch, AsyncMock
from fastapi.testclient import TestClient

# Import the FastAPI app
from src.api.main import app
from src.api.models import ContractGenerationRequest, EvaluationMetrics


class TestAPIEndpoints:
    """Test suite for API endpoints."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.client = TestClient(app)
    
    def test_root_endpoint(self):
        """Test root endpoint health check."""
        response = self.client.get("/")
        
        assert response.status_code == 200
        data = response.json()
        assert data["status"] == "healthy"
        assert "timestamp" in data
        assert "version" in data
        assert "services" in data
    
    def test_health_endpoint(self):
        """Test detailed health check endpoint."""
        response = self.client.get("/health")
        
        assert response.status_code == 200
        data = response.json()
        assert data["status"] == "healthy"
        assert "services" in data
        assert "document_processor" in data["services"]
        assert "quality_pipeline" in data["services"]
    
    def test_upload_skeleton_invalid_file_type(self):
        """Test skeleton upload with invalid file type."""
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp_file:
            tmp_file.write(b"Test content")
            tmp_file.flush()
            
            with open(tmp_file.name, "rb") as f:
                response = self.client.post(
                    "/upload/skeleton",
                    files={"file": ("test.txt", f, "text/plain")}
                )
        
        os.unlink(tmp_file.name)
        assert response.status_code == 400
        assert "Only .docx and .doc files are supported" in response.json()["detail"]
    
    @patch('src.api.main.document_processor')
    def test_upload_skeleton_valid_file(self, mock_processor):
        """Test skeleton upload with valid file."""
        # Mock the document processor
        mock_processor.find_placeholders.return_value = [
            Mock(field_name="client_name"),
            Mock(field_name="provider_name")
        ]
        mock_processor.validate_document_structure.return_value = {
            "is_valid": True,
            "issues": []
        }
        
        # Create a temporary .docx file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Test skeleton with {{client_name}}")
            doc.save(tmp_file.name)
            
            with open(tmp_file.name, "rb") as f:
                response = self.client.post(
                    "/upload/skeleton",
                    files={"file": ("skeleton.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
                )
        
        os.unlink(tmp_file.name)
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert "skeleton_id" in data
        assert "placeholders_found" in data
    
    @patch('src.api.main.document_processor')
    def test_upload_reference_contracts(self, mock_processor):
        """Test reference contracts upload."""
        mock_processor.extract_text_content.return_value = "Reference contract text content"
        
        # Create temporary .docx files
        files = []
        temp_files = []
        
        for i in range(2):
            tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            from docx import Document
            doc = Document()
            doc.add_paragraph(f"Reference contract {i+1}")
            doc.save(tmp_file.name)
            temp_files.append(tmp_file.name)
            
            with open(tmp_file.name, "rb") as f:
                files.append(("files", (f"ref{i+1}.docx", f.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")))
        
        try:
            response = self.client.post("/upload/references", files=files)
            
            assert response.status_code == 200
            data = response.json()
            assert data["success"] is True
            assert "reference_id" in data
            assert data["contracts_processed"] == 2
        finally:
            for temp_file in temp_files:
                os.unlink(temp_file)
    
    @patch('src.api.main.quality_pipeline')
    @patch('src.api.main.mlflow_tracker')
    def test_generate_contract_success(self, mock_mlflow, mock_pipeline):
        """Test successful contract generation."""
        # Setup mocks
        mock_mlflow.start_contract_generation_run.return_value = "run_123"
        mock_mlflow.log_generation_metrics.return_value = None
        mock_mlflow.end_run.return_value = None
        
        # Mock pipeline result
        from src.core.quality_pipeline import PipelineResult, QualityGateResult, QualityGateStatus
        
        mock_pipeline_result = PipelineResult(
            success=True,
            final_contract="Generated contract content",
            quality_scores={
                "overall": 4.7,
                "bleu": 0.85,
                "rouge": 0.87,
                "meteor": 0.92,
                "llm_judge": 4.6
            },
            quality_gates=[
                QualityGateResult("BLEU", QualityGateStatus.PASSED, 0.85, 0.8, {}, [])
            ],
            iterations=1,
            total_time=2.5,
            warnings=[],
            metadata={"generation_metadata": {}}
        )
        
        mock_pipeline.process_contract.return_value = mock_pipeline_result
        
        # First upload a skeleton
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Contract with {{client_name}}")
            doc.save(tmp_file.name)
            
            with patch('src.api.main.document_processor') as mock_doc_processor:
                mock_doc_processor.find_placeholders.return_value = [Mock(field_name="client_name")]
                mock_doc_processor.validate_document_structure.return_value = {"is_valid": True, "issues": []}
                
                with open(tmp_file.name, "rb") as f:
                    upload_response = self.client.post(
                        "/upload/skeleton",
                        files={"file": ("skeleton.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
                    )
        
        skeleton_id = upload_response.json()["skeleton_id"]
        
        # Generate contract
        contract_request = {
            "contract_data": {
                "client_name": "Test Client Corp",
                "provider_name": "Test Provider LLC",
                "contract_value": "$50,000"
            },
            "checklist": ["Include payment terms"],
            "quality_threshold": 4.5
        }
        
        response = self.client.post(
            f"/contracts/generate?skeleton_id={skeleton_id}",
            json=contract_request
        )
        
        os.unlink(tmp_file.name)
        
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert "contract_id" in data
        assert data["quality_score"] == 4.7
        assert data["meets_threshold"] is True
        assert "evaluation_metrics" in data
        assert "quality_gates" in data
    
    def test_generate_contract_missing_skeleton(self):
        """Test contract generation with missing skeleton."""
        contract_request = {
            "contract_data": {"client_name": "Test Client"},
            "quality_threshold": 4.5
        }
        
        response = self.client.post(
            "/contracts/generate?skeleton_id=nonexistent",
            json=contract_request
        )
        
        assert response.status_code == 404
        assert "Skeleton not found" in response.json()["detail"]
    
    @patch('src.api.main.quality_pipeline')
    def test_evaluate_contract(self, mock_pipeline):
        """Test contract evaluation endpoint."""
        mock_pipeline.evaluate_existing_contract.return_value = {
            "evaluation_results": {},
            "overall_score": 4.2
        }
        
        evaluation_request = {
            "contract_text": "Contract text to evaluate",
            "contract_context": {"contract_type": "service_agreement"}
        }
        
        response = self.client.post("/contracts/evaluate", json=evaluation_request)
        
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert "evaluation_metrics" in data
        assert "overall_assessment" in data
    
    def test_batch_generate_contracts(self):
        """Test batch contract generation."""
        batch_request = {
            "contracts": [
                {
                    "contract_data": {"client_name": "Client 1"},
                    "quality_threshold": 4.0
                },
                {
                    "contract_data": {"client_name": "Client 2"}, 
                    "quality_threshold": 4.0
                }
            ],
            "parallel_processing": False
        }
        
        response = self.client.post("/contracts/batch-generate", json=batch_request)
        
        assert response.status_code == 200
        data = response.json()
        assert data["total_contracts"] == 2
        assert data["successful_contracts"] >= 0
        assert "results" in data
        assert "summary" in data
    
    def test_download_contract_not_found(self):
        """Test downloading non-existent contract."""
        response = self.client.get("/contracts/download/nonexistent_id")
        
        assert response.status_code == 404
        assert "Contract not found" in response.json()["detail"]
    
    def test_quality_report_not_found(self):
        """Test quality report for non-existent contract."""
        response = self.client.get("/contracts/quality-report/nonexistent_id")
        
        assert response.status_code == 404
        assert "Contract not found" in response.json()["detail"]
    
    def test_api_error_handling(self):
        """Test API error handling."""
        # Test with invalid JSON
        response = self.client.post(
            "/contracts/evaluate",
            data="invalid json",
            headers={"Content-Type": "application/json"}
        )
        
        assert response.status_code == 422  # Validation error


class TestAPIModels:
    """Test suite for API models and validation."""
    
    def test_contract_generation_request_validation(self):
        """Test ContractGenerationRequest validation."""
        valid_data = {
            "contract_data": {
                "client_name": "Test Client",
                "provider_name": "Test Provider"
            },
            "quality_threshold": 4.5,
            "max_iterations": 3
        }
        
        request = ContractGenerationRequest(**valid_data)
        assert request.quality_threshold == 4.5
        assert request.max_iterations == 3
        assert request.contract_data["client_name"] == "Test Client"
    
    def test_contract_generation_request_invalid_threshold(self):
        """Test ContractGenerationRequest with invalid threshold."""
        invalid_data = {
            "contract_data": {"client_name": "Test"},
            "quality_threshold": 6.0  # Above maximum
        }
        
        with pytest.raises(ValueError):
            ContractGenerationRequest(**invalid_data)
    
    def test_evaluation_metrics_creation(self):
        """Test EvaluationMetrics model creation."""
        metrics = EvaluationMetrics(
            bleu_score=0.85,
            rouge_scores={"rouge1": 0.87, "rouge2": 0.83},
            meteor_score=0.91,
            llm_judge_score=4.6,
            overall_quality_score=4.5
        )
        
        assert metrics.bleu_score == 0.85
        assert metrics.rouge_scores["rouge1"] == 0.87
        assert metrics.overall_quality_score == 4.5


class TestAPIIntegration:
    """Integration tests for API functionality."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.client = TestClient(app)
    
    @patch('src.api.main.quality_pipeline')
    @patch('src.api.main.document_processor')
    def test_complete_contract_generation_workflow(self, mock_doc_processor, mock_pipeline):
        """Test complete workflow from skeleton upload to contract generation."""
        # Mock document processor
        mock_doc_processor.find_placeholders.return_value = [
            Mock(field_name="client_name"),
            Mock(field_name="provider_name")
        ]
        mock_doc_processor.validate_document_structure.return_value = {
            "is_valid": True,
            "issues": []
        }
        
        # Mock pipeline
        from src.core.quality_pipeline import PipelineResult, QualityGateStatus, QualityGateResult
        mock_pipeline.process_contract.return_value = PipelineResult(
            success=True,
            final_contract="Generated contract",
            quality_scores={"overall": 4.6},
            quality_gates=[QualityGateResult("TEST", QualityGateStatus.PASSED, 0.9, 0.8, {}, [])],
            iterations=1,
            total_time=2.0,
            warnings=[],
            metadata={}
        )
        
        # Step 1: Upload skeleton
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Contract with {{client_name}} and {{provider_name}}")
            doc.save(tmp_file.name)
            
            with open(tmp_file.name, "rb") as f:
                upload_response = self.client.post(
                    "/upload/skeleton",
                    files={"file": ("skeleton.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
                )
        
        assert upload_response.status_code == 200
        skeleton_id = upload_response.json()["skeleton_id"]
        
        # Step 2: Generate contract
        contract_request = {
            "contract_data": {
                "client_name": "Integration Test Client",
                "provider_name": "Integration Test Provider",
                "contract_value": "$75,000"
            },
            "checklist": ["Include payment terms", "Add confidentiality clause"],
            "quality_threshold": 4.5
        }
        
        generation_response = self.client.post(
            f"/contracts/generate?skeleton_id={skeleton_id}",
            json=contract_request
        )
        
        assert generation_response.status_code == 200
        generation_data = generation_response.json()
        assert generation_data["success"] is True
        contract_id = generation_data["contract_id"]
        
        # Step 3: Get quality report
        report_response = self.client.get(f"/contracts/quality-report/{contract_id}")
        assert report_response.status_code == 200
        report_data = report_response.json()
        assert "overall_quality_score" in report_data
        
        # Cleanup
        os.unlink(tmp_file.name)
    
    def test_api_documentation_accessible(self):
        """Test that API documentation is accessible."""
        response = self.client.get("/docs")
        assert response.status_code == 200
        
        response = self.client.get("/redoc")
        assert response.status_code == 200
    
    def test_openapi_schema(self):
        """Test OpenAPI schema generation."""
        response = self.client.get("/openapi.json")
        assert response.status_code == 200
        
        schema = response.json()
        assert "openapi" in schema
        assert "info" in schema
        assert "paths" in schema
        
        # Check that our main endpoints are documented
        assert "/contracts/generate" in schema["paths"]
        assert "/contracts/evaluate" in schema["paths"]
        assert "/upload/skeleton" in schema["paths"]