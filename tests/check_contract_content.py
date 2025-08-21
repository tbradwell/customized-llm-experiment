#!/usr/bin/env python3
"""Check contract content for השלם placeholders."""

from docx import Document
import re

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text.strip())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text.strip())
    
    return '\n'.join(full_text)

def check_content():
    """Check contract content."""
    contract_path = "experiments/amit_example_based_1755438792/outputs/example_based_contract.docx"
    
    try:
        text = extract_text_from_docx(contract_path)
        
        print(f"📄 Contract text length: {len(text)} characters")
        print(f"📋 First 1000 characters:")
        print("-" * 50)
        print(text[:1000])
        print("-" * 50)
        
        # Check for השלם instances
        haslem_count = text.count("השלם")
        print(f"\n🔍 Found {haslem_count} instances of 'השלם'")
        
        if haslem_count > 0:
            haslem_contexts = []
            lines = text.split('\n')
            for i, line in enumerate(lines):
                if 'השלם' in line:
                    haslem_contexts.append(f"Line {i+1}: {line.strip()}")
            
            print("📍 השלם contexts:")
            for context in haslem_contexts[:5]:
                print(f"  • {context}")
            if len(haslem_contexts) > 5:
                print(f"  ... and {len(haslem_contexts) - 5} more")
        
        # Check for placeholders
        placeholders = re.findall(r'\{[^}]*\}', text)
        print(f"\n🔧 Found {len(placeholders)} remaining placeholders: {placeholders[:3]}")
        
        # Check for real data indicators
        real_data = ['זואי', 'מור', 'הוד השרון', 'חיים הרצוג', '0542477683', 'יצחק אולשבנג']
        found_data = [data for data in real_data if data in text]
        print(f"\n✅ Found {len(found_data)} real data indicators: {found_data}")
        
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    check_content()