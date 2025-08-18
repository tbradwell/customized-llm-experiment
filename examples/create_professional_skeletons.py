#!/usr/bin/env python3
"""
Create professional .docx skeleton files with proper formatting,
styles, and structure for the Lawyer Contract Creation System.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def create_service_agreement_skeleton():
    """Create a professional service agreement skeleton with proper formatting."""
    
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('Contract Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(16)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading style
    heading_style = styles.add_style('Contract Heading', WD_STYLE_TYPE.PARAGRAPH)
    heading_font = heading_style.font
    heading_font.name = 'Arial'
    heading_font.size = Pt(12)
    heading_font.bold = True
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)
    
    # Body style
    body_style = styles.add_style('Contract Body', WD_STYLE_TYPE.PARAGRAPH)
    body_font = body_style.font
    body_font.name = 'Times New Roman'
    body_font.size = Pt(11)
    body_style.paragraph_format.space_after = Pt(6)
    body_style.paragraph_format.line_spacing = 1.15
    
    # Document title
    title = doc.add_paragraph('SERVICE AGREEMENT', style='Contract Title')
    
    # Opening paragraph
    opening = doc.add_paragraph(style='Contract Body')
    opening.add_run('This Service Agreement ("Agreement") is entered into on ')
    opening.add_run('{{contract_date}}').bold = True
    opening.add_run(' between:')
    
    # Add some spacing
    doc.add_paragraph()
    
    # Client information
    client_heading = doc.add_paragraph('CLIENT INFORMATION:', style='Contract Heading')
    
    client_para = doc.add_paragraph(style='Contract Body')
    client_para.add_run('Company: ')
    client_para.add_run('{{client_name}}').bold = True
    
    address_para = doc.add_paragraph(style='Contract Body')
    address_para.add_run('Address: ')
    address_para.add_run('{{client_address}}').bold = True
    
    contact_para = doc.add_paragraph(style='Contract Body')
    contact_para.add_run('Contact: ')
    contact_para.add_run('{{client_contact}}').bold = True
    
    doc.add_paragraph()
    
    # Provider information
    provider_heading = doc.add_paragraph('PROVIDER INFORMATION:', style='Contract Heading')
    
    provider_para = doc.add_paragraph(style='Contract Body')
    provider_para.add_run('Company: ')
    provider_para.add_run('{{provider_name}}').bold = True
    
    provider_address_para = doc.add_paragraph(style='Contract Body')
    provider_address_para.add_run('Address: ')
    provider_address_para.add_run('{{provider_address}}').bold = True
    
    provider_contact_para = doc.add_paragraph(style='Contract Body')
    provider_contact_para.add_run('Contact: ')
    provider_contact_para.add_run('{{provider_contact}}').bold = True
    
    doc.add_paragraph()
    
    # Contract sections
    sections = [
        {
            'title': '1. SCOPE OF SERVICES',
            'content': '{{service_description}}\n\nThe Provider agrees to perform the following services:\n{{detailed_services}}'
        },
        {
            'title': '2. COMPENSATION AND PAYMENT TERMS',
            'content': 'Total Contract Value: {{contract_value}}\nPayment Schedule: {{payment_terms}}\nPayment Method: {{payment_method}}\n\nLate payments may incur additional charges as specified herein.'
        },
        {
            'title': '3. TERM AND PERFORMANCE TIMELINE',
            'content': 'Start Date: {{start_date}}\nEnd Date: {{end_date}}\nKey Milestones: {{milestones}}\n\nTime is of the essence in the performance of this Agreement.'
        },
        {
            'title': '4. CONFIDENTIALITY PROVISIONS',
            'content': '{{confidentiality_clause}}\n\nThis confidentiality obligation shall survive the termination of this Agreement.'
        },
        {
            'title': '5. TERMINATION CONDITIONS',
            'content': '{{termination_conditions}}\n\nUpon termination, all outstanding payments shall become immediately due and payable.'
        },
        {
            'title': '6. GENERAL PROVISIONS',
            'content': '{{general_provisions}}\n\nThis Agreement constitutes the entire agreement between the parties and supersedes all prior negotiations, representations, or agreements relating to the subject matter hereof.'
        }
    ]
    
    for section in sections:
        # Section heading
        doc.add_paragraph(section['title'], style='Contract Heading')
        
        # Section content
        for paragraph in section['content'].split('\n\n'):
            para = doc.add_paragraph(paragraph.strip(), style='Contract Body')
    
    # Signature section
    doc.add_paragraph()
    doc.add_paragraph('SIGNATURES:', style='Contract Heading')
    
    # Create signature table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Client signature
    table.cell(0, 0).text = 'CLIENT:'
    table.cell(1, 0).text = 'Signature: _________________________'
    table.cell(2, 0).text = 'Name: {{client_name}}\nDate: _____________'
    
    # Provider signature  
    table.cell(0, 1).text = 'PROVIDER:'
    table.cell(1, 1).text = 'Signature: _________________________'
    table.cell(2, 1).text = 'Name: {{provider_name}}\nDate: _____________'
    
    return doc


def create_nda_skeleton():
    """Create a professional NDA skeleton with proper formatting."""
    
    doc = Document()
    
    # Use same styles as service agreement
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('NDA Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(16)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading style
    heading_style = styles.add_style('NDA Heading', WD_STYLE_TYPE.PARAGRAPH)
    heading_font = heading_style.font
    heading_font.name = 'Arial'
    heading_font.size = Pt(12)
    heading_font.bold = True
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)
    
    # Body style
    body_style = styles.add_style('NDA Body', WD_STYLE_TYPE.PARAGRAPH)
    body_font = body_style.font
    body_font.name = 'Times New Roman'
    body_font.size = Pt(11)
    body_style.paragraph_format.space_after = Pt(6)
    body_style.paragraph_format.line_spacing = 1.15
    
    # Document title
    title = doc.add_paragraph('NON-DISCLOSURE AGREEMENT', style='NDA Title')
    
    # Opening paragraph
    opening = doc.add_paragraph(style='NDA Body')
    opening.add_run('This Non-Disclosure Agreement ("Agreement") is made on ')
    opening.add_run('{{agreement_date}}').bold = True
    opening.add_run(' between:')
    
    doc.add_paragraph()
    
    # Parties section
    parties_heading = doc.add_paragraph('PARTIES:', style='NDA Heading')
    
    disclosing_para = doc.add_paragraph(style='NDA Body')
    disclosing_para.add_run('DISCLOSING PARTY: ')
    disclosing_para.add_run('{{disclosing_party_name}}').bold = True
    disclosing_para.add_run('\nAddress: ')
    disclosing_para.add_run('{{disclosing_party_address}}').bold = True
    disclosing_para.add_run('\nContact: ')
    disclosing_para.add_run('{{disclosing_party_contact}}').bold = True
    
    doc.add_paragraph()
    
    receiving_para = doc.add_paragraph(style='NDA Body')
    receiving_para.add_run('RECEIVING PARTY: ')
    receiving_para.add_run('{{receiving_party_name}}').bold = True
    receiving_para.add_run('\nAddress: ')
    receiving_para.add_run('{{receiving_party_address}}').bold = True
    receiving_para.add_run('\nContact: ')
    receiving_para.add_run('{{receiving_party_contact}}').bold = True
    
    # NDA sections
    sections = [
        {
            'title': '1. DEFINITION OF CONFIDENTIAL INFORMATION',
            'content': '{{confidential_info_definition}}\n\nConfidential Information shall not include information that is publicly available or independently developed.'
        },
        {
            'title': '2. OBLIGATIONS OF RECEIVING PARTY',
            'content': '{{receiving_party_obligations}}\n\nThe Receiving Party agrees to use the same degree of care to protect Confidential Information as it uses to protect its own confidential information, but in no event less than reasonable care.'
        },
        {
            'title': '3. EXCEPTIONS TO CONFIDENTIALITY',
            'content': 'The obligations of confidentiality shall not apply to information that:\n\n{{confidentiality_exceptions}}'
        },
        {
            'title': '4. DURATION OF AGREEMENT',
            'content': '{{agreement_duration}}\n\nThe obligations of confidentiality shall survive termination of this Agreement.'
        },
        {
            'title': '5. REMEDIES FOR BREACH',
            'content': '{{breach_remedies}}\n\nThe Receiving Party acknowledges that breach of this Agreement would cause irreparable harm for which monetary damages would be inadequate.'
        },
        {
            'title': '6. RETURN OF MATERIALS',
            'content': '{{return_of_materials}}\n\nUpon termination or upon request, all materials containing Confidential Information shall be returned or destroyed.'
        },
        {
            'title': '7. MISCELLANEOUS PROVISIONS',
            'content': '{{miscellaneous_provisions}}\n\nThis Agreement shall be governed by applicable state law and any disputes shall be resolved through binding arbitration.'
        }
    ]
    
    for section in sections:
        # Section heading
        doc.add_paragraph(section['title'], style='NDA Heading')
        
        # Section content
        for paragraph in section['content'].split('\n\n'):
            para = doc.add_paragraph(paragraph.strip(), style='NDA Body')
    
    # Signature section
    doc.add_paragraph()
    doc.add_paragraph('SIGNATURES:', style='NDA Heading')
    
    # Create signature table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Disclosing party signature
    table.cell(0, 0).text = 'DISCLOSING PARTY:'
    table.cell(1, 0).text = 'Signature: _________________________'
    table.cell(2, 0).text = 'Name: {{disclosing_party_name}}\nDate: _____________'
    
    # Receiving party signature
    table.cell(0, 1).text = 'RECEIVING PARTY:'
    table.cell(1, 1).text = 'Signature: _________________________'
    table.cell(2, 1).text = 'Name: {{receiving_party_name}}\nDate: _____________'
    
    return doc


def create_consulting_agreement_skeleton():
    """Create a professional consulting agreement skeleton."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('CONSULTING AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Opening
    opening = doc.add_paragraph()
    opening.add_run('This Consulting Agreement ("Agreement") is entered into on ')
    opening.add_run('{{agreement_date}}').bold = True
    opening.add_run(' between ')
    opening.add_run('{{client_name}}').bold = True
    opening.add_run(' ("Client") and ')
    opening.add_run('{{consultant_name}}').bold = True
    opening.add_run(' ("Consultant").')
    
    # Sections
    sections = [
        ('1. CONSULTING SERVICES', '{{service_description}}\n\nConsultant shall provide the services with professional competence and in accordance with industry standards.'),
        ('2. COMPENSATION', '{{compensation_terms}}\n\nPayment shall be made within thirty (30) days of receipt of invoice.'),
        ('3. INTELLECTUAL PROPERTY', '{{ip_provisions}}\n\nConsultant hereby assigns to Client all rights in work product created specifically for this engagement.'),
        ('4. TERM AND TERMINATION', '{{contract_duration}}\n\nEither party may terminate this Agreement with thirty (30) days written notice.'),
        ('5. CONFIDENTIALITY', '{{confidentiality_provisions}}\n\nConsultant acknowledges access to confidential information and agrees to maintain strict confidentiality.'),
        ('6. INDEPENDENT CONTRACTOR', 'Consultant is an independent contractor and not an employee of Client. Consultant is responsible for all taxes and benefits.')
    ]
    
    for title, content in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
    
    # Signatures
    doc.add_heading('SIGNATURES', level=1)
    
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Table Grid'
    
    sig_table.cell(0, 0).text = 'CLIENT:'
    sig_table.cell(1, 0).text = 'Signature: _________________________'
    sig_table.cell(2, 0).text = 'Name: {{client_name}}\nDate: _____________'
    
    sig_table.cell(0, 1).text = 'CONSULTANT:'
    sig_table.cell(1, 1).text = 'Signature: _________________________'
    sig_table.cell(2, 1).text = 'Name: {{consultant_name}}\nDate: _____________'
    
    return doc


def main():
    """Create all professional skeleton documents."""
    
    print("🏗️ Creating Professional .docx Skeleton Files")
    print("="*60)
    
    # Ensure directories exist
    skeletons_dir = Path("data/skeletons")
    skeletons_dir.mkdir(parents=True, exist_ok=True)
    
    # Create service agreement skeleton
    try:
        service_doc = create_service_agreement_skeleton()
        service_path = skeletons_dir / "professional_service_agreement_skeleton.docx"
        service_doc.save(str(service_path))
        logger.info(f"✅ Created professional service agreement skeleton: {service_path}")
    except Exception as e:
        logger.error(f"❌ Error creating service agreement: {e}")
    
    # Create NDA skeleton
    try:
        nda_doc = create_nda_skeleton()
        nda_path = skeletons_dir / "professional_nda_skeleton.docx"
        nda_doc.save(str(nda_path))
        logger.info(f"✅ Created professional NDA skeleton: {nda_path}")
    except Exception as e:
        logger.error(f"❌ Error creating NDA: {e}")
    
    # Create consulting agreement skeleton
    try:
        consulting_doc = create_consulting_agreement_skeleton()
        consulting_path = skeletons_dir / "professional_consulting_skeleton.docx"
        consulting_doc.save(str(consulting_path))
        logger.info(f"✅ Created professional consulting skeleton: {consulting_path}")
    except Exception as e:
        logger.error(f"❌ Error creating consulting agreement: {e}")
    
    print("\n🎉 Professional skeleton creation completed!")
    print("\n📁 Created files:")
    print("• professional_service_agreement_skeleton.docx - Full-featured service agreement")
    print("• professional_nda_skeleton.docx - Comprehensive NDA template") 
    print("• professional_consulting_skeleton.docx - Professional consulting agreement")
    
    print("\n💡 These skeletons include:")
    print("• Professional formatting and styles")
    print("• Proper document structure")
    print("• Signature tables")
    print("• Comprehensive placeholder coverage")
    print("• Legal document conventions")


if __name__ == "__main__":
    main()