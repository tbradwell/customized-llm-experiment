#!/usr/bin/env python3
"""
Script to create sample .docx skeleton files from text templates.
This converts the text skeletons to proper .docx format for testing.
"""

from docx import Document
from pathlib import Path
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def create_docx_from_text(text_file_path: str, output_path: str):
    """Create a .docx file from a text template."""
    try:
        # Read the text content
        with open(text_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create a new Document
        doc = Document()
        
        # Add title with styling
        lines = content.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # First non-empty line is typically the title
            if i == 0 or (i == 1 and not lines[0].strip()):
                title = doc.add_heading(line, level=0)
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
                # Section headings
                doc.add_heading(line, level=1)
            elif line.endswith(':') and len(line.split()) <= 3:
                # Subsection headings
                doc.add_heading(line, level=2)
            else:
                # Regular paragraph
                paragraph = doc.add_paragraph(line)
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Created {output_path}")
        
    except Exception as e:
        logger.error(f"Error creating {output_path}: {str(e)}")


def main():
    """Create all sample .docx files."""
    logger.info("Creating sample .docx skeleton files...")
    
    # Define source text files and target .docx files
    conversions = [
        ("data/skeletons/service_agreement_skeleton.txt", "data/skeletons/service_agreement_skeleton.docx"),
        ("data/skeletons/nda_skeleton.txt", "data/skeletons/nda_skeleton.docx"),
    ]
    
    # Ensure directories exist
    Path("data/skeletons").mkdir(parents=True, exist_ok=True)
    
    # Create .docx files
    for text_file, docx_file in conversions:
        if Path(text_file).exists():
            create_docx_from_text(text_file, docx_file)
        else:
            logger.warning(f"Source file not found: {text_file}")
    
    # Create a reference contract .docx
    if Path("data/references/high_quality_service_agreement.txt").exists():
        create_docx_from_text(
            "data/references/high_quality_service_agreement.txt",
            "data/references/high_quality_service_agreement.docx"
        )
    
    logger.info("Sample .docx file creation completed!")


if __name__ == "__main__":
    main()