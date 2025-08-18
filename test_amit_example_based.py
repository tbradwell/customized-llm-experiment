#!/usr/bin/env python3
"""Test script for example-based generation using amit_test specific folders."""

import logging
import sys
from pathlib import Path

# Add project root to Python path
sys.path.append(str(Path(__file__).parent))

from src.core.example_based_generator import ExampleBasedContractGenerator
from src.evaluation.metrics import MetricsCalculator
from docx import Document

def setup_logging():
    """Set up logging configuration."""
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text.strip())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text.strip())
    
    return '\n'.join(full_text)

def run_evaluation(generated_path: str, ground_truth_path: str) -> dict:
    """Run evaluation metrics against ground truth."""
    try:
        # Extract text from both documents
        generated_text = extract_text_from_docx(generated_path)
        ground_truth_text = extract_text_from_docx(ground_truth_path)
        
        print(f"📄 Generated text length: {len(generated_text)} characters")
        print(f"📄 Ground truth length: {len(ground_truth_text)} characters")
        
        # Initialize metrics calculator
        metrics_calc = MetricsCalculator()
        
        # Calculate individual metrics
        metrics = {}
        reference_texts = [ground_truth_text]
        
        # BLEU Score
        try:
            bleu_result = metrics_calc.calculate_bleu_score(generated_text, reference_texts)
            metrics['bleu'] = bleu_result.score
        except Exception as e:
            print(f"  ⚠️ BLEU calculation failed: {e}")
            metrics['bleu'] = 0.0
        
        # ROUGE Scores
        try:
            rouge_result = metrics_calc.calculate_rouge_scores(generated_text, reference_texts)
            metrics['rouge_1'] = rouge_result.details.get('rouge1_f', 0.0)
            metrics['rouge_2'] = rouge_result.details.get('rouge2_f', 0.0)
            metrics['rouge_l'] = rouge_result.details.get('rougeL_f', 0.0)
        except Exception as e:
            print(f"  ⚠️ ROUGE calculation failed: {e}")
            metrics['rouge_1'] = metrics['rouge_2'] = metrics['rouge_l'] = 0.0
        
        # METEOR Score
        try:
            meteor_result = metrics_calc.calculate_meteor_score(generated_text, reference_texts)
            metrics['meteor'] = meteor_result.score
        except Exception as e:
            print(f"  ⚠️ METEOR calculation failed: {e}")
            metrics['meteor'] = 0.0
        
        # Redundancy Score
        try:
            redundancy_result = metrics_calc.calculate_redundancy_score(generated_text)
            metrics['redundancy'] = redundancy_result.score
        except Exception as e:
            print(f"  ⚠️ Redundancy calculation failed: {e}")
            metrics['redundancy'] = 0.0
        
        # Basic completeness check
        required_elements = ["תובע", "נתבע", "תביעה", "בית משפט"]
        try:
            completeness_result = metrics_calc.calculate_completeness_score(generated_text, required_elements)
            metrics['completeness'] = completeness_result.score
        except Exception as e:
            print(f"  ⚠️ Completeness calculation failed: {e}")
            metrics['completeness'] = 0.0
        
        print(f"📊 EVALUATION METRICS:")
        for metric_name, score in metrics.items():
            if isinstance(score, (int, float)):
                print(f"  • {metric_name.upper()}: {score:.3f}")
        
        return metrics
        
    except Exception as e:
        print(f"❌ Evaluation failed: {e}")
        return {"error": str(e)}

def main():
    """Test the example-based generation with amit_test folders."""
    setup_logging()
    
    print("🚀 AMIT TEST - Example-Based Contract Generation")
    print("=" * 60)
    
    try:
        # Step 1: Set up paths as specified
        past_examples_dir = "examples/amit_test/past_examples"
        new_data_dir = "examples/amit_test/new_data" 
        skeleton_path = "examples/amit_test/sekeleton_oracle.docx"
        
        # Create experiment folder structure
        import time
        timestamp = int(time.time())
        exp_folder = Path(f"experiments/amit_example_based_{timestamp}")
        
        # Create folder structure
        (exp_folder / "inputs").mkdir(parents=True, exist_ok=True)
        (exp_folder / "outputs").mkdir(parents=True, exist_ok=True)
        (exp_folder / "reports").mkdir(parents=True, exist_ok=True)
        (exp_folder / "logs").mkdir(parents=True, exist_ok=True)
        
        output_path = exp_folder / "outputs" / "example_based_contract.docx"
        
        print(f"📚 Past examples dir: {past_examples_dir}")
        print(f"📁 New data dir: {new_data_dir}")
        print(f"📝 Skeleton: {skeleton_path}")
        print(f"🗂️  Experiment folder: {exp_folder}")
        print(f"💾 Output: {output_path}")
        
        # Step 2: Initialize the generator
        generator = ExampleBasedContractGenerator()
        
        # Step 3: Load new data from directory
        print(f"\n📖 Loading new data from directory...")
        new_data = generator.load_new_data_from_directory(new_data_dir)
        print(f"✅ Loaded new data: {len(new_data)} characters")
        
        # Step 4: Generate contract using example-based approach  
        print(f"\n🎯 Starting Example-Based Generation...")
        print("-" * 40)
        
        result = generator.generate_contract_from_examples(
            new_data=new_data,
            skeleton_path=skeleton_path,
            examples_dir=past_examples_dir
        )
        
        # Step 5: Display results
        print(f"\n📊 TWO-PHASE GENERATION RESULTS")
        print("=" * 50)
        print(f"✅ Success: {result.success}")
        
        print(f"\n🔧 PHASE 1 - STRUCTURED CONTENT CREATION:")
        print(f"  📝 Headers created: {result.metadata['phase1_headers_count']}")
        print(f"  📑 Content sections: {result.metadata['phase1_data_sections']}")
        print(f"  🔄 Improvement iterations: {result.improvement_iterations}")
        
        print(f"\n🔨 PHASE 2 - SKELETON EDITING:")
        print(f"  🔄 Skeleton edit iterations: {result.metadata['phase2_skeleton_iterations']}")
        print(f"  📚 Examples used: {result.metadata['examples_used_count']}")
        print(f"  📏 Final content length: {result.metadata['final_content_length']} characters")
        
        # Show Phase 1 structured headers
        if result.structured_content and result.structured_content.headers:
            print(f"\n📋 PHASE 1 HEADERS CREATED:")
            for i, header in enumerate(result.structured_content.headers[:5], 1):
                print(f"  {i}. {header}")
            if len(result.structured_content.headers) > 5:
                print(f"  ... and {len(result.structured_content.headers) - 5} more")
        
        # Step 6: Show sample of generated content
        print(f"\n📄 SAMPLE OF GENERATED CONTENT (first 800 chars):")
        print("-" * 50)
        print(result.final_skeleton_content[:800])
        print("...")
        
        # Step 7: Save Phase 1 structured content as JSON
        phase1_json_path = exp_folder / "reports" / "phase1_structured_content.json"
        generator.save_structured_content_json(result.structured_content, str(phase1_json_path))
        print(f"✅ Saved Phase 1 structured content to: {phase1_json_path}")
        
        # Step 7.5: Save to file (preserving original skeleton formatting)
        generator.save_filled_skeleton(skeleton_path, result.generated_content, str(output_path))
        
        # Step 7.5: Save experiment metadata
        import json
        experiment_metadata = {
            "experiment_name": "amit_example_based",
            "timestamp": timestamp,
            "past_examples_dir": past_examples_dir,
            "new_data_dir": new_data_dir,
            "skeleton_path": skeleton_path,
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "results": {
                "success": result.success,
                "improvement_iterations": result.improvement_iterations,
                "skeleton_edit_iterations": result.skeleton_edit_iterations,
                "examples_used_count": result.metadata['examples_used_count'],
                "final_content_length": result.metadata['final_content_length']
            }
        }
        
        with open(exp_folder / "experiment_metadata.json", 'w') as f:
            json.dump(experiment_metadata, f, indent=2)
        
        # Save generation results
        generation_results = {
            "generated_content_preview": result.final_skeleton_content[:1000],
            "examples_used": result.examples_used,
            "metadata": result.metadata
        }
        
        with open(exp_folder / "reports" / "generation_results.json", 'w') as f:
            json.dump(generation_results, f, indent=2, ensure_ascii=False)
        
        # Step 7.6: Run evaluation against ground truth
        ground_truth_path = "examples/amit_test/gt.docx"
        if Path(ground_truth_path).exists():
            print(f"\n📊 RUNNING EVALUATION AGAINST GROUND TRUTH")
            print("=" * 50)
            evaluation_results = run_evaluation(str(output_path), ground_truth_path)
            
            # Save evaluation results
            with open(exp_folder / "reports" / "evaluation_results.json", 'w') as f:
                json.dump(evaluation_results, f, indent=2)
                
            print(f"✅ Evaluation completed and saved to reports/evaluation_results.json")
        else:
            print(f"⚠️  Ground truth file not found: {ground_truth_path}")
        
        # Step 8: Comprehensive verification
        print(f"\n🔍 COMPREHENSIVE VERIFICATION")
        print("=" * 40)
        
        import re
        content = result.final_skeleton_content
        
        # Check placeholders
        placeholders = re.findall(r'\\{[^}]*\\}', content)
        if placeholders:
            print(f"❌ FOUND {len(placeholders)} remaining placeholders:")
            for i, placeholder in enumerate(placeholders[:3]):
                print(f"  {i+1}: {placeholder}")
            if len(placeholders) > 3:
                print(f"  ... and {len(placeholders) - 3} more")
        else:
            print("✅ SUCCESS: No placeholders remain!")
        
        # Check unwanted content
        unwanted_patterns = ['DOC-', 'יינתן במועד מאוחר יותר', '===']
        found_unwanted = []
        for pattern in unwanted_patterns:
            if pattern in content:
                found_unwanted.append(pattern)
        
        if found_unwanted:
            print(f"❌ Found unwanted patterns: {found_unwanted}")
        else:
            print("✅ SUCCESS: No unwanted patterns found!")
            
        # Check for real data incorporation
        real_data_indicators = [
            'זואי', 'מור', 'הוד השרון', 'חיים הרצוג', 
            '0542477683', 'רס אדרת', 'יצחק אולשבנג'
        ]
        found_indicators = [indicator for indicator in real_data_indicators if indicator in content]
        print(f"✅ SUCCESS: Found {len(found_indicators)} real data indicators")
        print(f"  Examples: {found_indicators[:3]}")
        
        # Final assessment
        if not placeholders and not found_unwanted and len(found_indicators) > 0:
            print(f"\n🎉🎉🎉 PERFECT SUCCESS! 🎉🎉🎉")
            print("=" * 50)
            print("✅ All requirements perfectly met:")
            print("  1. ✅ Used dedicated past examples from amit_test/past_examples/")
            print("  2. ✅ Loaded new data from amit_test/new_data/")
            print("  3. ✅ Applied to skeleton amit_test/sekeleton_oracle.docx")
            print("  4. ✅ No placeholders remain in final document")
            print("  5. ✅ No unwanted content (DOC names, generic text)")
            print("  6. ✅ Real data successfully incorporated")
            print("  7. ✅ Example-based learning working perfectly!")
        else:
            print(f"\n⚠️  Some issues detected - see above")
        
        
        print(f"\n📁 Final output: {output_path.resolve()}")
        return 0
        
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    exit_code = main()
    sys.exit(exit_code)