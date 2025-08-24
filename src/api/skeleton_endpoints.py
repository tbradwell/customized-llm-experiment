"""Skeleton processor API endpoints."""

import logging
import os
import uuid
import tempfile
import shutil
from datetime import datetime
from typing import List, Optional, Dict, Any

from fastapi import APIRouter, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse

from ..skeleton_processor import SkeletonProcessor
from ..skeleton_processor.config import SkeletonProcessorConfig
from .models import (
    SkeletonProcessingRequest, SkeletonProcessingResponse,
    SkeletonValidationResponse, SkeletonStatsResponse,
    DocumentUploadResponse, ErrorResponse
)

logger = logging.getLogger(__name__)

# Create router for skeleton processor endpoints
router = APIRouter(prefix="/skeleton", tags=["skeleton_processor"])

# Global storage for skeleton processor operations
skeleton_processor: Optional[SkeletonProcessor] = None
processing_jobs: Dict[str, Dict[str, Any]] = {}
uploaded_documents: Dict[str, List[str]] = {}


def get_skeleton_processor() -> SkeletonProcessor:
    """Get or initialize skeleton processor."""
    global skeleton_processor
    
    if skeleton_processor is None:
        # Initialize with environment configuration
        config = SkeletonProcessorConfig.from_environment()
        skeleton_processor = SkeletonProcessor(
            openai_api_key=config.openai_api_key,
            embedding_model=config.embedding_model,
            high_homogeneity_threshold=config.high_homogeneity_threshold,
            low_homogeneity_threshold=config.low_homogeneity_threshold,
            language=config.language
        )
    
    return skeleton_processor


@router.post("/upload-documents", response_model=DocumentUploadResponse)
async def upload_documents(files: List[UploadFile] = File(...)):
    """Upload DOCX documents for skeleton processing."""
    try:
        upload_id = str(uuid.uuid4())
        document_paths = []
        
        # Validate files
        for file in files:
            if not file.filename.endswith('.docx'):
                raise HTTPException(
                    status_code=400,
                    detail=f"File {file.filename} is not a DOCX file"
                )
        
        # Save uploaded files
        upload_dir = os.path.join(tempfile.gettempdir(), f"skeleton_upload_{upload_id}")
        os.makedirs(upload_dir, exist_ok=True)
        
        for file in files:
            file_path = os.path.join(upload_dir, file.filename)
            
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            document_paths.append(file_path)
        
        # Store document paths
        uploaded_documents[upload_id] = document_paths
        
        # Get processor for validation
        processor = get_skeleton_processor()
        
        # Validate documents
        is_valid, validation_issues = processor.validate_input_documents(document_paths)
        
        return DocumentUploadResponse(
            success=True,
            upload_id=upload_id,
            documents_count=len(document_paths),
            filenames=[file.filename for file in files],
            validation_passed=is_valid,
            validation_issues=validation_issues if not is_valid else [],
            total_size_mb=sum(os.path.getsize(path) for path in document_paths) / (1024 * 1024)
        )
        
    except Exception as e:
        logger.error(f"Error uploading documents: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


@router.post("/process", response_model=SkeletonProcessingResponse)
async def process_documents_to_skeleton(
    request: SkeletonProcessingRequest,
    background_tasks: BackgroundTasks
):
    """Process uploaded documents to generate skeleton template."""
    try:
        # Validate upload ID
        if request.upload_id not in uploaded_documents:
            raise HTTPException(
                status_code=404,
                detail="Upload ID not found. Please upload documents first."
            )
        
        document_paths = uploaded_documents[request.upload_id]
        
        # Generate job ID
        job_id = str(uuid.uuid4())
        
        # Initialize job tracking
        processing_jobs[job_id] = {
            "status": "processing",
            "started_at": datetime.now(),
            "upload_id": request.upload_id,
            "document_count": len(document_paths),
            "progress": 0
        }
        
        # Get processor
        processor = get_skeleton_processor()
        
        # Start processing in background
        if request.async_processing:
            background_tasks.add_task(
                process_skeleton_background,
                job_id, processor, document_paths
            )
            
            return SkeletonProcessingResponse(
                success=True,
                job_id=job_id,
                status="processing",
                message="Skeleton processing started in background",
                async_processing=True,
                status_check_url=f"/skeleton/status/{job_id}"
            )
        else:
            # Process synchronously
            try:
                skeleton_doc = processor.process_documents_to_skeleton(document_paths)
                
                # Update job status
                processing_jobs[job_id] = {
                    "status": "completed",
                    "started_at": processing_jobs[job_id]["started_at"],
                    "completed_at": datetime.now(),
                    "skeleton_doc": skeleton_doc,
                    "upload_id": request.upload_id,
                    "document_count": len(document_paths),
                    "progress": 100
                }
                
                return SkeletonProcessingResponse(
                    success=True,
                    job_id=job_id,
                    status="completed",
                    skeleton_id=skeleton_doc.id,
                    template_download_url=f"/skeleton/download/{skeleton_doc.id}",
                    stats_url=f"/skeleton/stats/{job_id}",
                    processing_time=(
                        processing_jobs[job_id]["completed_at"] - 
                        processing_jobs[job_id]["started_at"]
                    ).total_seconds(),
                    message="Skeleton processing completed successfully"
                )
                
            except Exception as e:
                # Update job status
                processing_jobs[job_id]["status"] = "failed"
                processing_jobs[job_id]["error"] = str(e)
                raise
        
    except Exception as e:
        logger.error(f"Error processing skeleton: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")


async def process_skeleton_background(job_id: str, processor: SkeletonProcessor, 
                                    document_paths: List[str]):
    """Background task for skeleton processing."""
    try:
        # Update progress
        processing_jobs[job_id]["progress"] = 10
        processing_jobs[job_id]["status"] = "preprocessing"
        
        # Process documents
        skeleton_doc = processor.process_documents_to_skeleton(document_paths)
        
        # Update job completion
        processing_jobs[job_id].update({
            "status": "completed",
            "completed_at": datetime.now(),
            "skeleton_doc": skeleton_doc,
            "progress": 100
        })
        
        logger.info(f"Background skeleton processing completed for job {job_id}")
        
    except Exception as e:
        logger.error(f"Background skeleton processing failed for job {job_id}: {str(e)}")
        processing_jobs[job_id].update({
            "status": "failed",
            "error": str(e),
            "failed_at": datetime.now()
        })


@router.get("/status/{job_id}")
async def get_processing_status(job_id: str):
    """Get processing status for a skeleton generation job."""
    if job_id not in processing_jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job_info = processing_jobs[job_id]
    
    response = {
        "job_id": job_id,
        "status": job_info["status"],
        "progress": job_info.get("progress", 0),
        "started_at": job_info["started_at"].isoformat(),
        "document_count": job_info.get("document_count", 0)
    }
    
    if job_info["status"] == "completed":
        skeleton_doc = job_info.get("skeleton_doc")
        if skeleton_doc:
            response.update({
                "completed_at": job_info["completed_at"].isoformat(),
                "skeleton_id": skeleton_doc.id,
                "template_download_url": f"/skeleton/download/{skeleton_doc.id}",
                "stats_url": f"/skeleton/stats/{job_id}",
                "processing_time": (
                    job_info["completed_at"] - job_info["started_at"]
                ).total_seconds()
            })
    
    elif job_info["status"] == "failed":
        response.update({
            "failed_at": job_info.get("failed_at", datetime.now()).isoformat(),
            "error": job_info.get("error", "Unknown error")
        })
    
    return response


@router.get("/download/{skeleton_id}")
async def download_skeleton_template(skeleton_id: str):
    """Download generated skeleton template."""
    # Find job with this skeleton ID
    job_info = None
    for job_data in processing_jobs.values():
        if (job_data.get("skeleton_doc") and 
            job_data["skeleton_doc"].id == skeleton_id):
            job_info = job_data
            break
    
    if not job_info:
        raise HTTPException(status_code=404, detail="Skeleton not found")
    
    skeleton_doc = job_info["skeleton_doc"]
    
    if not os.path.exists(skeleton_doc.template_path):
        raise HTTPException(status_code=404, detail="Template file not found")
    
    return FileResponse(
        skeleton_doc.template_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"skeleton_template_{skeleton_id}.docx"
    )


@router.get("/stats/{job_id}", response_model=SkeletonStatsResponse)
async def get_skeleton_statistics(job_id: str):
    """Get detailed statistics for skeleton processing job."""
    if job_id not in processing_jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job_info = processing_jobs[job_id]
    
    if job_info["status"] != "completed":
        raise HTTPException(
            status_code=400, 
            detail="Statistics only available for completed jobs"
        )
    
    skeleton_doc = job_info["skeleton_doc"]
    
    # Get detailed statistics
    processor = get_skeleton_processor()
    document_paths = uploaded_documents[job_info["upload_id"]]
    
    try:
        pipeline_stats = processor.get_processing_pipeline_stats(
            document_paths, skeleton_doc
        )
        
        return SkeletonStatsResponse(
            job_id=job_id,
            skeleton_id=skeleton_doc.id,
            processing_completed_at=job_info["completed_at"],
            total_processing_time=(
                job_info["completed_at"] - job_info["started_at"]
            ).total_seconds(),
            input_statistics={
                "document_count": len(document_paths),
                "total_paragraphs": skeleton_doc.total_paragraphs_processed,
                "total_clusters": skeleton_doc.total_clusters_found
            },
            algorithm_parameters=pipeline_stats.get("algorithm_parameters", {}),
            clustering_statistics={
                "clusters_created": skeleton_doc.total_clusters_found,
                "blocks_generated": skeleton_doc.blocks_generated,
                "uncertain_blocks": skeleton_doc.uncertain_blocks
            },
            output_information={
                "skeleton_id": skeleton_doc.id,
                "template_path": skeleton_doc.template_path,
                "delimiter_positions": len(skeleton_doc.delimiter_positions),
                "created_at": skeleton_doc.created_at.isoformat()
            }
        )
        
    except Exception as e:
        logger.error(f"Error generating statistics: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Statistics generation failed: {str(e)}")


@router.post("/validate", response_model=SkeletonValidationResponse)
async def validate_skeleton_template(file: UploadFile = File(...)):
    """Validate an existing skeleton template."""
    try:
        if not file.filename.endswith('.docx'):
            raise HTTPException(
                status_code=400,
                detail="File must be a DOCX document"
            )
        
        # Save file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            shutil.copyfileobj(file.file, tmp_file)
            tmp_path = tmp_file.name
        
        try:
            from docx import Document
            doc = Document(tmp_path)
            
            # Basic validation
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            
            # Check for delimiter patterns
            all_text = "\n".join(p.text for p in doc.paragraphs)
            
            delimiter_counts = {
                "block_start": all_text.count("%}"),
                "block_end": all_text.count("{%"),
                "uncertain": all_text.count("~}")
            }
            
            # Check if delimiters are balanced
            balanced_delimiters = delimiter_counts["block_start"] == delimiter_counts["block_end"]
            
            # Validation results
            validation_results = {
                "file_format": "valid",
                "paragraph_count": paragraph_count,
                "delimiter_counts": delimiter_counts,
                "balanced_delimiters": balanced_delimiters,
                "has_content": paragraph_count > 0
            }
            
            issues = []
            if not balanced_delimiters:
                issues.append("Block delimiters are not balanced")
            if paragraph_count == 0:
                issues.append("Document has no content")
            
            is_valid = len(issues) == 0
            
            return SkeletonValidationResponse(
                success=True,
                valid_skeleton=is_valid,
                validation_results=validation_results,
                issues=issues,
                filename=file.filename,
                file_size_mb=os.path.getsize(tmp_path) / (1024 * 1024)
            )
            
        finally:
            # Clean up temporary file
            os.unlink(tmp_path)
        
    except Exception as e:
        logger.error(f"Error validating skeleton: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Validation failed: {str(e)}")


@router.get("/jobs")
async def list_processing_jobs():
    """List all skeleton processing jobs."""
    jobs_summary = []
    
    for job_id, job_info in processing_jobs.items():
        summary = {
            "job_id": job_id,
            "status": job_info["status"],
            "started_at": job_info["started_at"].isoformat(),
            "document_count": job_info.get("document_count", 0)
        }
        
        if job_info["status"] == "completed":
            summary["completed_at"] = job_info["completed_at"].isoformat()
            summary["skeleton_id"] = job_info["skeleton_doc"].id
        elif job_info["status"] == "failed":
            summary["failed_at"] = job_info.get("failed_at", datetime.now()).isoformat()
            summary["error"] = job_info.get("error", "Unknown error")
        
        jobs_summary.append(summary)
    
    return {
        "total_jobs": len(jobs_summary),
        "jobs": jobs_summary
    }


@router.delete("/jobs/{job_id}")
async def delete_processing_job(job_id: str):
    """Delete a processing job and clean up associated files."""
    if job_id not in processing_jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    try:
        job_info = processing_jobs[job_id]
        
        # Clean up uploaded documents
        if job_info.get("upload_id") in uploaded_documents:
            document_paths = uploaded_documents[job_info["upload_id"]]
            for path in document_paths:
                try:
                    if os.path.exists(path):
                        os.unlink(path)
                except:
                    pass
            
            # Remove upload directory
            if document_paths:
                upload_dir = os.path.dirname(document_paths[0])
                try:
                    shutil.rmtree(upload_dir)
                except:
                    pass
            
            del uploaded_documents[job_info["upload_id"]]
        
        # Clean up skeleton template
        if job_info.get("skeleton_doc"):
            skeleton_doc = job_info["skeleton_doc"]
            try:
                if os.path.exists(skeleton_doc.template_path):
                    os.unlink(skeleton_doc.template_path)
            except:
                pass
        
        # Remove job record
        del processing_jobs[job_id]
        
        return {"success": True, "message": f"Job {job_id} deleted successfully"}
        
    except Exception as e:
        logger.error(f"Error deleting job {job_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Deletion failed: {str(e)}")


@router.get("/health")
async def skeleton_processor_health():
    """Health check for skeleton processor service."""
    try:
        processor = get_skeleton_processor()
        
        # Test OpenAI connection
        try:
            # Simple test - get embedding dimension
            dimension = processor.embedding_client.get_embedding_dimension()
            openai_status = "healthy"
        except Exception as e:
            openai_status = f"unhealthy: {str(e)}"
        
        return {
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "services": {
                "skeleton_processor": "healthy",
                "openai_embeddings": openai_status,
                "text_cleaner": "healthy",
                "clustering_engine": "healthy",
                "skeleton_generator": "healthy"
            },
            "active_jobs": len(processing_jobs),
            "uploaded_document_sets": len(uploaded_documents)
        }
        
    except Exception as e:
        logger.error(f"Skeleton processor health check failed: {str(e)}")
        raise HTTPException(status_code=503, detail="Service unhealthy")
