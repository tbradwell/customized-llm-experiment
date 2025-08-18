"""Document handling utilities for DOCX files."""

import logging
from typing import Dict, List, Set, Any
import copy
import re

from docx import Document

logger = logging.getLogger(__name__)


class DocHandler:
    """Utility class for handling DOCX document operations."""

    @staticmethod
    def replace_paragraph_text(paragraph, new_text):
        """
        Replace paragraph text while preserving basic formatting from the first run.
        
        Args:
            paragraph: The paragraph object to modify
            new_text: The new text content to set
            
        Returns:
            bool: True if successful, False otherwise
        """
        
        try:
            # Save formatting from first run
            first_run_format = {}
            if paragraph.runs:
                try:
                    first_run = paragraph.runs[0]
                    first_run_format = {
                        'bold': first_run.bold,
                        'italic': first_run.italic,
                        'underline': first_run.underline,
                        'font_name': first_run.font.name,
                        'font_size': first_run.font.size,
                    }
                except Exception:
                    # If we can't read formatting, continue without it
                    pass
            
            # Replace content
            paragraph.clear()
            new_run = paragraph.add_run(new_text)
            
            # Restore basic formatting
            try:
                if first_run_format.get('bold') is not None:
                    new_run.bold = first_run_format['bold']
                if first_run_format.get('italic') is not None:
                    new_run.italic = first_run_format['italic']
                if first_run_format.get('underline') is not None:
                    new_run.underline = first_run_format['underline']
                if first_run_format.get('font_name'):
                    new_run.font.name = first_run_format['font_name']
                if first_run_format.get('font_size'):
                    new_run.font.size = first_run_format['font_size']
            except Exception:
                # If formatting fails, at least we have the correct text
                pass
            
            return True
            
        except Exception:
            # If everything fails, return False
            return False



    # Simpler alternative - if the above is too complex, use this more direct approach:
    def replace_placeholders_preserve_formatting_simple(paragraph, old_text, new_text):
        """
        Simpler approach: if the text exists in the paragraph, replace it by 
        rebuilding with the first run's formatting.
        """
        
        current_text = paragraph.text
        
        if old_text not in current_text:
            return False
        
        # Save formatting from first run
        first_run_format = {}
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'bold': first_run.bold,
                'italic': first_run.italic, 
                'underline': first_run.underline,
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
            }
        
        # Replace text
        new_paragraph_text = current_text.replace(old_text, new_text)
        
        # Rebuild paragraph
        paragraph.clear()
        new_run = paragraph.add_run(new_paragraph_text)
        
        # Apply formatting
        try:
            if first_run_format.get('bold') is not None:
                new_run.bold = first_run_format['bold']
            if first_run_format.get('italic') is not None:
                new_run.italic = first_run_format['italic']
            if first_run_format.get('underline') is not None:
                new_run.underline = first_run_format['underline']
            if first_run_format.get('font_name'):
                new_run.font.name = first_run_format['font_name']
            if first_run_format.get('font_size'):
                new_run.font.size = first_run_format['font_size']
        except Exception as e:
            # If formatting fails, at least we have the text
            pass
        
        return True

    @staticmethod
    def clear_template_block_paragraphs(skeleton_doc: Document, start_idx: int, block_size: int) -> None:
        """Clear existing paragraphs in a template block."""
        for i in range(block_size):
            if start_idx + i < len(skeleton_doc.paragraphs):
                skeleton_doc.paragraphs[start_idx + i].clear()

    @staticmethod
    def get_reference_paragraphs(skeleton_doc: Document, start_idx: int, block_size: int) -> List:
        """Get reference paragraphs for duplication."""
        reference_paragraphs = []
        for i in range(block_size):
            if start_idx + i < len(skeleton_doc.paragraphs):
                reference_paragraphs.append(skeleton_doc.paragraphs[start_idx + i])
        return reference_paragraphs

    @staticmethod
    def insert_additional_template_blocks(skeleton_doc: Document, start_idx: int, 
                                        reference_paragraphs: List, list_changes: List) -> None:
        """Insert additional paragraph blocks for template duplication."""
        n_paragraphs = len(list_changes[0]) if list_changes else 0
        
        for block_idx in range(1, len(list_changes)):
            for par_index in range(n_paragraphs):
                if par_index < len(reference_paragraphs):
                    insert_position = start_idx + (n_paragraphs * block_idx) + par_index
                    DocHandler.insert_duplicate_paragraph(
                        skeleton_doc,
                        reference_paragraphs[par_index],
                        insert_position
                    )

    @staticmethod
    def add_text_to_template_blocks(skeleton_doc: Document, start_idx: int, list_changes: List) -> None:
        """Add text content to all template blocks."""
        current_position = start_idx
        for block_idx, new_paragraphs in enumerate(list_changes):
            for par_idx, text_content in enumerate(new_paragraphs):
                target_index = current_position + par_idx
                if target_index < len(skeleton_doc.paragraphs):
                    skeleton_doc.paragraphs[target_index].add_run(text_content)
            current_position += len(new_paragraphs)



    @staticmethod
    def insert_duplicate_paragraph(doc, paragraph_to_duplicate, insert_index):
        """
        Insert a duplicate paragraph at a specific index in the document.
        
        Args:
            doc (Document): The python-docx Document object
            paragraph_to_duplicate (Paragraph): The paragraph to duplicate
            insert_index (int): Index where to insert the duplicate (0-based)
                            Use -1 to append at the end
        
        Returns:
            Paragraph: The newly created duplicate paragraph
        
        Raises:
            IndexError: If insert_index is out of valid range
            TypeError: If parameters are not of correct type
        """
        
        # Validate inputs
        if not hasattr(doc, 'paragraphs'):
            raise TypeError("First parameter must be a Document object")
        
        if not hasattr(paragraph_to_duplicate, 'runs'):
            raise TypeError("Second parameter must be a Paragraph object")
        
        # Handle negative index (append at end)
        if insert_index == -1:
            insert_index = len(doc.paragraphs)
        
        # Validate index range
        if insert_index < 0 or insert_index > len(doc.paragraphs):
            raise IndexError(f"Insert index {insert_index} is out of range. "
                            f"Document has {len(doc.paragraphs)} paragraphs. "
                            f"Valid range: 0 to {len(doc.paragraphs)} (or -1 for end)")
        
        # Clone the paragraph's XML element
        source_xml = paragraph_to_duplicate._element
        new_xml = copy.deepcopy(source_xml)
        
        # Get the document body
        body = doc._body._body
        
        # Insert at the specified position
        if insert_index == len(doc.paragraphs):
            # Append at the end
            body.append(new_xml)
        else:
            # Insert at specific position
            target_paragraph = doc.paragraphs[insert_index]
            target_xml = target_paragraph._element
            body.insert(body.index(target_xml), new_xml)
        
        # Return the newly created paragraph
        # The paragraph list is updated automatically, so we can access it
        return doc.paragraphs[insert_index]


    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        doc = Document(file_path)
        full_text = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
        
        return '\n'.join(full_text)
    
    @staticmethod
    def extract_text_from_doc(doc: Document) -> str:
        """Extract text from document object.
        
        Args:
            doc: Document object
            
        Returns:
            Extracted text content
        """
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
        
        return '\n'.join(full_text)
    
    @staticmethod
    def replace_placeholders_in_doc(doc: Document, mappings: Dict[str, str]) -> None:
        """Replace placeholders in document while preserving formatting.
        
        Args:
            doc: Document object to modify
            mappings: Dictionary mapping placeholders to replacement text
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                original_text = paragraph.text
                updated_text = original_text
                
                # Replace each placeholder
                for placeholder, replacement in mappings.items():
                    updated_text = updated_text.replace(placeholder, replacement)
                
                # Update paragraph text if changed
                if updated_text != original_text:
                    paragraph.clear()
                    paragraph.add_run(updated_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            original_text = paragraph.text
                            updated_text = original_text
                            
                            # Replace each placeholder
                            for placeholder, replacement in mappings.items():
                                updated_text = updated_text.replace(placeholder, replacement)
                            
                            # Update paragraph text if changed
                            if updated_text != original_text:
                                paragraph.clear()
                                paragraph.add_run(updated_text)
    
    @staticmethod
    def extract_headers_from_docx(file_path: str) -> List[str]:
        """Extract headers from DOCX file using heuristic approach with formatting filters.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of extracted headers
        """
        try:
            doc = Document(file_path)
            headlines = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if paragraph meets basic heuristic criteria
                meets_heuristic = (len(text) < 40 or 
                                 any(word in text for word in ["חלק", "דין", "תביעה", "הצדדים", "סעדים"]))
                
                if not meets_heuristic:
                    continue
                
                # Check formatting requirements: underlined and center alignment
                has_underline = False
                is_center_aligned = False
                
                # Check for underlined text
                if para.runs:
                    for run in para.runs:
                        if run.underline:
                            has_underline = True
                            break
                
                # Check for center alignment
                if para.alignment is not None:
                    # Center alignment in python-docx is represented by WD_ALIGN_PARAGRAPH.CENTER (1)
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    is_center_aligned = (para.alignment == WD_ALIGN_PARAGRAPH.CENTER)
                
                # Only include if both underlined and center aligned
                if has_underline and is_center_aligned:
                    headlines.append(text)
            
            logger.info(f"Extracted {len(headlines)} headers from {file_path}")
            return headlines[:15]  # Limit to 15 headers
            
        except Exception as e:
            logger.error(f"Failed to extract headers from {file_path}: {e}")
            return []
    
    
    @staticmethod
    def extract_headers_from_multiple_docs(file_paths: List[str]) -> List[str]:
        """Extract and consolidate headers from multiple DOCX files.
        
        Args:
            file_paths: List of paths to DOCX files
            
        Returns:
            Consolidated list of unique headers
        """
        all_headers: Set[str] = set()
        
        for file_path in file_paths:
            headers = DocHandler.extract_headers_from_docx(file_path)
            all_headers.update(headers)
        
        # Convert to list and sort by frequency (if we had frequency data)
        # For now, just return as sorted list
        consolidated_headers = sorted(list(all_headers))
        
        logger.info(f"Consolidated {len(consolidated_headers)} unique headers from {len(file_paths)} documents")
        return consolidated_headers
    
    @staticmethod
    def extract_headers_from_examples(examples: List[Dict[str, Any]]) -> List[str]:
        """Extract common headers/sections from example documents using python-docx.
        
        Args:
            examples: List of example documents with 'source' file paths
            
        Returns:
            List of common header names found across examples
        """
        all_headers: Set[str] = set()
        header_counts: Dict[str, int] = {}
        
        # Extract headers from each example document
        for example in examples:
            if 'source' in example:
                try:
                    # Extract headers from the DOCX file
                    file_headers = DocHandler.extract_headers_from_docx(example['source'])
                    
                    # Add to all headers and count occurrences
                    for header in file_headers:
                        all_headers.add(header)
                        header_counts[header] = header_counts.get(header, 0) + 1
                        
                except Exception as e:
                    logger.warning(f"Could not extract headers from {example['source']}: {e}")
                    continue
        
        # Sort headers by frequency (most common first)
        sorted_headers = sorted(header_counts.items(), key=lambda x: x[1], reverse=True)
        
        # Return headers that appear in multiple examples or are structurally important
        common_headers = []
        for header, count in sorted_headers:
            # Include headers that appear in multiple examples OR contain legal keywords
            if count > 1 or any(keyword in header for keyword in ['טענות', 'סעדים', 'עובדות', 'רקע', 'פרטים', 'תיאור']):
                common_headers.append(header)
        
        logger.info(f"Extracted {len(common_headers)} common headers from {len(examples)} examples")
        return common_headers[:10]  # Limit to 10 headers
    
