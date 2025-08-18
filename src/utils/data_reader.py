"""Multi-format data reader for contract data extraction."""

import json
import logging
import mimetypes
import re
from pathlib import Path
from typing import Dict, Any, Optional, Union, List

from docx import Document
import PyPDF2
from PIL import Image
import pytesseract
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

# Internal imports
from config.settings import settings
from .error_handler import ProcessingError, error_handler

# CONFIGURATION CONSTANTS
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
MAX_IMAGE_SIZE = 20 * 1024 * 1024   # 20MB for images
SUPPORTED_IMAGE_FORMATS = {'.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif'}
SUPPORTED_DOC_FORMATS = {'.pdf', '.doc', '.docx', '.txt'}
SUPPORTED_EMAIL_FORMATS = {'.msg', '.eml'}

logger = logging.getLogger(__name__)


class DataReader:
    """Multi-format data reader for extracting contract information."""
    
    def __init__(self):
        """Initialize the data reader."""
        self.extraction_patterns = self._load_extraction_patterns()
    
    def read_contract_data(self, file_path: str) -> Dict[str, Any]:
        """Read contract data from various file formats.
        
        Args:
            file_path: Path to the data file
            
        Returns:
            Dictionary containing extracted contract data
            
        Raises:
            ProcessingError: If file cannot be read or parsed
        """
        try:
            file_path = Path(file_path)
            
            # Validate file
            self._validate_file(file_path)
            
            # Determine file type and extract content
            file_extension = file_path.suffix.lower()
            content = self._extract_content_by_type(file_path, file_extension)
            
            if not content or not content.strip():
                raise ProcessingError(
                    f"No readable content found in file: {file_path}",
                    suggestions=["Check if the file is corrupted", "Ensure the file contains text data"]
                )
            
            extracted_data = {'raw_content': content}
           
            
            # Add metadata
            extracted_data["_metadata"] = {
                "source_file": str(file_path),
                "file_type": file_extension,
                "content_length": len(content),
                "extraction_method": self._get_extraction_method(file_extension),
            }
            
            logger.info(f"Successfully extracted data from {file_path} ({file_extension})")
            return extracted_data
            
        except ProcessingError:
            raise
        except Exception as e:
            error_detail = error_handler.handle_error(e, {"file_path": str(file_path)})
            raise ProcessingError(
                f"Failed to read contract data from {file_path}: {str(e)}",
                details={"original_error": str(e), "file_path": str(file_path)},
                suggestions=[
                    "Check if the file format is supported",
                    "Verify the file is not corrupted",
                    "Ensure required dependencies are installed (tesseract for images)"
                ]
            )
    
    def _validate_file(self, file_path: Path) -> None:
        """Validate file exists and is within size limits."""
        if not file_path.exists():
            raise ProcessingError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            raise ProcessingError(f"Path is not a file: {file_path}")
        
        file_size = file_path.stat().st_size
        max_size = MAX_IMAGE_SIZE if file_path.suffix.lower() in SUPPORTED_IMAGE_FORMATS else MAX_FILE_SIZE
        
        if file_size > max_size:
            raise ProcessingError(
                f"File too large: {file_size / (1024*1024):.1f}MB (max {max_size / (1024*1024)}MB)"
            )
        
        # Check if format is supported
        file_extension = file_path.suffix.lower()
        all_supported = SUPPORTED_DOC_FORMATS | SUPPORTED_IMAGE_FORMATS | SUPPORTED_EMAIL_FORMATS | {'.json'}
        
        if file_extension not in all_supported:
            supported_formats = ', '.join(sorted(all_supported))
            raise ProcessingError(
                f"Unsupported file format: {file_extension}",
                details={"supported_formats": list(all_supported)},
                suggestions=[f"Supported formats: {supported_formats}"]
            )
    
    def _extract_content_by_type(self, file_path: Path, file_extension: str) -> str:
        """Extract content based on file type."""
        if file_extension == '.json':
            return self._read_json(file_path)
        elif file_extension == '.pdf':
            return self._read_pdf(file_path)
        elif file_extension in {'.doc', '.docx'}:
            return self._read_docx(file_path)
        elif file_extension == '.txt':
            return self._read_text(file_path)
        elif file_extension in SUPPORTED_IMAGE_FORMATS:
            return self._read_image_ocr(file_path)
        elif file_extension in SUPPORTED_EMAIL_FORMATS:
            return self._read_email(file_path)
        else:
            raise ProcessingError(f"Unsupported file type: {file_extension}")
    
    def _read_json(self, file_path: Path) -> str:
        """Read JSON file and return as formatted string."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Return JSON as formatted string for pattern extraction
            return json.dumps(data, indent=2, ensure_ascii=False)
            
        except json.JSONDecodeError as e:
            raise ProcessingError(f"Invalid JSON format: {e}")
    
    def _read_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file, with fallback to OCR for scanned PDFs."""
        try:
            content = []
            
            # First, try standard text extraction
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            content.append(f"--- Page {page_num + 1} ---\n{text}")
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                        continue
            
            # If no text found, try OCR on the PDF
            if not content:
                logger.info(f"No text found in PDF {file_path}, attempting OCR...")
                ocr_content = self._pdf_ocr_extraction(file_path)
                content.extend(ocr_content)
            
            return '\n\n'.join(content) if content else ""
            
        except Exception as e:
            raise ProcessingError(f"Error reading PDF: {e}")
    
    def _pdf_ocr_extraction(self, file_path: Path) -> List[str]:
        """Extract text from PDF using OCR."""
        try:
            if not PDF2IMAGE_AVAILABLE:
                logger.warning("pdf2image not available, cannot perform OCR on PDF")
                return []
            
            # Check if tesseract is available
            try:
                pytesseract.get_tesseract_version()
            except pytesseract.TesseractNotFoundError:
                logger.warning("Tesseract not available for PDF OCR")
                return []
            
            # Convert PDF pages to images
            pages = convert_from_path(file_path)
            content = []
            
            for page_num, page_image in enumerate(pages):
                try:
                    # Use Hebrew OCR if text appears to be Hebrew
                    text = pytesseract.image_to_string(page_image, lang='heb+eng', config='--psm 6')
                    if text.strip():
                        content.append(f"--- Page {page_num + 1} (OCR) ---\n{text}")
                    logger.info(f"OCR extracted {len(text)} characters from page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num + 1}: {e}")
                    continue
            
            return content
            
        except Exception as e:
            logger.warning(f"PDF OCR extraction failed: {e}")
            return []
        
    def _read_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file, including headers/footers, with fallback for image-based content."""
        try:
            doc = Document(file_path)
            content = []
            
            # Extract headers from all sections
            for section in doc.sections:
                # Extract from primary header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            content.append(f"[HEADER] {paragraph.text}")
                    
                    # Extract tables from header
                    for table in section.header.tables:
                        table_content = []
                        for row in table.rows:
                            row_content = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_content.append(cell.text.strip())
                            if row_content:
                                table_content.append(' | '.join(row_content))
                        if table_content:
                            content.append('\n--- Header Table ---\n' + '\n'.join(table_content))
            
            # Extract main document paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Extract main document tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_content.append(cell.text.strip())
                    if row_content:
                        table_content.append(' | '.join(row_content))
                if table_content:
                    content.append('\n--- Table ---\n' + '\n'.join(table_content))
            
            # Extract footers from all sections
            for section in doc.sections:
                # Extract from primary footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            content.append(f"[FOOTER] {paragraph.text}")
                    
                    # Extract tables from footer
                    for table in section.footer.tables:
                        table_content = []
                        for row in table.rows:
                            row_content = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_content.append(cell.text.strip())
                            if row_content:
                                table_content.append(' | '.join(row_content))
                        if table_content:
                            content.append('\n--- Footer Table ---\n' + '\n'.join(table_content))
            
            # If no text content found, check for images and try OCR
            if not content:
                logger.info(f"No text found in DOCX {file_path}, checking for images...")
                image_content = self._docx_image_extraction(file_path, doc)
                if image_content:
                    content.extend(image_content)
            
            return '\n\n'.join(content) if content else ""
            
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            return ""
    
    def _docx_image_extraction(self, file_path: Path, doc) -> List[str]:
        """Extract text from images in DOCX using OCR."""
        try:
            # Check if tesseract is available
            try:
                pytesseract.get_tesseract_version()
            except pytesseract.TesseractNotFoundError:
                logger.warning("Tesseract not available for DOCX image OCR")
                return []
            
            content = []
            
            # Extract images from DOCX file
            import zipfile
            import io
            
            # DOCX files are ZIP archives
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Look for image files in the media folder
                image_files = [f for f in zip_file.namelist() if f.startswith('word/media/') and 
                              any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'])]
                
                logger.info(f"Found {len(image_files)} images in DOCX file")
                
                for image_file in image_files:
                    try:
                        # Extract image data
                        image_data = zip_file.read(image_file)
                        
                        # Convert to PIL Image
                        image = Image.open(io.BytesIO(image_data))
                        
                        # Configure Tesseract for Hebrew and English
                        custom_config = r'--oem 3 --psm 6 -l heb+eng'
                        
                        # Run OCR on the image
                        ocr_text = pytesseract.image_to_string(image, config=custom_config)
                        
                        if ocr_text.strip():
                            content.append(f"--- Image: {image_file} ---\n{ocr_text.strip()}")
                            logger.info(f"Extracted {len(ocr_text)} characters from {image_file}")
                        else:
                            logger.warning(f"No text found in image: {image_file}")
                            
                    except Exception as img_error:
                        logger.warning(f"Failed to process image {image_file}: {img_error}")
                        continue
            
            if not content:
                logger.warning(f"No readable text found in DOCX images: {file_path}")
                return ["--- DOCX Image Content ---\nThis document contains images but no readable text was extracted."]
            
            return content
            
        except Exception as e:
            logger.warning(f"DOCX image extraction failed: {e}")
            return []
    
    def _read_text(self, file_path: Path) -> str:
        """Read plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise ProcessingError("Could not decode text file with any common encoding")
    
    def _read_image_ocr(self, file_path: Path) -> str:
        """Extract text from image using OCR."""
        try:
            # Check if tesseract is available
            try:
                pytesseract.get_tesseract_version()
            except pytesseract.TesseractNotFoundError:
                raise ProcessingError(
                    "Tesseract OCR not found. Please install tesseract to read image files.",
                    suggestions=[
                        "Install tesseract: sudo apt-get install tesseract-ocr (Ubuntu/Debian)",
                        "Install tesseract: brew install tesseract (macOS)",
                        "Download from: https://github.com/tesseract-ocr/tesseract"
                    ]
                )
            
            # Open and process image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode not in ('RGB', 'L'):
                image = image.convert('RGB')
            
            # Extract text using OCR with Hebrew and English support
            text = pytesseract.image_to_string(image, lang='heb+eng', config='--psm 6')
            
            return text
            
        except Exception as e:
            raise ProcessingError(f"Error processing image with OCR: {e}")
    
    def _read_email(self, file_path: Path) -> str:
        """Extract text from email files (.msg, .eml)."""
        try:
            if file_path.suffix.lower() == '.msg':
                return self._read_msg_file(file_path)
            elif file_path.suffix.lower() == '.eml':
                return self._read_eml_file(file_path)
            else:
                raise ProcessingError(f"Unsupported email format: {file_path.suffix}")
                
        except Exception as e:
            raise ProcessingError(f"Error reading email file: {e}")
    
    def _read_msg_file(self, file_path: Path) -> str:
        """Read MSG file (requires extract-msg library)."""
        try:
            import extract_msg
            
            msg = extract_msg.Message(file_path)
            content = []
            
            if msg.subject:
                content.append(f"Subject: {msg.subject}")
            if msg.sender:
                content.append(f"From: {msg.sender}")
            if msg.date:
                content.append(f"Date: {msg.date}")
            if msg.body:
                content.append(f"Body:\n{msg.body}")
            
            # Extract attachments info
            if msg.attachments:
                attachment_names = [att.longFilename or att.shortFilename for att in msg.attachments]
                content.append(f"Attachments: {', '.join(attachment_names)}")
            
            return '\n\n'.join(content)
            
        except ImportError:
            raise ProcessingError(
                "extract-msg library not found. Install with: pip install extract-msg",
                suggestions=["Install extract-msg: pip install extract-msg"]
            )
    
    def _read_eml_file(self, file_path: Path) -> str:
        """Read EML file using email library."""
        import email
        from email.parser import Parser
        
        with open(file_path, 'r', encoding='utf-8') as f:
            msg = email.message_from_file(f)
        
        content = []
        
        # Extract headers
        if msg['Subject']:
            content.append(f"Subject: {msg['Subject']}")
        if msg['From']:
            content.append(f"From: {msg['From']}")
        if msg['Date']:
            content.append(f"Date: {msg['Date']}")
        
        # Extract body
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    body = part.get_payload(decode=True)
                    if body:
                        content.append(f"Body:\n{body.decode('utf-8', errors='ignore')}")
        else:
            body = msg.get_payload(decode=True)
            if body:
                content.append(f"Body:\n{body.decode('utf-8', errors='ignore')}")
        
        return '\n\n'.join(content)
    
    def _extract_structured_data(self, content: str, file_extension: str) -> Dict[str, Any]:
        """Extract structured contract data from content."""
        # If it's JSON, parse directly
        if file_extension == '.json':
            try:
                return json.loads(content)
            except json.JSONDecodeError:
                pass
        
        # For other formats, use pattern extraction
        extracted_data = {}
        
        for field_name, patterns in self.extraction_patterns.items():
            value = self._extract_field_value(content, patterns)
            if value:
                extracted_data[field_name] = value
        
        # If no structured data found, include full content
        if not extracted_data:
            extracted_data["raw_content"] = content[:2000] + "..." if len(content) > 2000 else content
            extracted_data["extraction_note"] = "No structured fields detected. Please provide data in JSON format or update extraction patterns."
        
        return extracted_data
    
    def _extract_field_value(self, content: str, patterns: List[str]) -> Optional[str]:
        """Extract field value using regex patterns."""
        for pattern in patterns:
            try:
                match = re.search(pattern, content, re.IGNORECASE | re.MULTILINE)
                if match:
                    # Return first capturing group or full match
                    return match.group(1) if match.groups() else match.group(0)
            except re.error as e:
                logger.warning(f"Invalid regex pattern '{pattern}': {e}")
                continue
        
        return None
    
    def _load_extraction_patterns(self) -> Dict[str, List[str]]:
        """Load regex patterns for extracting common contract fields."""
        return {
            "client_name": [
                r"client[:\s]+([^,\n]+)",
                r"party\s*a[:\s]+([^,\n]+)",
                r"buyer[:\s]+([^,\n]+)",
                r"customer[:\s]+([^,\n]+)",
                r"company[:\s]+([^,\n]+)",
            ],
            "provider_name": [
                r"provider[:\s]+([^,\n]+)",
                r"party\s*b[:\s]+([^,\n]+)",
                r"seller[:\s]+([^,\n]+)",
                r"vendor[:\s]+([^,\n]+)",
                r"contractor[:\s]+([^,\n]+)",
            ],
            "contract_type": [
                r"agreement\s+type[:\s]+([^,\n]+)",
                r"contract\s+type[:\s]+([^,\n]+)",
                r"(service\s+agreement|consulting\s+agreement|nda|non-disclosure)",
            ],
            "contract_value": [
                r"total[:\s]+\$?([\d,]+(?:\.\d{2})?)",
                r"amount[:\s]+\$?([\d,]+(?:\.\d{2})?)",
                r"value[:\s]+\$?([\d,]+(?:\.\d{2})?)",
                r"\$[\d,]+(?:\.\d{2})?",
            ],
            "start_date": [
                r"start\s+date[:\s]+([^\n,]+)",
                r"effective\s+date[:\s]+([^\n,]+)",
                r"commence[s]?\s+on[:\s]+([^\n,]+)",
                r"\b\d{1,2}[/-]\d{1,2}[/-]\d{4}\b",
            ],
            "end_date": [
                r"end\s+date[:\s]+([^\n,]+)",
                r"expir[ey]\s+date[:\s]+([^\n,]+)",
                r"terminat[ei]s?\s+on[:\s]+([^\n,]+)",
            ],
            "payment_terms": [
                r"payment\s+terms[:\s]+([^\n.]+)",
                r"billing[:\s]+([^\n.]+)",
                r"invoic[ei]ng[:\s]+([^\n.]+)",
            ],
            "governing_law": [
                r"govern[ei]ng\s+law[:\s]+([^\n.]+)",
                r"jurisdiction[:\s]+([^\n.]+)",
                r"laws?\s+of\s+([^\n,.]+)",
            ]
        }
    
    def _get_extraction_method(self, file_extension: str) -> str:
        """Get description of extraction method used."""
        method_map = {
            '.json': 'Direct JSON parsing',
            '.pdf': 'PDF text extraction',
            '.docx': 'Word document parsing',
            '.txt': 'Plain text reading',
            '.jpg': 'OCR text recognition', '.jpeg': 'OCR text recognition',
            '.png': 'OCR text recognition', '.tiff': 'OCR text recognition',
            '.bmp': 'OCR text recognition', '.gif': 'OCR text recognition',
            '.msg': 'MSG email parsing', '.eml': 'EML email parsing'
        }
        return method_map.get(file_extension, 'Unknown method')
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get list of supported file formats by category."""
        return {
            "documents": list(SUPPORTED_DOC_FORMATS),
            "images": list(SUPPORTED_IMAGE_FORMATS),
            "emails": list(SUPPORTED_EMAIL_FORMATS),
            "structured": [".json"]
        }