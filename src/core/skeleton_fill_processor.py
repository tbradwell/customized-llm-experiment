"""Skeleton fill processor for filling skeleton documents with new data."""

import logging
import re
from typing import Dict, List, Any, Tuple
from docx import Document

from ..processors.paragraph_processor import ParagraphProcessor
from ..processors.template_processor import TemplateProcessor
from ..utils.doc_handler import DocHandler

logger = logging.getLogger(__name__)


class SkeletonFillProcessor:
    """Fills skeleton documents with new data using 4-phase approach."""
    
    def __init__(self):
        self.paragraph_processor = ParagraphProcessor()
        self.template_processor = TemplateProcessor()
        self.max_iterations = 10
    
    def fill_skeleton_with_data(self, skeleton_doc: Document, new_data: str) -> Tuple[Document, int]:
        """Fill skeleton document with new data using iterative processing.
        
        Args:
            skeleton_doc: Document skeleton to fill
            new_data: New data to use for filling placeholders
            
        Returns:
            Tuple of (filled_document, iterations_used)
        """
        logger.info("Starting skeleton filling process with new data")
        
        for iteration in range(self.max_iterations):
            logger.info(f"Filling iteration {iteration + 1}/{self.max_iterations}")
            
            changes_made = False
            
            # Execute the 4 phases in order
            regular_paragraphs, template_blocks = self._phase1_identify_fillable_content(skeleton_doc)
            
            template_changes = self._phase2_fill_template_blocks(skeleton_doc, template_blocks, new_data)
            if template_changes:
                changes_made = True
            
            regular_changes = self._phase3_fill_regular_paragraphs(skeleton_doc, regular_paragraphs, new_data)
            if regular_changes:
                changes_made = True
            
            table_changes = self._phase4_fill_table_content(skeleton_doc, new_data)
            if table_changes:
                changes_made = True
            
            # Check completion
            current_text = DocHandler.extract_text_from_doc(skeleton_doc)
            remaining_placeholders = re.findall(r'\{[^}]*\}', current_text)
            
            logger.info(f"Iteration {iteration + 1} complete. Remaining placeholders: {len(remaining_placeholders)}")
            
            if not remaining_placeholders:
                logger.info(f"All placeholders filled after {iteration + 1} iterations")
                break
            
            if not changes_made:
                logger.warning(f"No changes made in iteration {iteration + 1}, stopping")
                break
        
        logger.info(f"Skeleton filling completed in {iteration + 1} iterations")
        return skeleton_doc, iteration + 1
    
    def _phase1_identify_fillable_content(self, skeleton_doc: Document) -> Tuple[List[Dict], List[Dict]]:
        """Phase 1: Identify all content that needs to be filled with new data.
        
        Returns:
            Tuple of (regular_paragraphs, template_blocks)
        """
        logger.info("Phase 1: Identifying content to fill with new data...")
        
        regular_paragraphs = []      # Paragraphs with simple placeholders like {name}
        template_blocks = []         # Template blocks with {% %} that need duplication
        
        current_block = []
        block_start_idx = None
        
        # Scan through document to identify what needs filling
        for i, paragraph in enumerate(skeleton_doc.paragraphs):
            if not paragraph.text.strip():
                continue
            
            # Check for {% %} block placeholders - template loops that need multiple paragraphs
            if '%}' in paragraph.text and len(current_block) == 0 or \
               '{%' in paragraph.text and len(current_block) == 0:
                # Starting a template block
                if block_start_idx is None:
                    block_start_idx = i
                current_block.append(paragraph.text)
                
            elif '{%' in paragraph.text and len(current_block) > 0 or \
                 '%}' in paragraph.text and len(current_block) > 0:
                # Ending a template block
                current_block.append(paragraph.text)
                if current_block and block_start_idx is not None:
                    template_blocks.append({
                        'start_index': block_start_idx,
                        'paragraph_texts': current_block.copy(),
                        'block_size': len(current_block)
                    })
                # Reset for next block
                current_block = []
                block_start_idx = None
                
            elif len(current_block) > 0:
                # Continue building current template block
                current_block.append(paragraph.text)
                
            elif '}' in paragraph.text:
                # Regular placeholder paragraph - store additional identifying info
                regular_paragraphs.append({
                    'index': i,
                    'original_text': paragraph.text,  # Store original text for comparison
                    'text_hash': hash(paragraph.text.strip()),  # Hash for quick comparison
                    'char_count': len(paragraph.text.strip()),  # Length for validation
                    'has_placeholders': True  # Flag to indicate this needs filling
                })
        
        logger.info(f"Found {len(regular_paragraphs)} regular paragraphs and {len(template_blocks)} template blocks to fill")
        
        return regular_paragraphs, template_blocks
    
    def _phase2_fill_template_blocks(self, skeleton_doc: Document, template_blocks: List[Dict], new_data: str) -> bool:
        """Phase 2: Fill template blocks that need duplication with new data.
        
        Args:
            skeleton_doc: Document to process
            template_blocks: List of template block info from phase 1
            new_data: New data for filling placeholders
            
        Returns:
            bool: True if any changes were made
        """
        logger.info("Phase 2: Filling template blocks with new data...")
        
        changes_made = False
        
        # Process template blocks in reverse order so that earlier insertions 
        # don't affect the indices of later blocks
        for block_info in reversed(template_blocks):
            start_idx = block_info['start_index']
            block_size = block_info['block_size']
            
            logger.info(f"Filling template block at index {start_idx}, size {block_size}")
            
            # Use template processor to fill the block
            list_changes, block_changed = self.template_processor.fill_template_block(
                block_info['paragraph_texts'], new_data
            )
            
            if block_changed:
                # Clear existing paragraphs in the template block
                DocHandler.clear_template_block_paragraphs(skeleton_doc, start_idx, block_size)
                
                # Store references and insert additional blocks
                reference_paragraphs = DocHandler.get_reference_paragraphs(skeleton_doc, start_idx, block_size)
                DocHandler.insert_additional_template_blocks(skeleton_doc, start_idx, reference_paragraphs, list_changes)
                
                # Add text content to all blocks
                DocHandler.add_text_to_template_blocks(skeleton_doc, start_idx, list_changes)
                
                changes_made = True
        
        return changes_made
    
    def _phase3_fill_regular_paragraphs(self, skeleton_doc: Document, regular_paragraphs: List[Dict], new_data: str) -> bool:
        """Phase 3: Fill regular paragraphs with simple placeholders using new data.
        
        Args:
            skeleton_doc: Document to process
            regular_paragraphs: List of regular paragraph info from phase 1
            new_data: New data for filling placeholders
            
        Returns:
            bool: True if any changes were made
        """
        logger.info("Phase 3: Filling regular paragraphs with new data...")
        
        changes_made = False
        
        for para_info in regular_paragraphs:
            original_index = para_info['index']
            original_text = para_info['original_text']
            text_hash = para_info['text_hash']
            char_count = para_info['char_count']
            
            # First try to find paragraph at original position with safe text comparison
            target_paragraph = None
            
            if original_index < len(skeleton_doc.paragraphs):
                current_para = skeleton_doc.paragraphs[original_index]
                # Safe comparison using text content and validation
                if (current_para.text.strip() == original_text.strip() and 
                    len(current_para.text.strip()) == char_count and
                    hash(current_para.text.strip()) == text_hash):
                    target_paragraph = current_para
                    logger.debug(f"Found paragraph at original index {original_index}")
            
            # If not found at original position, search through document
            if target_paragraph is None:
                logger.info(f"Paragraph moved from index {original_index}, searching by content...")
                target_paragraph = self._find_moved_paragraph(skeleton_doc, original_text, text_hash, char_count)
            
            # Fill the paragraph if found
            if target_paragraph:
                if self.paragraph_processor.fill_paragraph_with_data(target_paragraph, new_data):
                    changes_made = True
            else:
                logger.warning(f"Could not locate paragraph with text: {original_text[:50]}...")
        
        return changes_made
    
    def _phase4_fill_table_content(self, skeleton_doc: Document, new_data: str) -> bool:
        """Phase 4: Fill table cell paragraphs with new data.
        
        Args:
            skeleton_doc: Document to process
            new_data: New data for filling placeholders
            
        Returns:
            bool: True if any changes were made
        """
        logger.info("Phase 4: Filling table content with new data...")
        
        changes_made = False
        
        for table in skeleton_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{' in paragraph.text:
                            if self.paragraph_processor.fill_paragraph_with_data(paragraph, new_data):
                                changes_made = True
        
        return changes_made
    
    def _find_moved_paragraph(self, skeleton_doc: Document, target_text: str, text_hash: int, char_count: int):
        """Helper method to find a paragraph that may have moved due to template processing.
        
        Args:
            skeleton_doc: Document to search
            target_text: Original text to find
            text_hash: Hash of the text for validation
            char_count: Character count for validation
            
        Returns:
            Paragraph object if found, None otherwise
        """
        for i, paragraph in enumerate(skeleton_doc.paragraphs):
            if (paragraph.text.strip() == target_text.strip() and 
                len(paragraph.text.strip()) == char_count and
                hash(paragraph.text.strip()) == text_hash):
                logger.debug(f"Found moved paragraph at index {i}")
                return paragraph
        
        return None
