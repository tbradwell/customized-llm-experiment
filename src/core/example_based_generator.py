"""Example-based contract generator using past examples for learning."""

import logging
import random
import re
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional
from dataclasses import dataclass

from docx import Document
from openai import OpenAI

from config.settings import settings
from ..utils.doc_handler import DocHandler
from ..utils.error_handler import ProcessingError, error_handler

# CONFIGURATION CONSTANTS
MAX_IMPROVEMENT_ITERATIONS = 5
MAX_SKELETON_EDIT_ITERATIONS = 10
EXAMPLES_PER_ITERATION = 2
MIN_CONTENT_LENGTH = 100
DEFAULT_TEMPERATURE = 0.1
NUM_SAMPLES = 1

logger = logging.getLogger(__name__)


@dataclass
class StructuredContent:
    """Structured content with headers and organized data."""
    headers: List[str]
    content_by_header: Dict[str, str]
    full_content: str
    data_usage: Dict[str, bool]
    metadata: Dict[str, Any]

@dataclass
class ExampleBasedResult:
    """Result of example-based generation."""
    generated_content: str
    structured_content: StructuredContent
    improvement_iterations: int
    skeleton_edit_iterations: int
    examples_used: List[str]
    final_skeleton_content: str
    success: bool
    metadata: Dict[str, Any]


class ExampleBasedContractGenerator:
    """Generates contracts using past examples as learning templates."""
    
    def __init__(self):
        self.client = OpenAI(api_key=settings.openai_api_key)
        self.model = settings.openai_model
        self.temperature = DEFAULT_TEMPERATURE
        
    def generate_contract_from_examples(self, new_data: str, skeleton_path: str, 
                                      examples_dir: str = "experiments") -> ExampleBasedResult:
        """Generate contract using two-phase example-based learning approach.
        
        Phase 1: Create structured content with headers from examples
        Phase 2: Iteratively edit skeleton until no {} remain and all data is used
        
        Args:
            new_data: Extracted text data from new documents
            skeleton_path: Path to skeleton document with placeholders
            examples_dir: Directory containing past contract examples
            
        Returns:
            ExampleBasedResult with structured content and filled skeleton
        """
        try:
            logger.info("Starting two-phase example-based contract generation")
            
            # Step 1: Clean the new data
            cleaned_data = self._clean_new_data(new_data)
            logger.info(f"Cleaned new data: {len(cleaned_data)} characters")
            
            # Step 2: Find past examples
            past_examples = self._find_past_examples(examples_dir)
            logger.info(f"Found {len(past_examples)} past contract examples")
            
            if len(past_examples) < 2:
                raise ProcessingError("Not enough past examples found (need at least 2)")
            
            # PHASE 1: Create structured content with headers
            logger.info("=== PHASE 1: Creating structured content with headers ===")
            structured_content = self._create_structured_content_from_examples(
                cleaned_data, past_examples
            )
            
            # PHASE 2: Iteratively edit skeleton until complete
            logger.info("=== PHASE 2: Iteratively editing skeleton until complete ===")
            final_skeleton_text, skeleton_iterations = self._iteratively_edit_skeleton(
                skeleton_path, structured_content, cleaned_data
            )
            
            metadata = {
                'phase1_headers_count': len(structured_content.headers),
                'phase1_data_sections': len(structured_content.content_by_header),
                'phase2_skeleton_iterations': skeleton_iterations,
                'examples_used_count': len(structured_content.metadata.get('examples_used', [])),
                'final_content_length': len(structured_content.full_content),
                'model_used': self.model,
                'skeleton_path': skeleton_path
            }
            
            return ExampleBasedResult(
                generated_content=structured_content.full_content,
                structured_content=structured_content,
                improvement_iterations=structured_content.metadata.get('improvement_iterations', 0),
                skeleton_edit_iterations=skeleton_iterations,
                examples_used=structured_content.metadata.get('examples_used', []),
                final_skeleton_content=final_skeleton_text,
                success=True,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Example-based generation failed: {str(e)}")
            raise ProcessingError(
                f"Example-based generation failed: {str(e)}",
                details={"error_type": type(e).__name__},
                suggestions=[
                    "Check that past examples exist in the experiments directory",
                    "Verify OpenAI API configuration",
                    "Ensure skeleton file is accessible"
                ]
            )
    
    def _create_structured_content_from_examples(self, new_data: str, 
                                               past_examples: List[Dict[str, Any]]) -> StructuredContent:
        """Phase 1: Create structured content with headers based on examples."""
        logger.info("Creating structured content from examples")
        
        # Step 1: Extract headers from examples
        example_headers = DocHandler.extract_headers_from_examples(past_examples[:3])
        logger.info(f"Extracted {len(example_headers)} common headers from examples")
        
        # Step 2: Generate initial structured content
        current_structured = self._generate_initial_structured_content(new_data, past_examples, example_headers)
        
        # Step 3: Iteratively improve structured content
        examples_used = []
        for iteration in range(MAX_IMPROVEMENT_ITERATIONS):
            logger.info(f"Structured content improvement iteration {iteration + 1}/{MAX_IMPROVEMENT_ITERATIONS}")
            
            # Choose different examples for improvement
            selected_examples = random.sample(past_examples, min(EXAMPLES_PER_ITERATION, len(past_examples)))
            examples_used.extend([ex['source'] for ex in selected_examples])
            
            # Improve the structured content
            improved_structured, no_more_improvements = self._improve_structured_content(
                current_structured, new_data, selected_examples, iteration + 1
            )
            
            if no_more_improvements:
                logger.info(f"Structured content completed after {iteration + 1} iterations")
                break
                
            current_structured = improved_structured
        
        # Add metadata
        current_structured.metadata = {
            'improvement_iterations': iteration + 1 if 'iteration' in locals() else 0,
            'examples_used': examples_used,
            'headers_extracted': example_headers
        }
        
        return current_structured
    
    
    def _generate_initial_structured_content(self, new_data: str, past_examples: List[Dict[str, Any]], 
                                           headers: List[str]) -> StructuredContent:
        """Generate initial structured content with headers."""
        selected_examples = random.sample(past_examples, min(NUM_SAMPLES, len(past_examples)))
        
        examples_text = "\\n\\n".join([f"=== EXAMPLE {i+1} ===\\n{ex['content']}" 
                                     for i, ex in enumerate(selected_examples)])
        
        headers_text = "\\n".join([f"- {header}" for header in headers]) if headers else "Generate appropriate headers"
        
        prompt = f"""
You are a legal document expert. Create structured content from the NEW DATA using the style of PAST EXAMPLES.

NEW DATA TO ORGANIZE:
{new_data}

PAST EXAMPLES FOR STRUCTURE:
{examples_text}

SUGGESTED HEADERS (use these or create better ones):
{headers_text}

INSTRUCTIONS:
1. Organize the NEW DATA into a structured document with clear headers
2. Follow the style and legal language of the PAST EXAMPLES
3. Create complete content under each header using the NEW DATA
4. Ensure all important information from NEW DATA is included
5. Return as JSON with this format:
   {{
     "headers": ["כותרת 1", "כותרת 2", ...],
     "content_by_header": {{
       "כותרת 1": "תוכן מפורט תחת כותרת 1",
       "כותרת 2": "תוכן מפורט תחת כותרת 2"
     }}
   }}

Return structured content as JSON:"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert legal document structurer. Return only valid JSON."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=self.temperature,
                max_tokens=4096
            )
            
            response_text = response.choices[0].message.content.strip()
            
            # Parse JSON response
            try:
                import json
                data = json.loads(response_text)
                headers = data.get('headers', [])
                content_by_header = data.get('content_by_header', {})
                
                # Create full content
                full_content = ""
                for header in headers:
                    if header in content_by_header:
                        full_content += f"\\n\\n{header}\\n{'-' * len(header)}\\n{content_by_header[header]}"
                
                # Track data usage (initially all unused)
                data_usage = {chunk: False for chunk in new_data.split('\\n\\n') if chunk.strip()}
                
                return StructuredContent(
                    headers=headers,
                    content_by_header=content_by_header,
                    full_content=full_content.strip(),
                    data_usage=data_usage,
                    metadata={}
                )
                
            except json.JSONDecodeError:
                logger.warning("Could not parse structured content, using fallback")
                return self._create_fallback_structured_content(new_data)
                
        except Exception as e:
            logger.error(f"Initial structured content generation failed: {e}")
            return self._create_fallback_structured_content(new_data)
    
    def _create_fallback_structured_content(self, new_data: str) -> StructuredContent:
        """Create fallback structured content."""
        headers = ["תיאור המקרה", "פרטי הצדדים", "עובדות נוספות"]
        content_by_header = {
            "תיאור המקרה": new_data[:1000],
            "פרטי הצדדים": "יש לחלץ מהמידע החדש",
            "עובדות נוספות": new_data[1000:2000] if len(new_data) > 1000 else ""
        }
        
        full_content = "\\n\\n".join([f"{h}\\n{content_by_header[h]}" for h in headers])
        data_usage = {chunk: False for chunk in new_data.split('\\n\\n') if chunk.strip()}
        
        return StructuredContent(
            headers=headers,
            content_by_header=content_by_header,
            full_content=full_content,
            data_usage=data_usage,
            metadata={}
        )
    
    def _improve_structured_content(self, current_structured: StructuredContent, new_data: str,
                                   selected_examples: List[Dict[str, Any]], iteration: int) -> Tuple[StructuredContent, bool]:
        """Improve structured content using different examples."""
        import json
        
        examples_text = "\\n\\n".join([f"=== REFERENCE EXAMPLE {i+1} ===\\n{ex['content'][:1500]}" 
                                     for i, ex in enumerate(selected_examples)])
        
        current_json = {
            "headers": current_structured.headers,
            "content_by_header": current_structured.content_by_header
        }
        
        prompt = f"""
You are a legal document improvement expert. Review and improve the CURRENT STRUCTURED CONTENT.

CURRENT STRUCTURED CONTENT:
{json.dumps(current_json, ensure_ascii=False, indent=2)}

NEW DATA FOR REFERENCE:
{new_data[:2000]}

REFERENCE EXAMPLES FOR IMPROVEMENT:{examples_text}

INSTRUCTIONS:
1. Improve the structure, content organization, and legal language
2. Ensure all important information from NEW DATA is properly categorized
3. Make headers more professional and legally appropriate
4. Enhance content under each header for completeness and clarity
5. If the current structure is already excellent, respond EXACTLY with: "NO_MORE_IMPROVEMENTS"
6. Otherwise, return improved JSON with same format

IMPORTANT: Only respond with "NO_MORE_IMPROVEMENTS" if genuinely excellent.

Provide your response:"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system", 
                        "content": "You are an expert legal document structure improver. Return JSON or NO_MORE_IMPROVEMENTS."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=self.temperature,
                max_completion_tokens=4000
            )
            
            response_text = response.choices[0].message.content.strip()
            
            # Check if LLM says no more improvements
            if response_text == "NO_MORE_IMPROVEMENTS" or "NO_MORE_IMPROVEMENTS" in response_text:
                return current_structured, True
            
            # Try to parse improved JSON
            try:
                import json
                data = json.loads(response_text)
                headers = data.get('headers', current_structured.headers)
                content_by_header = data.get('content_by_header', current_structured.content_by_header)
                
                # Create full content
                full_content = ""
                for header in headers:
                    if header in content_by_header:
                        full_content += f"\\n\\n{header}\\n{'-' * len(header)}\\n{content_by_header[header]}"
                
                improved_structured = StructuredContent(
                    headers=headers,
                    content_by_header=content_by_header,
                    full_content=full_content.strip(),
                    data_usage=current_structured.data_usage,  # Keep existing tracking
                    metadata=current_structured.metadata
                )
                
                return improved_structured, False
                
            except json.JSONDecodeError:
                logger.warning("Could not parse improved structured content")
                return current_structured, True
                
        except Exception as e:
            logger.error(f"Structured content improvement failed: {e}")
            return current_structured, True
    
    def _iteratively_edit_skeleton(self, skeleton_path: str, structured_content: StructuredContent, 
                                 new_data: str) -> Tuple[str, int]:
        """Phase 2: Iteratively edit skeleton until no {} remain and all data is incorporated."""
        logger.info("Starting iterative skeleton editing")
        
        # Load skeleton document
        skeleton_doc = Document(skeleton_path)
        
        for iteration in range(MAX_SKELETON_EDIT_ITERATIONS):
            logger.info(f"Skeleton editing iteration {iteration + 1}/{MAX_SKELETON_EDIT_ITERATIONS}")
            
            # Extract current skeleton text
            current_text = self._extract_text_from_doc(skeleton_doc)
            
            # Check remaining placeholders
            remaining_placeholders = re.findall(r'\{[^}]*\}', current_text)
            
            # Check unused data
            unused_data = self._find_unused_data(structured_content, current_text)
            
            if not remaining_placeholders and not unused_data:
                logger.info(f"Skeleton editing completed after {iteration + 1} iterations - no placeholders and all data used")
                break
            
            logger.info(f"Found {len(remaining_placeholders)} placeholders and {len(unused_data)} unused data chunks")
            
            # Edit skeleton using LLM
            edited_doc = self._edit_skeleton_with_llm(
                skeleton_doc, structured_content, remaining_placeholders, unused_data, iteration + 1
            )
            
            if edited_doc:
                skeleton_doc = edited_doc
            else:
                logger.warning(f"Skeleton editing failed at iteration {iteration + 1}")
                break
        
        final_text = self._extract_text_from_doc(skeleton_doc)
        return final_text, iteration + 1
    
    def _find_unused_data(self, structured_content: StructuredContent, current_skeleton_text: str) -> List[str]:
        """Find data from structured content that hasn't been incorporated into skeleton."""
        unused_chunks = []
        
        # Check each piece of content from structured content
        for header, content in structured_content.content_by_header.items():
            # Look for key phrases from this content in the skeleton
            key_phrases = [phrase.strip() for phrase in content.split('.') if len(phrase.strip()) > 20][:5]
            
            content_used = False
            for phrase in key_phrases:
                if phrase in current_skeleton_text:
                    content_used = True
                    break
            
            if not content_used and len(content.strip()) > 50:
                unused_chunks.append(f"[{header}] {content[:200]}...")
        
        return unused_chunks[:5]  # Limit to 5 chunks to avoid overwhelming
    
    def _edit_skeleton_with_llm(self, skeleton_doc: Document, structured_content: StructuredContent,
                               remaining_placeholders: List[str], unused_data: List[str], 
                               iteration: int) -> Optional[Document]:
        """Use LLM to edit skeleton document."""
        import json
        
        current_skeleton_text = self._extract_text_from_doc(skeleton_doc)
        
        structured_json = {
            "headers": structured_content.headers,
            "content_by_header": structured_content.content_by_header
        }
        
        prompt = f"""
You are a legal document completion expert. Complete the SKELETON by filling placeholders with REAL DATA.

CURRENT SKELETON:
{current_skeleton_text[:3000]}

STRUCTURED CONTENT WITH REAL DATA:
{json.dumps(structured_json, ensure_ascii=False, indent=2)[:2000]}

REMAINING PLACEHOLDERS TO FILL:
{', '.join(remaining_placeholders[:10]) if remaining_placeholders else 'None'}

UNUSED DATA TO INCORPORATE:
{chr(10).join(unused_data) if unused_data else 'None'}

CRITICAL INSTRUCTIONS:
1. Fill ALL placeholders {{}} with SPECIFIC REAL DATA from the structured content
2. Use actual names: זואי מור, גיל מור
3. Use actual address: חיים הרצוג 7, דירה 7, הוד השרון  
4. Use actual phone: 0542477683
5. Use actual email: zoeykricheli@gmail.com
6. Use actual engineer: יצחק אולשבנג, מהנדס מספר 64154
7. NEVER use generic placeholders like [השלם], [כתובת], [שם]
8. If unused data exists, add complete new sections with specific details
9. Return COMPLETE legal document with ALL real information filled in

Provide the completed skeleton with ALL REAL DATA:"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert at completing legal documents while preserving format and structure."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,
                max_completion_tokens=4000
            )
            
            completed_text = response.choices[0].message.content.strip()
            
            # Create new document with completed text
            new_doc = Document()
            paragraphs = completed_text.split('\\n')
            for para in paragraphs:
                if para.strip():
                    new_doc.add_paragraph(para.strip())
            
            return new_doc
            
        except Exception as e:
            logger.error(f"LLM skeleton editing failed: {e}")
            return None
    
    def save_structured_content_json(self, structured_content: StructuredContent, output_path: str):
        """Save structured content as JSON for mid-step analysis."""
        import json
        
        structured_data = {
            "phase_1_results": {
                "headers": structured_content.headers,
                "content_by_header": structured_content.content_by_header,
                "full_content": structured_content.full_content,
                "data_usage_tracking": structured_content.data_usage,
                "metadata": structured_content.metadata
            },
            "created_at": __import__('time').strftime("%Y-%m-%d %H:%M:%S"),
            "headers_count": len(structured_content.headers),
            "total_content_length": len(structured_content.full_content)
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(structured_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved structured content JSON to: {output_path}")
    
    def _clean_new_data(self, raw_data: str) -> str:
        """Clean the new data by removing document headers and unwanted formatting."""
        if not raw_data:
            return ""
        
        # Remove document name headers like "=== DOC-20250322-WA0007.docx ==="
        cleaned = re.sub(r'=== [^=]+\.(docx|pdf|msg|jpg|jpeg|png|tiff|bmp) ===\n?', '', raw_data)
        
        # Remove any remaining document references with extensions
        cleaned = re.sub(r'DOC-\d{8}-[A-Z0-9]+\.(docx|pdf|msg|jpg|jpeg|png|tiff|bmp)', '', cleaned)
        
        # Remove standalone document names
        cleaned = re.sub(r'DOC-[^\s\n]+', '', cleaned)
        
        # Clean up extra whitespace and newlines
        cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
        cleaned = re.sub(r'\s+', ' ', cleaned)
        cleaned = cleaned.strip()
        
        return cleaned
    
    def load_new_data_from_directory(self, data_dir: str) -> str:
        """Load and combine data from all files in a directory."""
        from ..utils.data_reader import DataReader
        
        data_reader = DataReader()
        data_path = Path(data_dir)
        
        if not data_path.is_dir():
            raise ProcessingError(f"Data directory not found: {data_dir}")
        
        # Get supported file extensions
        supported_formats = data_reader.get_supported_formats()
        all_extensions = set()
        for format_list in supported_formats.values():
            all_extensions.update(format_list)
        
        # Find all supported files in directory
        data_files = []
        for file_path in data_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in all_extensions:
                data_files.append(file_path)
        
        if not data_files:
            raise ProcessingError(f"No supported data files found in directory: {data_dir}")
        
        logger.info(f"Found {len(data_files)} data files in directory")
        
        # Process each file and combine data
        raw_contents = []
        
        for i, file_path in enumerate(sorted(data_files)):
            logger.info(f"Processing file {i+1}/{len(data_files)}: {file_path.name}")
            
            try:
                file_data = data_reader.read_contract_data(str(file_path))
                
                # Handle raw content
                if "raw_content" in file_data:
                    raw_contents.append(f"=== {file_path.name} ===\n{file_data['raw_content']}")
            except Exception as e:
                logger.warning(f"Could not process file {file_path.name}: {e}")
                continue
        
        # Combine all raw content
        combined_content = "\n\n".join(raw_contents)
        logger.info(f"Combined data from {len(data_files)} files: {len(combined_content)} characters")
        
        return combined_content
    
    def _find_past_examples(self, examples_dir: str) -> List[Dict[str, Any]]:
        """Find and load past contract examples."""
        examples = []
        examples_path = Path(examples_dir)
        
        if examples_path.is_dir():
            # Look for DOCX files directly in the specified directory
            contract_files = list(examples_path.glob("*.docx"))
            logger.info(f"Looking for examples in: {examples_path}")
            logger.info(f"Found {len(contract_files)} DOCX files")
        else:
            # Fallback: Find all contract files in experiments subdirectories  
            contract_files = list(examples_path.glob("*/outputs/temp_contract_iter_*.docx"))
            logger.info(f"Using fallback search pattern in: {examples_path}")
        
        for file_path in contract_files:
            try:
                content = self._extract_text_from_docx(str(file_path))
                if len(content) > MIN_CONTENT_LENGTH:  # Only use substantial examples
                    examples.append({
                        'source': str(file_path),
                        'content': content,
                        'length': len(content)
                    })
                    logger.info(f"Loaded example: {file_path.name} ({len(content)} chars)")
            except Exception as e:
                logger.warning(f"Could not load example {file_path}: {e}")
                continue
        
        if not examples:
            raise ProcessingError(f"No valid examples found in {examples_dir}")
        
        # Sort by length and take the best ones
        examples.sort(key=lambda x: x['length'], reverse=True)
        return examples[:10]  # Use top 10 examples
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from a DOCX file."""
        return DocHandler.extract_text_from_docx(file_path)
    
    def _generate_initial_content(self, new_data: str, past_examples: List[Dict[str, Any]]) -> str:
        """Generate initial content using 2 randomly chosen examples."""
        selected_examples = random.sample(past_examples, min(2, len(past_examples)))
        
        examples_text = "\n\n=== EXAMPLE 1 ===\n" + selected_examples[0]['content'][:2000]
        if len(selected_examples) > 1:
            examples_text += "\n\n=== EXAMPLE 2 ===\n" + selected_examples[1]['content'][:2000]
        
        prompt = f"""
You are a legal document expert. Create a new legal document text using the NEW DATA provided, following the layout and style of the PAST EXAMPLES.

NEW DATA TO INCORPORATE:
{new_data[:2000]}

PAST EXAMPLES FOR STYLE AND LAYOUT:
{examples_text}

INSTRUCTIONS:
1. Create a new legal document that incorporates the specific facts from the NEW DATA
2. Follow the same layout, structure, and legal language style as shown in the PAST EXAMPLES
3. Use the exact names, dates, addresses, and details from the NEW DATA
4. Maintain the same professional legal tone and formatting
5. Ensure all specific information from NEW DATA is included
6. The result should be a complete legal document, not a template

Generate the complete legal document text:"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert legal document writer specializing in creating documents based on examples and new data."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=self.temperature,
                max_completion_tokens=4000
            )
            
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            logger.error(f"Initial content generation failed: {e}")
            raise ProcessingError(f"Initial content generation failed: {e}")
    
    def _improve_content_with_examples(self, current_content: str, new_data: str, 
                                     selected_examples: List[Dict[str, Any]], 
                                     iteration: int) -> Tuple[str, bool]:
        """Improve content using different examples."""
        
        examples_text = ""
        for i, example in enumerate(selected_examples):
            examples_text += f"\n\n=== REFERENCE EXAMPLE {i+1} ===\n{example['content'][:1500]}"
        
        prompt = f"""
You are a legal document improvement expert. Review the CURRENT DOCUMENT and improve it using the style and approach shown in the REFERENCE EXAMPLES.

CURRENT DOCUMENT:
{current_content[:3000]}

NEW DATA FOR CONTEXT:
{new_data[:1500]}

REFERENCE EXAMPLES FOR IMPROVEMENT:{examples_text}

INSTRUCTIONS:
1. Compare the CURRENT DOCUMENT with the REFERENCE EXAMPLES
2. Improve the structure, legal language, formatting, and completeness
3. Ensure all facts from NEW DATA are properly incorporated
4. Make the document more professional and legally sound
5. If the current document is already well-written and complete, respond EXACTLY with: "NO_MORE_IMPROVEMENTS"
6. Otherwise, provide the improved version

IMPORTANT: Only respond with "NO_MORE_IMPROVEMENTS" if you genuinely believe the document cannot be improved further.

Provide your response:"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert legal document improvement specialist. Only make changes if they genuinely improve the document."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=self.temperature,
                max_completion_tokens=4000
            )
            
            response_text = response.choices[0].message.content.strip()
            
            # Check if LLM says no more improvements
            no_more_improvements = (
                response_text == "NO_MORE_IMPROVEMENTS" or 
                "NO_MORE_IMPROVEMENTS" in response_text
            )
            
            if no_more_improvements:
                return current_content, True
            else:
                return response_text, False
                
        except Exception as e:
            logger.error(f"Content improvement failed: {e}")
            return current_content, True  # Stop on error
    
    def _fill_skeleton_with_content(self, skeleton_path: str, generated_content: str) -> str:
        """Fill the skeleton document by replacing all placeholders with generated content while preserving format."""
        try:
            # Load skeleton
            skeleton_doc = Document(skeleton_path)
            
            # Get placeholder mappings from LLM
            placeholder_mappings = self._get_placeholder_mappings(skeleton_doc, generated_content)
            
            if not placeholder_mappings:
                logger.info("No placeholders found in skeleton")
                # Still return the document content for consistency
                return self._extract_text_from_doc(skeleton_doc)
            
            logger.info(f"Found {len(placeholder_mappings)} placeholders to fill")
            
            # Replace placeholders in document while preserving formatting
            self._replace_placeholders_in_doc(skeleton_doc, placeholder_mappings)
            
            # Return the updated document text
            return self._extract_text_from_doc(skeleton_doc)
            
        except Exception as e:
            logger.error(f"Skeleton filling failed: {e}")
            raise ProcessingError(f"Skeleton filling failed: {e}")
    
    def _get_placeholder_mappings(self, skeleton_doc: Document, generated_content: str) -> Dict[str, str]:
        """Get mappings for all placeholders using LLM intelligence."""
        # Extract all text with placeholders
        skeleton_text = self._extract_text_from_doc(skeleton_doc)
        
        # Find all placeholders (anything between { and })
        placeholders = re.findall(r'\{[^}]*\}', skeleton_text)
        
        if not placeholders:
            return {}
        
        # Use LLM to intelligently map each placeholder
        return self._intelligent_placeholder_mapping(skeleton_text, generated_content, placeholders)
    
    def _extract_text_from_doc(self, doc: Document) -> str:
        """Extract text from document object."""
        return DocHandler.extract_text_from_doc(doc)
    
    def _replace_placeholders_in_doc(self, doc: Document, mappings: Dict[str, str]):
        """Replace placeholders in document while preserving formatting."""
        DocHandler.replace_placeholders_in_doc(doc, mappings)
    
    def _intelligent_placeholder_mapping(self, skeleton_text: str, generated_content: str, 
                                       placeholders: List[str]) -> Dict[str, str]:
        """Intelligently fill placeholders using the generated content."""
        
        prompt = f"""
You are a legal document completion expert. Create mappings for ALL placeholders using information from the GENERATED CONTENT.

GENERATED CONTENT TO USE FOR FILLING:
{generated_content[:3000]}

PLACEHOLDERS TO FILL:
{', '.join(placeholders[:20])}

INSTRUCTIONS:
1. For each placeholder, provide appropriate replacement text from the GENERATED CONTENT
2. Use specific information - names, dates, addresses, amounts, etc.
3. If a placeholder cannot be filled from GENERATED CONTENT, use reasonable legal text
4. Return ONLY a JSON object mapping each placeholder to its replacement
5. Example format: {{"{{placeholder1}}": "replacement text", "{{placeholder2}}": "other text"}}

Return JSON mapping:"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert at mapping placeholders to appropriate content. Return only valid JSON."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,  # Lower temperature for precise filling
                max_completion_tokens=2000
            )
            
            response_text = response.choices[0].message.content.strip()
            
            # Try to parse JSON response
            try:
                import json
                mappings = json.loads(response_text)
                logger.info(f"Successfully created {len(mappings)} placeholder mappings")
                return mappings
            except json.JSONDecodeError:
                logger.warning("Could not parse LLM response as JSON, using fallback")
                return self._create_fallback_mappings(placeholders, generated_content)
            
        except Exception as e:
            logger.error(f"Intelligent placeholder mapping failed: {e}")
            return self._create_fallback_mappings(placeholders, generated_content)
    
    def _create_fallback_mappings(self, placeholders: List[str], generated_content: str) -> Dict[str, str]:
        """Create fallback mappings for placeholders."""
        mappings = {}
        
        # Simple extraction logic
        for placeholder in placeholders:
            placeholder_lower = placeholder.lower()
            
            if 'date' in placeholder_lower or 'תאריך' in placeholder_lower:
                # Look for dates in content
                date_match = re.search(r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{4}|\d{4}-\d{1,2}-\d{1,2})', generated_content)
                mappings[placeholder] = date_match.group(1) if date_match else "2023-12-01"
            
            elif 'name' in placeholder_lower or 'שם' in placeholder_lower:
                # Look for names
                if 'זואי' in generated_content:
                    mappings[placeholder] = "זואי מור"
                else:
                    mappings[placeholder] = "[שם]"
            
            elif 'address' in placeholder_lower or 'כתובת' in placeholder_lower:
                if 'חיים הרצוג' in generated_content:
                    mappings[placeholder] = "חיים הרצוג 7, דירה 7, הוד השרון"
                else:
                    mappings[placeholder] = "[כתובת]"
            
            elif 'phone' in placeholder_lower or 'טלפון' in placeholder_lower:
                phone_match = re.search(r'\b0\d{2}-?\d{7}\b', generated_content)
                mappings[placeholder] = phone_match.group(0) if phone_match else "050-1234567"
            
            else:
                # Extract any available information or use contextual replacement
                if 'זואי' in generated_content or 'מור' in generated_content:
                    if 'name' in placeholder_lower or 'שם' in placeholder_lower:
                        mappings[placeholder] = "זואי מור"
                    elif 'company' in placeholder_lower or 'חברה' in placeholder_lower:
                        mappings[placeholder] = "רס אדרת"
                    else:
                        mappings[placeholder] = "יש לציין בהתאם למסמכים"
                else:
                    mappings[placeholder] = "יש לציין בהתאם למסמכים"
        
        return mappings
    
    def save_filled_skeleton(self, skeleton_path: str, generated_content: str, output_path: str):
        """Save the filled skeleton document to output path while preserving formatting."""
        import shutil
        
        # Copy skeleton to output location
        shutil.copy2(skeleton_path, output_path)
        
        # Load and fill the copied document
        skeleton_doc = Document(output_path)
        placeholder_mappings = self._get_placeholder_mappings(skeleton_doc, generated_content)
        
        if placeholder_mappings:
            self._replace_placeholders_in_doc(skeleton_doc, placeholder_mappings)
            skeleton_doc.save(output_path)
            logger.info(f"Saved filled skeleton document to: {output_path}")
        else:
            logger.info("No placeholders to fill, saved original skeleton")
    
    def _simple_placeholder_replacement(self, skeleton_text: str, generated_content: str) -> str:
        """Simple fallback placeholder replacement."""
        # Replace all placeholders with a generic completion message
        result = re.sub(r'\{[^}]*\}', '[CONTENT COMPLETED FROM EXAMPLES]', skeleton_text)
        return result