"""Document processor for handling .docx contract skeletons."""

import re
import copy
from pathlib import Path
from typing import Dict, List, Tuple, Any, Optional
from dataclasses import dataclass
import logging

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

logger = logging.getLogger(__name__)


@dataclass
class PlaceholderInfo:
    """Information about a placeholder found in the document."""
    text: str
    field_name: str
    paragraph: Optional[Paragraph] = None
    table_cell: Optional[_Cell] = None
    run_index: Optional[int] = None


class DocumentProcessor:
    """Handles .docx document processing for contract generation."""
    
    def __init__(self):
        # Regex patterns for different placeholder formats
        self.placeholder_patterns = [
            r'\{\{(\w+)\}\}',  # {{field_name}}
            r'\[([A-Z_]+)\]',  # [FIELD_NAME]
            r'\{([^}]+)\}',    # {field_name} - enhanced to handle dots, spaces, special chars
            r'<(\w+)>',        # <field_name>
        ]
    
    def create_document_copy(self, skeleton_path: str, output_path: str) -> Document:
        """Create a copy of the skeleton document.
        
        Args:
            skeleton_path: Path to the skeleton .docx file
            output_path: Path where the copy should be saved
            
        Returns:
            Document object of the copied file
        """
        try:
            # Load the skeleton document
            skeleton_doc = Document(skeleton_path)
            
            # Save as copy to preserve original
            skeleton_doc.save(output_path)
            
            # Load the copy for manipulation
            copy_doc = Document(output_path)
            
            logger.info(f"Created document copy from {skeleton_path} to {output_path}")
            return copy_doc
            
        except Exception as e:
            logger.error(f"Error creating document copy: {str(e)}")
            raise
    
    def find_placeholders(self, document: Document) -> List[PlaceholderInfo]:
        """Find all placeholders in the document.
        
        Args:
            document: Document object to search
            
        Returns:
            List of PlaceholderInfo objects containing placeholder details
        """
        placeholders = []
        
        # Search in paragraphs
        for paragraph in document.paragraphs:
            placeholders.extend(self._find_placeholders_in_paragraph(paragraph))
        
        # Search in tables
        for table in document.tables:
            placeholders.extend(self._find_placeholders_in_table(table))
        
        # Search in headers and footers
        for section in document.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    placeholders.extend(self._find_placeholders_in_paragraph(paragraph))
            
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    placeholders.extend(self._find_placeholders_in_paragraph(paragraph))
        
        logger.info(f"Found {len(placeholders)} placeholders in document")
        return placeholders
    
    def _find_placeholders_in_paragraph(self, paragraph: Paragraph) -> List[PlaceholderInfo]:
        """Find placeholders in a paragraph."""
        placeholders = []
        text = paragraph.text
        
        for pattern in self.placeholder_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                placeholder_info = PlaceholderInfo(
                    text=match.group(0),
                    field_name=match.group(1),
                    paragraph=paragraph
                )
                placeholders.append(placeholder_info)
        
        return placeholders
    
    def _find_placeholders_in_table(self, table: Table) -> List[PlaceholderInfo]:
        """Find placeholders in table cells."""
        placeholders = []
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    cell_placeholders = self._find_placeholders_in_paragraph(paragraph)
                    # Update each placeholder to include cell information
                    for placeholder in cell_placeholders:
                        placeholder.table_cell = cell
                    placeholders.extend(cell_placeholders)
        
        return placeholders
    
    def replace_placeholders(self, document: Document, data: Dict[str, Any]) -> Document:
        """Replace placeholders in the document with actual data.
        
        Args:
            document: Document object to process
            data: Dictionary mapping field names to values
            
        Returns:
            Document with placeholders replaced
        """
        placeholders = self.find_placeholders(document)
        replaced_count = 0
        
        for placeholder in placeholders:
            field_name = placeholder.field_name.lower()
            
            # Look for matching data (case-insensitive)
            replacement_value = None
            for key, value in data.items():
                if key.lower() == field_name:
                    replacement_value = str(value) if value is not None else ""
                    break
            
            if replacement_value is not None:
                if placeholder.paragraph:
                    self._replace_in_paragraph(placeholder.paragraph, placeholder.text, replacement_value)
                    replaced_count += 1
                else:
                    logger.warning(f"Could not replace placeholder {placeholder.text} - no paragraph reference")
            else:
                logger.warning(f"No data found for placeholder {placeholder.text}")
        
        logger.info(f"Replaced {replaced_count} placeholders in document")
        return document
    
    def _replace_in_paragraph(self, paragraph: Paragraph, placeholder_text: str, replacement_value: str):
        """Replace placeholder in a specific paragraph while preserving formatting."""
        # Get the paragraph text
        full_text = paragraph.text
        
        if placeholder_text in full_text:
            # Clear the paragraph
            paragraph.clear()
            
            # Replace the placeholder in the text
            new_text = full_text.replace(placeholder_text, replacement_value)
            
            # Add the new text back to the paragraph
            paragraph.add_run(new_text)
    
    def validate_document_structure(self, document: Document) -> Dict[str, Any]:
        """Validate the document structure and content.
        
        Args:
            document: Document to validate
            
        Returns:
            Dictionary containing validation results
        """
        validation_results = {
            "is_valid": True,
            "issues": [],
            "statistics": {
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "section_count": len(document.sections),
                "remaining_placeholders": 0
            }
        }
        
        # Check for remaining placeholders
        remaining_placeholders = self.find_placeholders(document)
        validation_results["statistics"]["remaining_placeholders"] = len(remaining_placeholders)
        
        if remaining_placeholders:
            validation_results["is_valid"] = False
            validation_results["issues"].append(
                f"Document contains {len(remaining_placeholders)} unfilled placeholders"
            )
        
        # Check for empty critical sections
        if len(document.paragraphs) == 0:
            validation_results["is_valid"] = False
            validation_results["issues"].append("Document has no paragraphs")
        
        # Check for proper document structure
        has_content = any(p.text.strip() for p in document.paragraphs)
        if not has_content:
            validation_results["is_valid"] = False
            validation_results["issues"].append("Document appears to be empty")
        
        return validation_results
    
    def extract_text_content(self, document: Document) -> str:
        """Extract all text content from the document.
        
        Args:
            document: Document to extract text from
            
        Returns:
            String containing all text content
        """
        content_parts = []
        
        # Extract from paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                content_parts.append(paragraph.text.strip())
        
        # Extract from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            content_parts.append(paragraph.text.strip())
        
        return "\n".join(content_parts)
    
    def get_document_metadata(self, document: Document) -> Dict[str, Any]:
        """Extract metadata from the document.
        
        Args:
            document: Document to analyze
            
        Returns:
            Dictionary containing document metadata
        """
        return {
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "section_count": len(document.sections),
            "word_count": len(self.extract_text_content(document).split()),
            "character_count": len(self.extract_text_content(document)),
            "has_headers": any(section.header for section in document.sections),
            "has_footers": any(section.footer for section in document.sections),
        }
