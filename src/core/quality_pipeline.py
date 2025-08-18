"""Quality assurance pipeline with quality gates for contract generation."""

import logging
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from enum import Enum
import time

from config.settings import settings
from .document_processor import DocumentProcessor
from .content_generator import IntelligentContentGenerator, GenerationContext, GenerationResult
from ..evaluation.metrics import MetricsCalculator, COMETEvaluator, EvaluationResult
from ..evaluation.llm_judge import LLMJudge
from ..utils.error_handler import ProcessingError, QualityGateError, error_handler
from ..utils.mlflow_tracker import MLflowTracker

logger = logging.getLogger(__name__)


class QualityGateStatus(Enum):
    """Status of quality gate evaluation."""
    PASSED = "passed"
    FAILED = "failed"
    WARNING = "warning"
    RETRY = "retry"


@dataclass
class QualityGateResult:
    """Result of a quality gate evaluation."""
    gate_name: str
    status: QualityGateStatus
    score: float
    threshold: float
    details: Dict[str, Any]
    recommendations: List[str]


@dataclass
class PipelineResult:
    """Complete result of the quality pipeline."""
    success: bool
    final_contract: str
    quality_scores: Dict[str, float]
    quality_gates: List[QualityGateResult]
    iterations: int
    total_time: float
    warnings: List[str]
    metadata: Dict[str, Any]


class QualityAssurancePipeline:
    """Multi-stage quality assurance pipeline with quality gates."""
    
    def __init__(self, enable_mlflow: bool = True):
        self.document_processor = DocumentProcessor()
        self.content_generator = IntelligentContentGenerator()
        self.metrics_calculator = MetricsCalculator()
        self.comet_evaluator = COMETEvaluator()
        self.llm_judge = LLMJudge()
        
        # MLflow tracking
        self.enable_mlflow = enable_mlflow
        self.mlflow_tracker = MLflowTracker() if enable_mlflow else None
        
        # Quality gate thresholds from settings
        self.quality_thresholds = {
            "bleu": settings.min_bleu_score,
            "rouge": settings.min_rouge_average,
            "meteor": settings.min_meteor_score,
            "comet": settings.min_comet_score,
            "llm_judge": settings.min_llm_judge_score,
            "redundancy": settings.max_redundancy_score,
            "completeness": settings.min_completeness_score
        }
        
        self.max_iterations = settings.max_regeneration_attempts
        self.quality_gates_enabled = settings.quality_gate_enabled
    
    def process_contract(self, skeleton_path: str, contract_data: Dict[str, Any],
                        checklist: Optional[List[str]] = None,
                        reference_contracts: Optional[List[str]] = None,
                        output_path: Optional[str] = None,
                        experiment_name: Optional[str] = None,
                        experiment_tags: Optional[Dict[str, str]] = None,
                        experiment_folder: Optional[str] = None) -> PipelineResult:
        """Process a contract through the complete quality pipeline.
        
        Args:
            skeleton_path: Path to the contract skeleton .docx file
            contract_data: Data to fill into the contract
            checklist: Optional checklist of requirements
            reference_contracts: Optional reference contracts for comparison
            output_path: Optional output path for the final contract
            experiment_name: Optional custom name for MLflow experiment run
            experiment_tags: Optional tags to add to the MLflow run
            
        Returns:
            PipelineResult containing the final contract and quality assessment
        """
        start_time = time.time()
        logger.info("Starting quality assurance pipeline")
        
        # Start MLflow tracking
        mlflow_run_id = None
        if self.enable_mlflow and self.mlflow_tracker:
            custom_name = experiment_name or f"pipeline_run_{int(start_time)}"
            mlflow_run_id = self.mlflow_tracker.start_contract_generation_run(
                contract_data, run_name=custom_name
            )
            # Add custom tags
            if experiment_tags:
                import mlflow
                for key, value in experiment_tags.items():
                    mlflow.set_tag(key, value)
        
        try:
            # Stage 1: Pre-generation validation
            pre_validation_result = self._pre_generation_validation(
                skeleton_path, contract_data, checklist
            )
            
            if not pre_validation_result["valid"]:
                error_message = f"Pre-generation validation failed: {'; '.join(pre_validation_result['errors'])}"
                raise ProcessingError(
                    error_message,
                    details={
                        "validation_errors": pre_validation_result["errors"],
                        "validation_warnings": pre_validation_result.get("warnings", [])
                    },
                    suggestions=[
                        "Check skeleton file exists and is valid",
                        "Ensure contract data contains required fields",
                        "Verify checklist format is correct"
                    ]
                )
            
            # Stage 2: Iterative generation with quality gates
            generation_result = self._iterative_generation_with_quality_gates(
                skeleton_path, contract_data, checklist, reference_contracts, experiment_folder
            )
            
            # Stage 3: Final validation and output preparation
            if generation_result.success and output_path:
                final_contract_path = self._prepare_final_output(
                    generation_result.final_contract, output_path
                )
                generation_result.metadata["output_path"] = final_contract_path
            
            total_time = time.time() - start_time
            generation_result.total_time = total_time
            
            # Log to MLflow
            if self.enable_mlflow and self.mlflow_tracker:
                self.mlflow_tracker.log_generation_metrics(
                    generation_result.quality_scores,
                    {"generation_time": total_time, "iterations": generation_result.iterations}
                )
                self.mlflow_tracker.log_quality_gates(generation_result.quality_gates)
                
                # Log artifacts if paths exist
                if output_path and final_contract_path:
                    self.mlflow_tracker.log_contract_artifacts(
                        final_contract_path, skeleton_path,
                        {"pipeline_result": generation_result.metadata}
                    )
                
                # Log code artifacts for reproducibility
                self.mlflow_tracker.log_code_artifacts()
                
                self.mlflow_tracker.end_run("FINISHED")
            
            logger.info(f"Pipeline completed in {total_time:.2f} seconds with {generation_result.iterations} iterations")
            return generation_result
            
        except (ProcessingError, QualityGateError) as e:
            # Log error to MLflow
            if self.enable_mlflow and self.mlflow_tracker:
                self.mlflow_tracker.log_error(str(e), "QualityPipelineError")
                self.mlflow_tracker.end_run("FAILED")
            
            # Re-raise custom errors to be handled by the API layer
            error_handler.handle_error(e, {"stage": "quality_pipeline", "skeleton_path": skeleton_path})
            raise
        except Exception as e:
            # Log error to MLflow
            if self.enable_mlflow and self.mlflow_tracker:
                self.mlflow_tracker.log_error(str(e), "UnexpectedError")
                self.mlflow_tracker.end_run("FAILED")
            
            logger.error(f"Unexpected pipeline error: {str(e)}")
            error_detail = error_handler.handle_error(e, {"stage": "quality_pipeline", "skeleton_path": skeleton_path})
            raise ProcessingError(
                f"Quality pipeline failed: {str(e)}",
                details={"original_error": str(e), "error_type": type(e).__name__},
                suggestions=[
                    "Check system logs for detailed error information",
                    "Verify all dependencies are properly installed",
                    "Contact support if the issue persists"
                ]
            )
    
    def _pre_generation_validation(self, skeleton_path: str, contract_data: Dict[str, Any],
                                 checklist: Optional[List[str]]) -> Dict[str, Any]:
        """Validate inputs before generation."""
        validation_result = {
            "valid": True,
            "errors": [],
            "warnings": []
        }
        
        # Validate skeleton file
        try:
            from docx import Document
            skeleton_doc = Document(skeleton_path)
            placeholders = self.document_processor.find_placeholders(skeleton_doc)
            
            if not placeholders:
                validation_result["warnings"].append("No placeholders found in skeleton document")
            
        except Exception as e:
            validation_result["valid"] = False
            validation_result["errors"].append(f"Invalid skeleton document: {str(e)}")
        
        # Validate contract data
        if not contract_data:
            validation_result["valid"] = False
            validation_result["errors"].append("No contract data provided")
        
        # Check for required fields
        required_fields = ["client_name", "provider_name"]
        for field in required_fields:
            if field not in contract_data or not contract_data[field]:
                validation_result["warnings"].append(f"Missing or empty required field: {field}")
        
        # Validate checklist
        if checklist and not isinstance(checklist, list):
            validation_result["errors"].append("Checklist must be a list of strings")
        
        return validation_result
    
    def _iterative_generation_with_quality_gates(self, skeleton_path: str,
                                               contract_data: Dict[str, Any],
                                               checklist: Optional[List[str]],
                                               reference_contracts: Optional[List[str]],
                                               experiment_folder: Optional[str] = None) -> PipelineResult:
        """Generate contract with iterative quality improvement."""
        quality_gates = []
        quality_scores = {}
        warnings = []
        
        # Extract skeleton text and create generation context
        from docx import Document
        skeleton_doc = Document(skeleton_path)
        skeleton_text = self.document_processor.extract_text_content(skeleton_doc)
        placeholders = [p.field_name for p in self.document_processor.find_placeholders(skeleton_doc)]
        
        context = GenerationContext(
            contract_type=contract_data.get("contract_type", "service_agreement"),
            skeleton_text=skeleton_text,
            placeholders=placeholders,
            contract_data=contract_data,
            checklist=checklist
        )
        
        iteration = 0
        best_contract = ""
        best_scores = {}
        
        while iteration < self.max_iterations:
            iteration += 1
            logger.info(f"Generation iteration {iteration}/{self.max_iterations}")
            
            # Generate contract content with iterative refinement
            generation_result = self.content_generator.generate_complete_contract(context)
            
            if not generation_result.success:
                warnings.extend(generation_result.warnings)
                continue
            
            # Create document copy and apply content in experiment folder
            if experiment_folder:
                output_path = str(experiment_folder / "outputs" / f"temp_contract_iter_{iteration}.docx")
            else:
                output_path = f"temp_contract_iter_{iteration}.docx"
            contract_doc = self.document_processor.create_document_copy(skeleton_path, output_path)
            final_doc = self.document_processor.replace_placeholders(contract_doc, generation_result.filled_placeholders)
            final_doc.save(output_path)
            
            # Extract final contract text
            final_contract_text = self.document_processor.extract_text_content(final_doc)
            
            # Run quality gates
            iteration_gates, iteration_scores = self._run_quality_gates(
                final_contract_text, reference_contracts, context
            )
            
            quality_gates.extend(iteration_gates)
            quality_scores.update(iteration_scores)
            
            # Check if quality gates pass
            if self._check_quality_gates_pass(iteration_gates):
                logger.info(f"Quality gates passed on iteration {iteration}")
                return PipelineResult(
                    success=True,
                    final_contract=final_contract_text,
                    quality_scores=iteration_scores,
                    quality_gates=quality_gates,
                    iterations=iteration,
                    total_time=0,  # Will be set by caller
                    warnings=warnings,
                    metadata={
                        "generation_metadata": generation_result.generation_metadata,
                        "final_iteration": iteration,
                        "temp_file": output_path
                    }
                )
            else:
                # Save best result so far
                overall_score = self._calculate_overall_quality_score(iteration_scores)
                best_overall_score = self._calculate_overall_quality_score(best_scores)
                
                if overall_score > best_overall_score:
                    best_contract = final_contract_text
                    best_scores = iteration_scores.copy()
                
                # Prepare feedback for next iteration
                failed_gates = [gate for gate in iteration_gates if gate.status == QualityGateStatus.FAILED]
                if failed_gates:
                    feedback = self._prepare_improvement_feedback(failed_gates)
                    context.checklist = (context.checklist or []) + feedback
        
        # If max iterations reached, return best result
        logger.warning(f"Quality gates not passed after {self.max_iterations} iterations")
        return PipelineResult(
            success=False,
            final_contract=best_contract,
            quality_scores=best_scores,
            quality_gates=quality_gates,
            iterations=iteration,
            total_time=0,  # Will be set by caller
            warnings=warnings + ["Quality gates not passed within maximum iterations"],
            metadata={"max_iterations_reached": True}
        )
    
    def _run_quality_gates(self, contract_text: str, reference_contracts: Optional[List[str]],
                          context: GenerationContext) -> Tuple[List[QualityGateResult], Dict[str, float]]:
        """Run all quality gates on the generated contract."""
        gates = []
        scores = {}
        
        if not self.quality_gates_enabled:
            logger.info("Quality gates disabled, skipping evaluation")
            return gates, scores
        
        # Gate 1: BLEU Score
        if reference_contracts:
            bleu_result = self.metrics_calculator.calculate_bleu_score(
                contract_text, reference_contracts, self.quality_thresholds["bleu"]
            )
            gates.append(self._create_quality_gate_result("BLEU", bleu_result))
            scores["bleu"] = bleu_result.score
        
        # Gate 2: ROUGE Scores
        if reference_contracts:
            rouge_result = self.metrics_calculator.calculate_rouge_scores(
                contract_text, reference_contracts, self.quality_thresholds["rouge"]
            )
            gates.append(self._create_quality_gate_result("ROUGE", rouge_result))
            scores["rouge"] = rouge_result.score
        
        # Gate 3: METEOR Score
        if reference_contracts:
            meteor_result = self.metrics_calculator.calculate_meteor_score(
                contract_text, reference_contracts, self.quality_thresholds["meteor"]
            )
            gates.append(self._create_quality_gate_result("METEOR", meteor_result))
            scores["meteor"] = meteor_result.score
        
        # Gate 4: COMET Score
        if reference_contracts:
            comet_result = self.comet_evaluator.calculate_comet_score(
                contract_text, reference_contracts, context.skeleton_text, self.quality_thresholds["comet"]
            )
            gates.append(self._create_quality_gate_result("COMET", comet_result))
            scores["comet"] = comet_result.score
        
        # Gate 5: Redundancy Check
        redundancy_result = self.metrics_calculator.calculate_redundancy_score(
            contract_text, self.quality_thresholds["redundancy"]
        )
        gates.append(self._create_quality_gate_result("Redundancy", redundancy_result))
        scores["redundancy"] = redundancy_result.score
        
        # Gate 6: Completeness Check
        required_elements = context.checklist or []
        completeness_result = self.metrics_calculator.calculate_completeness_score(
            contract_text, required_elements, self.quality_thresholds["completeness"]
        )
        gates.append(self._create_quality_gate_result("Completeness", completeness_result))
        scores["completeness"] = completeness_result.score
        
        # Gate 7: LLM Judge
        llm_judge_result = self.llm_judge.evaluate_contract(
            contract_text, context.contract_data, reference_contracts, self.quality_thresholds["llm_judge"]
        )
        gates.append(self._create_quality_gate_result("LLM_Judge", llm_judge_result))
        scores["llm_judge"] = llm_judge_result.score
        
        return gates, scores
    
    def _create_quality_gate_result(self, gate_name: str, evaluation_result: EvaluationResult) -> QualityGateResult:
        """Create a quality gate result from an evaluation result."""
        status = QualityGateStatus.PASSED if evaluation_result.passed_threshold else QualityGateStatus.FAILED
        
        recommendations = []
        if not evaluation_result.passed_threshold:
            recommendations.append(f"Score {evaluation_result.score:.3f} below threshold {evaluation_result.threshold}")
            if "recommendations" in evaluation_result.details:
                recommendations.extend(evaluation_result.details["recommendations"])
        
        return QualityGateResult(
            gate_name=gate_name,
            status=status,
            score=evaluation_result.score,
            threshold=evaluation_result.threshold,
            details=evaluation_result.details,
            recommendations=recommendations
        )
    
    def _check_quality_gates_pass(self, gates: List[QualityGateResult]) -> bool:
        """Check if all quality gates pass."""
        if not gates:
            return False
        
        # All gates must pass
        return all(gate.status == QualityGateStatus.PASSED for gate in gates)
    
    def _calculate_overall_quality_score(self, scores: Dict[str, float]) -> float:
        """Calculate overall quality score from individual metrics."""
        if not scores:
            return 0.0
        
        # Weighted average based on metric importance
        weights = {
            "bleu": 0.15,
            "rouge": 0.15,
            "meteor": 0.15,
            "comet": 0.15,
            "redundancy": 0.1,  # Lower is better for redundancy
            "completeness": 0.15,
            "llm_judge": 0.15
        }
        
        weighted_sum = 0.0
        total_weight = 0.0
        
        for metric, score in scores.items():
            if metric in weights:
                weight = weights[metric]
                # For redundancy, invert the score (lower is better)
                if metric == "redundancy":
                    score = 1.0 - score
                weighted_sum += score * weight
                total_weight += weight
        
        return weighted_sum / total_weight if total_weight > 0 else 0.0
    
    def _prepare_improvement_feedback(self, failed_gates: List[QualityGateResult]) -> List[str]:
        """Prepare feedback for improving the next iteration."""
        feedback = []
        
        for gate in failed_gates:
            feedback.append(f"Improve {gate.gate_name.lower()}: {gate.recommendations[0] if gate.recommendations else 'Quality below threshold'}")
        
        return feedback
    
    def _prepare_final_output(self, contract_text: str, output_path: str) -> str:
        """Prepare the final contract output."""
        try:
            # Here you would create the final .docx file
            # For now, just save as text
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(contract_text)
            
            logger.info(f"Final contract saved to {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error saving final contract: {str(e)}")
            return ""
    
    def evaluate_existing_contract(self, contract_path: str, 
                                 reference_contracts: Optional[List[str]] = None) -> Dict[str, Any]:
        """Evaluate an existing contract without regeneration.
        
        Args:
            contract_path: Path to the contract file to evaluate
            reference_contracts: Optional reference contracts for comparison
            
        Returns:
            Dictionary containing evaluation results
        """
        try:
            # Load contract
            from docx import Document
            contract_doc = Document(contract_path)
            contract_text = self.document_processor.extract_text_content(contract_doc)
            
            # Run evaluation metrics
            evaluation_results = {}
            
            if reference_contracts:
                bleu_result = self.metrics_calculator.calculate_bleu_score(
                    contract_text, reference_contracts
                )
                evaluation_results["bleu"] = bleu_result
                
                rouge_result = self.metrics_calculator.calculate_rouge_scores(
                    contract_text, reference_contracts
                )
                evaluation_results["rouge"] = rouge_result
                
                meteor_result = self.metrics_calculator.calculate_meteor_score(
                    contract_text, reference_contracts
                )
                evaluation_results["meteor"] = meteor_result
            
            redundancy_result = self.metrics_calculator.calculate_redundancy_score(contract_text)
            evaluation_results["redundancy"] = redundancy_result
            
            # LLM Judge evaluation
            contract_context = {"contract_type": "unknown"}  # Would extract from contract
            llm_judge_result = self.llm_judge.evaluate_contract(
                contract_text, contract_context, reference_contracts
            )
            evaluation_results["llm_judge"] = llm_judge_result
            
            return {
                "contract_path": contract_path,
                "evaluation_results": evaluation_results,
                "overall_score": self._calculate_overall_quality_score({
                    name: result.score for name, result in evaluation_results.items()
                })
            }
            
        except Exception as e:
            logger.error(f"Error evaluating existing contract: {str(e)}")
            return {"error": str(e)}
