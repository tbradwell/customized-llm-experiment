"""Delimiter formatting utilities for Hebrew RTL skeleton documents."""

import logging
from typing import Dict, List, Tuple, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class DelimiterFormatter:
    """Utility class for managing Hebrew RTL delimiters in skeleton documents."""
    
    # Hebrew RTL delimiters as per algorithm
    BLOCK_START = "%}"  # Opening delimiter (RTL)
    BLOCK_END = "{%"    # Closing delimiter (RTL)
    UNCERTAIN = "~}"    # Uncertain content delimiter
    
    def __init__(self):
        """Initialize delimiter formatter."""
        self.delimiter_positions = {}
        self.block_stack = []  # Track open blocks for validation
        
    def add_block_start_delimiter(self, doc: Document, position: int, 
                                cluster_id: int) -> None:
        """Add block start delimiter at specified position.
        
        Args:
            doc: Document to modify
            position: Paragraph position to add delimiter
            cluster_id: ID of the cluster this block represents
        """
        if position >= len(doc.paragraphs):
            logger.warning(f"Position {position} out of range, document has {len(doc.paragraphs)} paragraphs")
            return
        
        paragraph = doc.paragraphs[position]
        current_text = paragraph.text
        
        # Add delimiter before the paragraph content
        delimiter_text = f"{self.BLOCK_START}\n{current_text}"
        
        # Replace paragraph text while preserving formatting
        self._replace_paragraph_text_preserve_format(paragraph, delimiter_text)
        
        # Track delimiter
        self.delimiter_positions[position] = {
            'type': 'block_start',
            'cluster_id': cluster_id,
            'delimiter': self.BLOCK_START
        }
        
        # Track open block
        self.block_stack.append({
            'position': position,
            'cluster_id': cluster_id
        })
        
        logger.debug(f"Added block start delimiter at position {position} for cluster {cluster_id}")
    
    def add_block_end_delimiter(self, doc: Document, position: int) -> None:
        """Add block end delimiter at specified position.
        
        Args:
            doc: Document to modify
            position: Paragraph position to add delimiter
        """
        if position >= len(doc.paragraphs):
            logger.warning(f"Position {position} out of range, document has {len(doc.paragraphs)} paragraphs")
            return
        
        paragraph = doc.paragraphs[position]
        current_text = paragraph.text
        
        # Add delimiter after the paragraph content
        delimiter_text = f"{current_text}\n{self.BLOCK_END}"
        
        # Replace paragraph text while preserving formatting
        self._replace_paragraph_text_preserve_format(paragraph, delimiter_text)
        
        # Track delimiter
        self.delimiter_positions[position] = {
            'type': 'block_end',
            'delimiter': self.BLOCK_END
        }
        
        # Close block from stack
        if self.block_stack:
            closed_block = self.block_stack.pop()
            logger.debug(f"Closed block for cluster {closed_block['cluster_id']} at position {position}")
        else:
            logger.warning(f"Block end delimiter at position {position} has no matching start")
        
        logger.debug(f"Added block end delimiter at position {position}")
    
    def add_uncertain_delimiter(self, doc: Document, position: int) -> None:
        """Add uncertain content delimiter at specified position.
        
        Args:
            doc: Document to modify
            position: Paragraph position to add delimiter
        """
        if position >= len(doc.paragraphs):
            logger.warning(f"Position {position} out of range, document has {len(doc.paragraphs)} paragraphs")
            return
        
        paragraph = doc.paragraphs[position]
        current_text = paragraph.text
        
        # Add uncertain delimiter at the end of paragraph
        delimiter_text = f"{current_text} {self.UNCERTAIN}"
        
        # Replace paragraph text while preserving formatting
        self._replace_paragraph_text_preserve_format(paragraph, delimiter_text)
        
        # Track delimiter
        self.delimiter_positions[position] = {
            'type': 'uncertain',
            'delimiter': self.UNCERTAIN
        }
        
        logger.debug(f"Added uncertain delimiter at position {position}")
    
    def close_final_block(self, doc: Document) -> None:
        """Close any remaining open blocks at the end of the document.
        
        Args:
            doc: Document to modify
        """
        if self.block_stack:
            last_position = len(doc.paragraphs) - 1
            if last_position >= 0:
                # Close all remaining blocks at the end of the document
                remaining_blocks = len(self.block_stack)
                for _ in range(remaining_blocks):
                    self.add_block_end_delimiter(doc, last_position)
                
                logger.info(f"Closed {remaining_blocks} remaining open blocks")
            else:
                # No paragraphs in document, just clear the stack
                remaining_blocks = len(self.block_stack)
                self.block_stack.clear()
                logger.warning(f"Document has no paragraphs, cleared {remaining_blocks} unclosed blocks")
    
    def validate_delimiters(self) -> Tuple[bool, List[str]]:
        """Validate that all delimiters are properly paired.
        
        Returns:
            Tuple of (is_valid, list_of_issues)
        """
        issues = []
        
        # Check for unclosed blocks
        if self.block_stack:
            for block in self.block_stack:
                issues.append(f"Unclosed block at position {block['position']} for cluster {block['cluster_id']}")
        
        # Count delimiter types
        block_starts = sum(1 for pos_data in self.delimiter_positions.values() 
                          if pos_data['type'] == 'block_start')
        block_ends = sum(1 for pos_data in self.delimiter_positions.values() 
                        if pos_data['type'] == 'block_end')
        
        if block_starts != block_ends:
            issues.append(f"Unmatched delimiters: {block_starts} starts, {block_ends} ends")
        
        is_valid = len(issues) == 0
        
        if is_valid:
            logger.info("All delimiters are properly paired")
        else:
            logger.warning(f"Delimiter validation failed: {issues}")
        
        return is_valid, issues
    
    def get_delimiter_summary(self) -> Dict[str, int]:
        """Get summary of delimiter usage.
        
        Returns:
            Dictionary with delimiter counts
        """
        summary = {
            'block_starts': 0,
            'block_ends': 0,
            'uncertain': 0,
            'total_positions': len(self.delimiter_positions)
        }
        
        for pos_data in self.delimiter_positions.values():
            if pos_data['type'] == 'block_start':
                summary['block_starts'] += 1
            elif pos_data['type'] == 'block_end':
                summary['block_ends'] += 1
            elif pos_data['type'] == 'uncertain':
                summary['uncertain'] += 1
        
        return summary
    
    def _replace_paragraph_text_preserve_format(self, paragraph, new_text: str) -> None:
        """Replace paragraph text while preserving formatting.
        
        Args:
            paragraph: Paragraph object to modify
            new_text: New text content
        """
        # Save formatting from first run if available
        first_run_format = {}
        if paragraph.runs:
            try:
                first_run = paragraph.runs[0]
                first_run_format = {
                    'bold': first_run.bold,
                    'italic': first_run.italic,
                    'underline': first_run.underline,
                    'font_name': first_run.font.name,
                    'font_size': first_run.font.size,
                }
            except Exception:
                # If we can't read formatting, continue without it
                pass
        
        # Clear paragraph and add new text
        paragraph.clear()
        new_run = paragraph.add_run(new_text)
        
        # Restore formatting
        try:
            if first_run_format.get('bold') is not None:
                new_run.bold = first_run_format['bold']
            if first_run_format.get('italic') is not None:
                new_run.italic = first_run_format['italic']
            if first_run_format.get('underline') is not None:
                new_run.underline = first_run_format['underline']
            if first_run_format.get('font_name'):
                new_run.font.name = first_run_format['font_name']
            if first_run_format.get('font_size'):
                new_run.font.size = first_run_format['font_size']
        except Exception:
            # If formatting fails, at least we have the text
            pass
    
    def format_rtl_document(self, doc: Document) -> None:
        """Apply RTL formatting to the entire document.
        
        Args:
            doc: Document to format
        """
        try:
            # Set document direction to RTL if possible
            for paragraph in doc.paragraphs:
                try:
                    # Note: python-docx has limited RTL support
                    # This is a basic implementation
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                except Exception:
                    pass
            
            logger.info("Applied RTL formatting to document")
            
        except Exception as e:
            logger.warning(f"Failed to apply RTL formatting: {e}")
    
    def export_delimiter_map(self) -> Dict[int, Dict[str, str]]:
        """Export delimiter positions for external use.
        
        Returns:
            Dictionary mapping positions to delimiter information
        """
        return dict(self.delimiter_positions)
    
    def import_delimiter_map(self, delimiter_map: Dict[int, Dict[str, str]]) -> None:
        """Import delimiter positions from external source.
        
        Args:
            delimiter_map: Dictionary mapping positions to delimiter information
        """
        self.delimiter_positions = dict(delimiter_map)
        
        # Rebuild block stack for validation
        self.block_stack = []
        for pos, data in delimiter_map.items():
            if data['type'] == 'block_start':
                self.block_stack.append({
                    'position': pos,
                    'cluster_id': data.get('cluster_id', 0)
                })
        
        logger.info(f"Imported {len(delimiter_map)} delimiter positions")
