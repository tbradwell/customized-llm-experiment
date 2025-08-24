"""Skeleton generator for algorithm steps 4-9."""

import logging
from typing import List, Dict, Set, Tuple, Optional
import random
from docx import Document
import uuid
import os


from .models.paragraph import Paragraph
from .models.cluster import Cluster
from .models.skeleton_document import SkeletonDocument
from .utils.delimiter_formatter import DelimiterFormatter
from .repositories.skeleton_document_repository import SkeletonDocumentRepository
from ..utils.doc_handler import DocHandler

logger = logging.getLogger(__name__)


class SkeletonGenerator:
    """Handles skeleton document generation from clustered paragraphs."""
    
    def __init__(self, skeleton_document_repository: Optional[SkeletonDocumentRepository] = None):
        """Initialize skeleton generator.
        
        Args:
            skeleton_document_repository: Optional repository for skeleton document database operations
        """
        self.delimiter_formatter = DelimiterFormatter()
        self.doc_handler = DocHandler()
        self.skeleton_document_repository = skeleton_document_repository
        self.use_database = skeleton_document_repository is not None
        
    def generate_skeleton(self, paragraphs: List[Paragraph], 
                         clusters: Dict[int, Cluster]) -> SkeletonDocument:
        """Generate skeleton document according to algorithm steps 4-9.
        
        Args:
            paragraphs: List of processed paragraphs with cluster assignments
            clusters: Dictionary of clusters with representative members
            
        Returns:
            Generated skeleton document
        """
        logger.info(f"Generating skeleton from {len(paragraphs)} paragraphs and {len(clusters)} clusters")
        
        # Step 4: Take a random document
        source_document = self._select_source_document(paragraphs)
        
        if not source_document:
            logger.error("No source document found")
            raise ValueError("No source document available for skeleton generation")
        
        logger.info(f"Selected source document: {source_document}")
        
        # Get paragraphs from the selected document
        source_paragraphs = [p for p in paragraphs if p.document_id == source_document]
        source_paragraphs.sort(key=lambda x: x.absolute_position)
        
        logger.info(f"Processing {len(source_paragraphs)} paragraphs from source document")
        
        # Steps 5-9: Process paragraphs and generate skeleton
        skeleton_doc = self._generate_skeleton_document(source_paragraphs, clusters)
        
        return skeleton_doc
    
    def _select_source_document(self, paragraphs: List[Paragraph]) -> Optional[str]:
        """Select a random document as per algorithm step 4.
        
        Args:
            paragraphs: List of paragraphs
            
        Returns:
            Document ID of selected document
        """
        document_ids = list(set(p.document_id for p in paragraphs))
        
        if not document_ids:
            return None
        
        # Select random document
        selected_doc = random.choice(document_ids)
        logger.debug(f"Selected document {selected_doc} from {len(document_ids)} available documents")
        
        return selected_doc
    
    def _generate_skeleton_document(self, source_paragraphs: List[Paragraph], 
                                  clusters: Dict[int, Cluster]) -> SkeletonDocument:
        """Generate skeleton document from source paragraphs.
        
        Args:
            source_paragraphs: Paragraphs from selected source document
            clusters: Dictionary of clusters
            
        Returns:
            Generated skeleton document
        """
        # Create skeleton document object
        skeleton_doc = SkeletonDocument(
            id=str(uuid.uuid4()),
            source_document_ids=[source_paragraphs[0].document_id] if source_paragraphs else []
        )
        
        # Step 5: For each paragraph in document
        paragraph_to_replace = self._create_paragraph_replacement_structure(source_paragraphs, clusters)
        
        # Steps 6-9: Process replacements and create document
        template_doc = self._create_template_document(source_paragraphs, paragraph_to_replace)
        
        # Save template and update skeleton document
        output_dir = os.path.join(os.getcwd(), "skeleton_outputs")
        os.makedirs(output_dir, exist_ok=True)
        
        template_path = os.path.join(output_dir, f"skeleton_template_{skeleton_doc.id}.docx")
        template_doc.save(template_path)
        skeleton_doc.template_path = template_path
        
        logger.info(f"Skeleton template saved to: {template_path}")
        
        # Update statistics
        skeleton_doc.total_paragraphs_processed = len(source_paragraphs)
        skeleton_doc.total_clusters_found = len(clusters)
        skeleton_doc.delimiter_positions = self.delimiter_formatter.export_delimiter_map()
        
        # Save skeleton document to database if available
        if self.use_database:
            saved_skeleton_id = self.skeleton_document_repository.save_skeleton_document(skeleton_doc)
            logger.info(f"Saved skeleton document to database: {saved_skeleton_id}")
        
        logger.info(f"Generated skeleton document with {len(paragraph_to_replace)} replacements")
        
        return skeleton_doc
    
    def _create_paragraph_replacement_structure(self, paragraphs: List[Paragraph], 
                                              clusters: Dict[int, Cluster]) -> Dict[str, Dict]:
        """Create paragraph-to-replacement mapping structure (steps 5.1-5.9).
        
        Args:
            paragraphs: Source document paragraphs
            clusters: Dictionary of clusters
            
        Returns:
            Dictionary mapping paragraph IDs to replacement information
        """
        logger.info("Creating paragraph replacement structure")
        
        paragraph_to_replace = {}
        
        # Step 6: Initialize empty set: used_cluster_ids = set()
        # Step 7: Initialize previous_paragraph = None
        used_cluster_ids = set()
        previous_paragraph = None
        
        for paragraph in paragraphs:
            try:
                # Steps 5.1-5.2: Take clean version and embedding from db
                # (Already available in paragraph object)
                
                # Step 5.3: Find closest cluster to embedded paragraph
                # (Already done during clustering - cluster_id is assigned)
                
                if paragraph.cluster_id is None or paragraph.cluster_id not in clusters:
                    logger.warning(f"Paragraph {paragraph.id} has no valid cluster assignment")
                    continue
                
                cluster = clusters[paragraph.cluster_id]
                
                # Steps 5.4-5.6: Get representative members
                replacement_candidates = self._get_replacement_candidates(cluster, paragraphs)
                
                # Step 5.7: Assign block assignment based on previous paragraph
                block_assignment = self._assign_block_assignment(paragraph, previous_paragraph)
                paragraph.block_assignment = block_assignment
                paragraph.is_block = block_assignment is not None
                
                # Store replacement information
                paragraph_to_replace[paragraph.id] = {
                    'paragraph': paragraph,
                    'cluster_id': paragraph.cluster_id,
                    'block_assignment': block_assignment,
                    'certainty_type': paragraph.certainty_type,
                    'replacement_candidates': replacement_candidates,
                    'used_cluster_ids': used_cluster_ids.copy(),
                    'previous_paragraph': previous_paragraph
                }
                
                # Update tracking variables
                if paragraph.is_block and paragraph.cluster_id not in used_cluster_ids:
                    used_cluster_ids.add(paragraph.cluster_id)
                
                previous_paragraph = paragraph
                
                logger.debug(f"Added replacement for paragraph {paragraph.id}: "
                           f"cluster={paragraph.cluster_id}, block={block_assignment}, "
                           f"type={paragraph.certainty_type}")
                
            except Exception as e:
                logger.error(f"Failed to process paragraph {paragraph.id}: {e}")
                continue
        
        logger.info(f"Created replacement structure for {len(paragraph_to_replace)} paragraphs")
        return paragraph_to_replace
    
    def _get_replacement_candidates(self, cluster: Cluster, 
                                  all_paragraphs: List[Paragraph]) -> List[Dict[str, str]]:
        """Get replacement candidates for a cluster based on homogeneity score (step 5.6).
        
        Args:
            cluster: Cluster object with representative member IDs and homogeneity score
            all_paragraphs: All paragraphs to search for candidates
            
        Returns:
            List of replacement candidate dictionaries with text and source info
        """
        candidates = []
        
        # Create mapping from paragraph ID to paragraph object
        paragraph_map = {p.id: p for p in all_paragraphs}
        
        # Step 5.6: Check homogeneity score threshold (changed from 0.8 to 0.9)
        if cluster.homogeneity_score > 0.9:
            # High homogeneity (> 0.9): Take only medoid's original text
            if cluster.medoid_id and cluster.medoid_id in paragraph_map:
                medoid = paragraph_map[cluster.medoid_id]
                candidates.append({
                    'text': medoid.original_text,
                    'text_type': 'original',
                    'source_paragraph_id': medoid.id,
                    'source_type': 'medoid'
                })
                logger.debug(f"High homogeneity cluster {cluster.id} (score: {cluster.homogeneity_score:.3f}): using medoid original text only")
        else:
            # Low homogeneity (≤ 0.9): Take medoid and 2 farthest members' clean text
            
            # Add medoid clean text
            if cluster.medoid_id and cluster.medoid_id in paragraph_map:
                medoid = paragraph_map[cluster.medoid_id]
                candidates.append({
                    'text': medoid.clean_text or medoid.original_text,
                    'text_type': 'clean',
                    'source_paragraph_id': medoid.id,
                    'source_type': 'medoid'
                })
            
            # Add farthest members clean text
            for i, member_id in enumerate(cluster.farthest_member_ids[:2]):  # Take only first 2
                if member_id in paragraph_map:
                    farthest = paragraph_map[member_id]
                    candidates.append({
                        'text': farthest.clean_text or farthest.original_text,
                        'text_type': 'clean',
                        'source_paragraph_id': farthest.id,
                        'source_type': f'farthest_{i+1}'
                    })
            
            logger.debug(f"Low homogeneity cluster {cluster.id} (score: {cluster.homogeneity_score:.3f}): using medoid + {len(cluster.farthest_member_ids[:2])} farthest members clean text")
        
        logger.debug(f"Generated {len(candidates)} replacement candidates for cluster {cluster.id}")
        return candidates
    
    def _assign_block_assignment(self, current_paragraph: Paragraph, 
                               previous_paragraph: Optional[Paragraph]) -> Optional[str]:
        """Assign block assignment based on clustering (algorithm step 5.7).
        
        Args:
            current_paragraph: Current paragraph being processed
            previous_paragraph: Previous paragraph (or None)
            
        Returns:
            Block assignment string or None
        """
        if current_paragraph.cluster_id is None:
            return None
        
        # If current paragraph is mapped to same cluster as previous one, assign block
        if (previous_paragraph is not None and 
            previous_paragraph.cluster_id == current_paragraph.cluster_id):
            block_name = f"block_{current_paragraph.cluster_id}"
            logger.debug(f"Assigned block {block_name} to paragraph {current_paragraph.id}")
            return block_name
        
        return None
    
    def _create_template_document(self, source_paragraphs: List[Paragraph],
                                paragraph_to_replace: Dict[str, Dict]) -> Document:
        """Create template document with alternatives and delimiters (steps 8-9).
        
        Args:
            source_paragraphs: Original paragraphs from source document
            paragraph_to_replace: Replacement mapping structure
            
        Returns:
            Template document with alternatives
        """
        logger.info("Creating template document")
        
        # Create new document
        template_doc = Document()
        
        # Step 8: For each member in the paragraph_to_replace
        for paragraph in source_paragraphs:
            if paragraph.id not in paragraph_to_replace:
                # Add original paragraph if no replacement
                self._add_original_paragraph(template_doc, paragraph)
                continue
            
            replacement_info = paragraph_to_replace[paragraph.id]
            
            # Step 8.1: Clean the text in original document paragraph (already available in paragraph.clean_text)
            
            # Steps 8.2-8.4: Add delimiters based on conditions
            self._add_delimiters_to_document(template_doc, replacement_info)
            
            # Step 8.5: Add alternatives to document
            self._add_paragraph_alternatives(template_doc, replacement_info)
        
        # Step 9: Close final block if needed
        self.delimiter_formatter.close_final_block(template_doc)
        
        # Apply RTL formatting
        self.delimiter_formatter.format_rtl_document(template_doc)
        
        # Validate delimiters
        is_valid, issues = self.delimiter_formatter.validate_delimiters()
        if not is_valid:
            logger.warning(f"Delimiter validation issues: {issues}")
        
        logger.info("Template document created successfully")
        return template_doc
    
    def _add_delimiters_to_document(self, doc: Document, replacement_info: Dict) -> None:
        """Add delimiters to document based on paragraph type and state.
        
        Args:
            doc: Document to modify
            replacement_info: Replacement information for current paragraph
        """
        paragraph = replacement_info['paragraph']
        previous_paragraph = replacement_info['previous_paragraph']
        used_cluster_ids = replacement_info['used_cluster_ids']
        current_position = len(doc.paragraphs)
        
        # Step 8.2: If previous paragraph was block and conditions met
        if (previous_paragraph is not None and 
            previous_paragraph.is_block and 
            (not paragraph.is_block or previous_paragraph.cluster_id != paragraph.cluster_id)):
            
            self.delimiter_formatter.add_block_end_delimiter(doc, current_position - 1)
        
        # Step 8.3: If type is 'block' and cluster not seen before
        if (paragraph.is_block and 
            paragraph.cluster_id not in used_cluster_ids):
            
            self.delimiter_formatter.add_block_start_delimiter(doc, current_position, paragraph.cluster_id)
        
        # Step 8.4: If type is 'uncertain' and conditions met
        if (paragraph.certainty_type == 'uncertain' and 
            (not paragraph.is_block or paragraph.cluster_id not in used_cluster_ids)):
            
            self.delimiter_formatter.add_uncertain_delimiter(doc, current_position)
    
    def _add_paragraph_alternatives(self, doc: Document, replacement_info: Dict) -> None:
        """Add paragraph alternatives to document (step 8.5).
        
        Args:
            doc: Document to modify
            replacement_info: Replacement information
        """
        candidates = replacement_info['replacement_candidates']
        
        if not candidates:
            # Fallback: add original paragraph
            original_paragraph = replacement_info['paragraph']
            self._add_original_paragraph(doc, original_paragraph)
            return
        
        # Add each alternative as a separate paragraph
        for i, candidate in enumerate(candidates):
            # Extract text from candidate dictionary
            text_content = candidate['text']
            source_paragraph_id = candidate['source_paragraph_id']
            text_type = candidate['text_type']
            source_type = candidate['source_type']
            
            # Create paragraph
            para = doc.add_paragraph()
            
            # For now, add as simple text run
            # TODO: Preserve original font styling from source paragraph
            para.add_run(text_content)
            
            logger.debug(f"Added alternative {i+1} ({source_type}, {text_type} text) from paragraph {source_paragraph_id}")
    
    def _add_original_paragraph(self, doc: Document, paragraph: Paragraph) -> None:
        """Add original paragraph to document.
        
        Args:
            doc: Document to modify
            paragraph: Paragraph to add
        """
        para = doc.add_paragraph()
        
        # Use original text
        text_content = paragraph.original_text
        
        # Apply original formatting if available
        font_styles = paragraph.font_style.get('runs', [])
        
        if font_styles:
            for style in font_styles:
                run = para.add_run(style.get('text', ''))
                self._apply_font_style(run, style)
        else:
            para.add_run(text_content)
    
    def _apply_font_style(self, run, style: Dict) -> None:
        """Apply font styling to a run.
        
        Args:
            run: Document run object
            style: Style dictionary
        """
        try:
            if style.get('bold') is not None:
                run.bold = style['bold']
            if style.get('italic') is not None:
                run.italic = style['italic']
            if style.get('underline') is not None:
                run.underline = style['underline']
            if style.get('font_name'):
                run.font.name = style['font_name']
            if style.get('font_size'):
                run.font.size = style['font_size']
        except Exception as e:
            logger.warning(f"Failed to apply font style: {e}")
    
    def get_generation_stats(self, skeleton_doc: SkeletonDocument) -> Dict[str, any]:
        """Get statistics about skeleton generation.
        
        Args:
            skeleton_doc: Generated skeleton document
            
        Returns:
            Dictionary with generation statistics
        """
        delimiter_summary = self.delimiter_formatter.get_delimiter_summary()
        
        return {
            'skeleton_id': skeleton_doc.id,
            'source_documents': len(skeleton_doc.source_document_ids),
            'total_paragraphs_processed': skeleton_doc.total_paragraphs_processed,
            'total_clusters_found': skeleton_doc.total_clusters_found,
            'blocks_generated': skeleton_doc.blocks_generated,
            'uncertain_blocks': skeleton_doc.uncertain_blocks,
            'delimiter_summary': delimiter_summary,
            'template_path': skeleton_doc.template_path,
            'created_at': skeleton_doc.created_at.isoformat()
        }
    
    def validate_skeleton(self, skeleton_doc: SkeletonDocument) -> Tuple[bool, List[str]]:
        """Validate generated skeleton document.
        
        Args:
            skeleton_doc: Skeleton document to validate
            
        Returns:
            Tuple of (is_valid, list_of_issues)
        """
        issues = []
        
        # Check basic requirements
        if not skeleton_doc.id:
            issues.append("Skeleton document missing ID")
        
        if not skeleton_doc.template_path:
            issues.append("Skeleton document missing template path")
        
        if skeleton_doc.total_paragraphs_processed == 0:
            issues.append("No paragraphs were processed")
        
        # Validate delimiters
        delimiter_valid, delimiter_issues = self.delimiter_formatter.validate_delimiters()
        if not delimiter_valid:
            issues.extend(delimiter_issues)
        
        # Check template file exists
        try:
            import os
            if not os.path.exists(skeleton_doc.template_path):
                issues.append(f"Template file not found: {skeleton_doc.template_path}")
        except Exception:
            issues.append("Could not verify template file existence")
        
        is_valid = len(issues) == 0
        
        if is_valid:
            logger.info("Skeleton validation passed")
        else:
            logger.warning(f"Skeleton validation failed: {issues}")
        
        return is_valid, issues
