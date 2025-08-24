"""Main skeleton processor orchestrator."""

# External imports
import logging
import os
import time
from typing import List, Dict, Tuple, Any, Optional

from docx import Document

# Internal imports
from .document_preprocessor import DocumentPreprocessor
from .clustering_engine import ClusteringEngine
from .skeleton_generator import SkeletonGenerator
from .models.paragraph import Paragraph
from .models.cluster import Cluster
from .models.skeleton_document import SkeletonDocument
from .utils.embedding_client import EmbeddingClient
from .utils.text_cleaner import TextCleaner
from .repositories.document_repository import DocumentRepository
from .repositories.paragraph_repository import ParagraphRepository
from .repositories.cluster_repository import ClusterRepository
from .repositories.skeleton_document_repository import SkeletonDocumentRepository

# CONFIGURATION CONSTANTS
DEFAULT_EMBEDDING_MODEL = "text-embedding-3-small"
DEFAULT_HIGH_HOMOGENEITY_THRESHOLD = 0.9
DEFAULT_LOW_HOMOGENEITY_THRESHOLD = 0.9
DEFAULT_LANGUAGE = "he"
MINIMUM_DOCUMENTS_FOR_CLUSTERING = 2

logger = logging.getLogger(__name__)


class SkeletonProcessor:
    """Main orchestrator for the skeleton processing algorithm."""
    
    def __init__(self, 
                 openai_api_key: str,
                 embedding_model: str = DEFAULT_EMBEDDING_MODEL,
                 high_homogeneity_threshold: float = DEFAULT_HIGH_HOMOGENEITY_THRESHOLD,
                 low_homogeneity_threshold: float = DEFAULT_LOW_HOMOGENEITY_THRESHOLD,
                 language: str = DEFAULT_LANGUAGE,
                 database_url: Optional[str] = None):
        """Initialize skeleton processor with algorithm parameters."""
        self.embedding_model = embedding_model
        self.high_homogeneity_threshold = high_homogeneity_threshold
        self.low_homogeneity_threshold = low_homogeneity_threshold
        self.database_url = database_url
        self.use_database = database_url is not None
        
        # Initialize database repositories if database URL provided
        if self.use_database:
            logger.info(f"Initializing with database: {database_url.split('@')[1] if '@' in database_url else database_url}")
            self.document_repository = DocumentRepository(database_url)
            self.paragraph_repository = ParagraphRepository(database_url)
            self.cluster_repository = ClusterRepository(database_url)
            self.skeleton_document_repository = SkeletonDocumentRepository(database_url)
            
            # Validate database setup
            try:
                is_valid, issues = self.document_repository.validate_database_setup()
                if not is_valid:
                    logger.warning(f"Database validation issues: {issues}")
                else:
                    logger.info("Database validation passed")
            except Exception as e:
                logger.warning(f"Could not validate database setup: {e}")
        else:
            logger.info("Initializing without database (in-memory processing only)")
            self.document_repository = None
            self.paragraph_repository = None
            self.cluster_repository = None
            self.skeleton_document_repository = None
        
        # Initialize processing components
        self.embedding_client = EmbeddingClient(openai_api_key, embedding_model)
        self.text_cleaner = TextCleaner(language)
        
        self.document_preprocessor = DocumentPreprocessor(
            self.embedding_client, self.text_cleaner, self.document_repository, self.paragraph_repository
        )
        
        self.clustering_engine = ClusteringEngine(
            high_homogeneity_threshold, low_homogeneity_threshold, 
            self.cluster_repository, self.paragraph_repository
        )
        
        self.skeleton_generator = SkeletonGenerator(self.skeleton_document_repository)
        
        logger.info(f"Initialized SkeletonProcessor with model {embedding_model}, database: {self.use_database}")
        
        # Validate that we're using real OpenAI API
        self._validate_real_openai_setup(openai_api_key)
    
    def _validate_real_openai_setup(self, api_key: str) -> None:
        """Validate that we're using a real OpenAI API key, not mock data."""
        if not api_key:
            raise ValueError("OpenAI API key is required - no mock embeddings allowed")
        
        if api_key.lower().startswith(('test', 'mock', 'fake', 'dummy', 'sk-test')):
            raise ValueError(f"Mock/test API key detected: {api_key[:15]}... - only real OpenAI API keys allowed")
        
        if len(api_key) < 20:  # Real OpenAI keys are much longer
            raise ValueError(f"API key too short: {len(api_key)} characters - likely a test key")
        
        logger.info("✅ Real OpenAI API key validated - no mock embeddings will be used")
    
    def process_documents_to_skeleton(self, document_paths: List[str]) -> SkeletonDocument:
        """Process documents and generate skeleton template."""
        logger.info(f"Starting skeleton processing for {len(document_paths)} documents")
        start_time = time.time()
        
        try:
            # Steps 1.1-1.2.6: Document preprocessing
            logger.info("Phase 1: Document preprocessing and embedding creation")
            paragraphs = self.document_preprocessor.process_documents(document_paths)
            
            if not paragraphs:
                raise ValueError("No paragraphs extracted from documents")
            
            # Steps 2-3: Clustering and homogeneity scoring
            logger.info("Phase 2: Clustering and homogeneity analysis")
            clusters = self.clustering_engine.cluster_paragraphs(paragraphs)
            
            if not clusters:
                raise ValueError("No clusters created from paragraphs")
            
            # Assign certainty types based on homogeneity
            self.clustering_engine.assign_certainty_types(paragraphs, clusters)
            
            # Note: Paragraph cluster assignments are now updated directly in clustering engine
            
            # Steps 4-9: Skeleton generation
            logger.info("Phase 3: Skeleton document generation")
            skeleton_doc = self.skeleton_generator.generate_skeleton(paragraphs, clusters)
            
            # Calculate processing time
            processing_time = time.time() - start_time
            
            logger.info(f"Skeleton processing completed in {processing_time:.2f} seconds")
            
            return skeleton_doc
            
        except Exception as e:
            logger.error(f"Skeleton processing failed: {e}")
            raise
    
    def validate_input_documents(self, document_paths: List[str]) -> Tuple[bool, List[str]]:
        """Validate input documents before processing."""
        issues = []
        
        if not document_paths:
            issues.append("No document paths provided")
            return False, issues
        
        for path in document_paths:
            if not os.path.exists(path):
                issues.append(f"Document not found: {path}")
            elif not path.lower().endswith('.docx'):
                issues.append(f"Document is not DOCX format: {path}")
        
        if len(document_paths) < MINIMUM_DOCUMENTS_FOR_CLUSTERING:
            issues.append(f"At least {MINIMUM_DOCUMENTS_FOR_CLUSTERING} documents required for meaningful clustering")
        
        is_valid = len(issues) == 0
        return is_valid, issues
    
    def get_processing_pipeline_stats(self, document_paths: List[str], skeleton_doc: SkeletonDocument) -> Dict[str, Any]:
        """Get comprehensive statistics about the processing pipeline."""
        
        # Basic document statistics
        total_documents = len(document_paths)
        valid_documents = 0
        total_file_size = 0
        
        # Text content statistics
        total_paragraphs_in_docs = 0
        total_text_length = 0
        
        for doc_path in document_paths:
            try:
                if os.path.exists(doc_path):
                    valid_documents += 1
                    total_file_size += os.path.getsize(doc_path)
                    
                    # Extract text statistics
                    doc = Document(doc_path)
                    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                    total_paragraphs_in_docs += len(paragraphs)
                    total_text_length += sum(len(p) for p in paragraphs)
                    
                    # Include table text
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if paragraph.text.strip():
                                        total_text_length += len(paragraph.text.strip())
                        
            except Exception:
                continue
        
        # Processing efficiency statistics
        processing_rate = (skeleton_doc.total_paragraphs_processed / total_paragraphs_in_docs 
                          if total_paragraphs_in_docs > 0 else 0.0)
        
        cluster_efficiency = (skeleton_doc.total_clusters_found / skeleton_doc.total_paragraphs_processed 
                             if skeleton_doc.total_paragraphs_processed > 0 else 0.0)
        
        # Template statistics
        template_exists = os.path.exists(skeleton_doc.template_path) if skeleton_doc.template_path else False
        template_size = os.path.getsize(skeleton_doc.template_path) if template_exists else 0
        
        return {
            # Input document statistics
            'input_documents': {
                'total_count': total_documents,
                'valid_count': valid_documents,
                'total_file_size_bytes': total_file_size,
                'total_paragraphs': total_paragraphs_in_docs,
                'total_text_length': total_text_length
            },
            
            # Processing statistics
            'processing': {
                'paragraphs_processed': skeleton_doc.total_paragraphs_processed,
                'clusters_found': skeleton_doc.total_clusters_found,
                'processing_rate': round(processing_rate, 3),
                'cluster_efficiency': round(cluster_efficiency, 3)
            },
            
            # Output statistics
            'output': {
                'template_created': template_exists,
                'template_size_bytes': template_size,
                'content_blocks': skeleton_doc.blocks_generated,
                'uncertain_blocks': skeleton_doc.uncertain_blocks,
                'block_utilization': (skeleton_doc.blocks_generated / skeleton_doc.total_clusters_found 
                                     if skeleton_doc.total_clusters_found > 0 else 0.0)
            },
            
            # Metadata
            'metadata': {
                'skeleton_id': skeleton_doc.id,
                'algorithm_version': skeleton_doc.algorithm_version,
                'created_at': skeleton_doc.created_at.isoformat(),
                'embedding_model': self.embedding_model,
                'homogeneity_thresholds': {
                    'high': self.high_homogeneity_threshold,
                    'low': self.low_homogeneity_threshold
                }
            }
        }
